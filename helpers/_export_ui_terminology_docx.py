#!/usr/bin/env python3
"""
Export DAAB-Website-UI-Terminology-Reference.md to a branded Word document.

Run from repo root:
  python helpers/_export_ui_terminology_docx.py
"""
from __future__ import annotations

import sys
from pathlib import Path

try:
    from _paths import ROOT
except ImportError:
    ROOT = Path(__file__).resolve().parents[1]

from daab_docx_export import (  # noqa: E402
    add_code_block,
    add_formatted_run,
    add_image,
    export_markdown_to_docx,
)

MD_PATH = ROOT / "documents" / "DAAB-Website-UI-Terminology-Reference.md"
OUT_PATH = ROOT / "documents" / "docx" / "DAAB-Website-UI-Terminology-Reference.docx"

LAYOUT_WIREFRAME = """\
+-------------------------------------------------------------+
|  SKIP LINK (keyboard only)                                  |
+-------------------------------------------------------------+
|  TOP CHROME (#daab-top-chrome)                              |
|    [ Logo ]  Brand text     Home  Activities v  Scientists v |
|                             About v  Membership v  [Search] AZ/EN |
|  -----------------------------------------------------------|
|  Breadcrumbs: Home > Section > Current page                 |
+-------------------------------------------------------------+
|  HERO                                                       |
|  [ Page title (h1)              ] [ Hero summary panel     ]|
|  [ Page hero subtitle           ] [ .panel-card / lead text ]|
+-------------------------------------------------------------+
|  (optional) Section nav strip:  Tab | Tab | Tab | Tab       |
+-------------------------------------------------------------+
|  [ Sidebar widget  ]  |  Main content (#content)            |
|  - timeline / links  |  cards, text, forms, gallery         |
+-------------------------------------------------------------+
|  SITE FOOTER (.footer-pro)                                  |
+-------------------------------------------------------------+
|                                    [ Back-to-top ] (fixed)  |
+-------------------------------------------------------------+"""


def _append_visual_reference_page(docx_path: Path) -> None:
    """Re-open saved doc and append layout wireframe + optional logo images."""
    from docx import Document

    document = Document(str(docx_path))
    document.add_page_break()
    document.add_heading("Appendix A — Layout wireframe", level=1)
    intro = document.add_paragraph()
    add_formatted_run(
        intro,
        "ASCII overview of the standard page stack. "
        "Open any inner page at http://localhost:8010/ for a live reference.",
    )
    add_code_block(document, LAYOUT_WIREFRAME.splitlines())

    document.add_heading("Appendix B — Visual assets on site", level=1)
    note = document.add_paragraph()
    add_formatted_run(
        note,
        "These images appear in the live UI and help identify branding elements. "
        "SVG logos may not embed in all Word viewers; open the site if previews are missing.",
    )

    for rel_path, caption in (
        ("images/daab-logo.png", "DAAB / WAAS brand logo (.nav-brand-logo)"),
        ("images/didk-logo.svg", "DİDK partner logo (Forum hub)"),
        ("images/etn-logo.svg", "ETN partner logo (Forum hub)"),
    ):
        image_path = ROOT / rel_path
        if image_path.is_file():
            document.add_heading(caption, level=3)
            add_image(document, image_path, caption)

    document.save(str(docx_path))


def main() -> int:
    if not MD_PATH.is_file():
        print(f"Missing source: {MD_PATH.relative_to(ROOT)}", file=sys.stderr)
        return 1

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    export_markdown_to_docx(MD_PATH, OUT_PATH)
    _append_visual_reference_page(OUT_PATH)

    rel = OUT_PATH.relative_to(ROOT)
    print(f"OK  {rel}")
    print(
        "\nTip: Open in Word and press F9 to refresh the table of contents.",
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
