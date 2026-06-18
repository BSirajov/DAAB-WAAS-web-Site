#!/usr/bin/env python3
"""
Build publication-ready AZ and EN Word documents for the new daab-waas.com site.

Outputs:
  documents/docx/Yeni daab-waas.com saytı (AZ).docx
  documents/docx/Yeni daab-waas.com saytı (EN).docx

Run from repo root:
  python helpers/_format_yeni_site_docx.py
"""
from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

try:
    from _paths import ROOT
except ImportError:
    ROOT = Path(__file__).resolve().parents[1]

sys.path.insert(0, str(ROOT / "helpers"))
from daab_docx_export import (  # noqa: E402
    add_table_of_contents,
    add_title_page,
    set_update_fields_on_open,
    setup_header_footer,
    setup_styles,
)

OUT_AZ = ROOT / "documents" / "docx" / "Yeni daab-waas.com saytı (AZ).docx"
OUT_EN = ROOT / "documents" / "docx" / "Yeni daab-waas.com saytı (EN).docx"

DOC_TITLE_AZ = "Yeni daab-waas.com saytı"
DOC_TITLE_EN = "New daab-waas.com website"


def add_body_paragraphs(document: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        p = document.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)


def add_mono_section(document: Document, title: str, paragraphs: list[str]) -> None:
    document.add_heading(title, level=2)
    add_body_paragraphs(document, paragraphs)


INTRO_AZ = [
    "Hörmətli həmkarlar,",
    (
        "Son bir neçə həftə ərzində veb-saytımızın yenilənməsi üzərində çalışıram. "
        "Tələbələrimizin həyata keçirdiyi layihə hələ tamamlanmadığından bu işi onlardan "
        "asılı olmayaraq, paralel şəkildə özüm yerinə yetirirəm. Düşünürəm ki, saytımız daha "
        "müasir dizayn və funksionallıqlara malik olamaqla istiadəçilər üçün daha cəlbedici olmalıdır."
    ),
    (
        "Süni intellektin dəstəyindən istifadə etməklə mövcud saytımızı yeni "
        "https://daab-waas.com platformasına miqrasiya etmiş, onun strukturunu, dizaynını və "
        "məzmununu əsaslı şəkildə yeniləmişəm."
    ),
    (
        "Bununla yanaşı, DAAB-ın gələcək fəaliyyətində sponsorluq, ianə və tərəfdaşlıq "
        "dəstəyinin mühüm rol oynayacağını nəzərə alaraq, bu istiqamətdə ayrıca səhifələr də "
        "hazırlamışam. Hesab edirəm ki, potensial sponsorlar və tərəfdaşlarla əlaqə qurarkən "
        "yeni saytımızı, xüsusilə də sponsorluq və ianə dəstəyi ilə bağlı səhifələri onlarla "
        "paylaşmağımız vacibdir."
    ),
    (
        "Vaxtınız və imkanınız olarsa, lütfən yeni saytı nəzərdən keçirin. Hər hansı uyğunsuzluq, "
        "texniki səhv, məzmun çatışmazlığı və ya təklifiniz olarsa, mənə bildirin ki, lazımi "
        "düzəlişləri edim."
    ),
]

INTRO_EN = [
    "Dear colleagues,",
    (
        "Over the past few weeks I have been working on updating our website. Because the student "
        "project is not yet complete, I am carrying out this work in parallel and independently. "
        "I believe the site should be more attractive to users through a more modern design and "
        "broader functionality."
    ),
    (
        "With the support of artificial intelligence, I have migrated our existing site to the new "
        "https://daab-waas.com platform and substantially renewed its structure, design, and content."
    ),
    (
        "Bearing in mind that sponsorship, donations, and partnership support will play an important "
        "role in DAAB’s future activities, I have also prepared dedicated pages in this direction. "
        "I believe it is important to share our new site—especially the sponsorship and donation "
        "pages—with potential sponsors and partners when we reach out to them."
    ),
    (
        "If you have time, please review the new site. If you notice any inconsistency, technical "
        "error, missing content, or have suggestions, let me know so I can make the necessary "
        "corrections."
    ),
]

FEATURE_SECTIONS: list[dict] = [
    {
        "az_title": "Üzvlük Müraciəti",
        "en_title": "Membership Application",
        "az": [
            "Üzvlük Müraciəti funksionallığı DAAB-a qoşulmaq istəyən şəxslər üçün sadə, təhlükəsiz və tam onlayn müraciət prosesini təmin edir. Xüsusi hazırlanmış elektron müraciət forması vasitəsilə namizədlər şəxsi, akademik və peşəkar məlumatlarını, fəaliyyət sahələrini, yaşadıqları ölkə və şəhəri, əlaqə məlumatlarını və digər zəruri məlumatları təqdim edə bilərlər. Sistem bütün məlumatların strukturlaşdırılmış şəkildə toplanmasını təmin edərək üzvlük müraciətlərinin operativ və səmərəli qiymətləndirilməsinə şərait yaradır. Bu funksionallıq DAAB-ın dünyanın müxtəlif ölkələrində yaşayan azərbaycanlı alim və mütəxəssisləri vahid şəbəkədə birləşdirmək, beynəlxalq əməkdaşlığı genişləndirmək, elm, təhsil, innovasiya və cəmiyyətin inkişafına töhfə vermək missiyasına xidmət edir."
        ],
        "en": [
            "The Membership Application functionality provides a simple, secure, and fully online process for individuals who wish to join our association. Through a dedicated application form, prospective members can submit their personal, academic, and professional information, areas of expertise, country and city of residence, contact details, and other relevant credentials for review. The system streamlines the membership process by collecting all required information in a structured format, enabling efficient evaluation by the Association. This functionality supports WAAS’s mission of connecting Azerbaijani scientists and professionals worldwide, expanding its global network, and fostering collaboration in science, education, innovation, and societal development."
        ],
    },
    {
        "az_title": "Sponsorluq İmkanları",
        "en_title": "Sponsorship Opportunities",
        "az": [
            "Sponsorluq funksionallığı təşkilatlar, şirkətlər və fərdi dəstəkçilər üçün birliyimizin fəaliyyətlərinə və təşəbbüslərinə töhfə vermək imkanı yaradır. Bu bölmə vasitəsilə potensial sponsorlar təşkilatın həyata keçirdiyi layihələr, forumlar, elmi tədbirlər, təhsil proqramları və digər fəaliyyət istiqamətləri haqqında məlumat əldə edə, müxtəlif sponsorluq paketləri ilə tanış ola və əməkdaşlıq maraqlarını birbaşa DAAB-a çatdıra bilərlər. Sponsorluq mexanizmi elm, təhsil və innovasiyanın inkişafına dəstək verməklə yanaşı, sponsor təşkilatların sosial məsuliyyət təşəbbüslərini və ictimai nüfuzunu gücləndirməyə xidmət edir."
        ],
        "en": [
            "The Sponsorship functionality provides organizations, companies, and individual supporters with an opportunity to contribute to the activities and initiatives of our Association. Through this section, potential sponsors can learn about the projects, forums, scientific events, educational programs, and other activities carried out by the organization, explore various sponsorship packages, and express their interest in collaboration directly to WAAS. By supporting science, education, and innovation, the sponsorship mechanism not only contributes to the advancement of these fields but also helps sponsoring organizations strengthen their corporate social responsibility initiatives and enhance their public reputation."
        ],
    },
    {
        "az_title": "İanə Et",
        "en_title": "Donate",
        "az": [
            "İanə funksionallığı elmə, təhsilə və beynəlxalq elmi əməkdaşlığın inkişafına töhfə vermək istəyən fərdi şəxslər və təşkilatlar üçün nəzərdə tutulmuşdur. Bu bölmə vasitəsilə istifadəçilər DAAB-ın həyata keçirdiyi layihələrin, elmi təşəbbüslərin, gənc tədqiqatçıların dəstəklənməsi proqramlarının, beynəlxalq forumların və digər fəaliyyətlərin maliyyələşdirilməsinə könüllü şəkildə dəstək göstərə bilərlər. Toplanan vəsaitlər təşkilatın məqsəd və missiyasına uyğun şəkildə istifadə olunur və Azərbaycan elminin qlobal miqyasda inkişafına, eləcə də dünyanın müxtəlif ölkələrində yaşayan azərbaycanlı alimlərin birgə fəaliyyətinin gücləndirilməsinə xidmət edir."
        ],
        "en": [
            "The Donation functionality is designed for individuals and organizations who wish to contribute to the advancement of science, education, and international scientific cooperation. Through this section, users can voluntarily support the funding of WAAS’s projects, scientific initiatives, young researcher support programs, international forums, and other activities. The funds collected are utilized in accordance with the Association’s mission and objectives, contributing to the global development of Azerbaijani science and strengthening collaboration among Azerbaijani scientists living and working in different countries around the world."
        ],
    },
    {
        "az_title": "Axtarış və Filtrləmə İmkanları",
        "en_title": "Advanced Search and Filtering",
        "az": [
            "Saytın inkişaf etdirilmiş axtarış və filtrləmə funksionallığı istifadəçilərə geniş həcmdə məlumat arasında lazım olan məzmunu sürətli və rahat şəkildə tapmağa imkan verir. Məlumatlar müxtəlif meyarlar üzrə – ölkə, şəhər, fəaliyyət sahəsi, elmi istiqamət, təşkilat, tədbir və digər kateqoriyalar əsasında filtr edilə bilər. Bu imkan istifadəçilərin vaxtına qənaət etməklə yanaşı, onlara maraq dairələrinə uyğun məlumatları daha dəqiq və effektiv şəkildə əldə etməyə şərait yaradır."
        ],
        "en": [
            "The enhanced search and filtering functionality enables users to quickly and efficiently locate relevant information within the website’s extensive content repository. Information can be filtered using a variety of criteria, including country, city, field of expertise, scientific discipline, organization, event, and other categories. This capability helps users find the most relevant content with greater accuracy while significantly improving the overall browsing experience."
        ],
    },
    {
        "az_title": "Sıralama (Sorting) Funksionallığı",
        "en_title": "Sorting Capabilities",
        "az": [
            "Sıralama funksionallığı məlumatların istifadəçi ehtiyaclarına uyğun şəkildə təşkil edilməsini təmin edir. Siyahılar əlifba sırası, tarix, ölkə, fəaliyyət sahəsi və digər parametrlər üzrə artan və ya azalan qaydada sıralana bilər. Bu imkan xüsusilə böyük məlumat bazalarında naviqasiyanı asanlaşdırır və istifadəçilərin istədikləri məlumatlara daha operativ çıxışını təmin edir."
        ],
        "en": [
            "The sorting functionality allows users to organize information according to their specific preferences and needs. Content can be sorted alphabetically, by date, country, field of activity, and other available parameters in ascending or descending order. This feature is particularly valuable when navigating large datasets, making information discovery faster and more intuitive."
        ],
    },
    {
        "az_title": "Qruplaşdırma (Grouping) İmkanları",
        "en_title": "Grouping and Categorization",
        "az": [
            "Qruplaşdırma funksionallığı məlumatların məntiqi və strukturlaşdırılmış şəkildə təqdim olunmasına xidmət edir. Məsələn, alimlər ölkələr, fəaliyyət sahələri və ya elmi istiqamətlər üzrə qruplaşdırıla bilər. Bu yanaşma istifadəçilərə məlumatları daha sistemli şəkildə araşdırmağa, əlaqələri daha yaxşı görməyə və maraqlandıqları sahələr üzrə ümumi mənzərəni daha aydın şəkildə dərk etməyə imkan verir."
        ],
        "en": [
            "The grouping functionality presents information in a structured and logical manner by organizing related records into meaningful categories. For example, scientists can be grouped by country, scientific discipline, professional field, or other classification criteria. This approach helps users better understand relationships between records and gain a clearer overview of the available information."
        ],
    },
    {
        "az_title": "Təkmilləşdirilmiş Naviqasiya Sistemi",
        "en_title": "Enhanced Navigation System",
        "az": [
            "Yeni naviqasiya sistemi saytın bölmələri arasında daha sürətli və intuitiv keçidləri təmin edir. Təkmilləşdirilmiş menyular, alt menyular, səhifədaxili keçidlər və “çörək qırıntıları” (breadcrumbs) istifadəçilərə saytın strukturunda rahat istiqamətlənməyə kömək edir. Bu xüsusiyyət xüsusilə çoxsaylı bölmələr və geniş məzmun təqdim edən platformalar üçün əhəmiyyətlidir."
        ],
        "en": [
            "The redesigned navigation system provides a more intuitive and user-friendly experience across the website. Improved menus, hierarchical submenus, contextual navigation elements, and breadcrumb trails allow users to move seamlessly between sections and quickly understand their current location within the site structure. This significantly improves accessibility and ease of use."
        ],
    },
    {
        "az_title": "Kart və Siyahı Görünüşləri",
        "en_title": "Card and List Views",
        "az": [
            "İstifadəçilər məlumatları öz seçimlərinə uyğun olaraq kart (card) və ya siyahı (list) formatında görüntüləyə bilərlər. Kart görünüşü vizual təqdimata üstünlük verən istifadəçilər üçün daha rahatdır, siyahı görünüşü isə daha çox məlumatın eyni anda nəzərdən keçirilməsinə imkan yaradır. Bu çevik yanaşma fərqli istifadəçi ehtiyaclarını nəzərə alır və ümumi istifadəçi təcrübəsini yaxşılaşdırır."
        ],
        "en": [
            "Users can choose between card-based and list-based presentation modes according to their preferences. The card view offers a visually engaging layout that highlights key information, while the list view enables users to review larger amounts of information efficiently. This flexibility ensures an optimal experience for different user needs and browsing styles."
        ],
    },
    {
        "az_title": "Adaptiv və Mobil cihazlara uyğunlaşdırılmış Dizayn",
        "en_title": "Responsive and Mobile-Friendly Design",
        "az": [
            "Saytın bütün funksionallıqları müxtəlif cihazlarda – kompüter, planşet və mobil telefonlarda tam işlək şəkildə təqdim olunur. Adaptiv dizayn texnologiyaları ekran ölçüsündən asılı olmayaraq məzmunun düzgün göstərilməsini, rahat oxunmasını və funksiyaların problemsiz istifadəsini təmin edir. Bu xüsusiyyət istifadəçilərə istənilən vaxt və istənilən yerdən sayta rahat çıxış imkanı yaradır."
        ],
        "en": [
            "All website features have been designed and optimized to function seamlessly across desktops, tablets, and mobile devices. Responsive design principles ensure that content remains accessible, readable, and easy to navigate regardless of screen size. This allows users to access information and services anytime and from anywhere."
        ],
    },
    {
        "az_title": "Ikidilli Məzmun Dəstəyi",
        "en_title": "Multilingual Content Support",
        "az": [
            "Ikidilli dəstək funksionallığı istifadəçilərə məlumatları Azərbaycan və ingilis dillərində əldə etmək imkanı verir. Bu xüsusiyyət dünyanın müxtəlif ölkələrində yaşayan azərbaycanlı alimlər, tədqiqatçılar və tərəfdaş təşkilatlar üçün məlumatların daha əlçatan və anlaşıqlı olmasını təmin edir, beynəlxalq əməkdaşlığın inkişafına əlavə töhfə verir."
        ],
        "en": [
            "The multilingual functionality enables users to access information both in Azerbaijani and english languages, making the platform more accessible to Azerbaijani scientists, researchers, professionals, and partners worldwide. This capability supports international collaboration and ensures effective communication across diverse audiences."
        ],
    },
    {
        "az_title": "Məlumatların Strukturlaşdırılmış Təqdimatı",
        "en_title": "Structured Information Presentation",
        "az": [
            "Yeni texniki imkanlar sayəsində məlumatlar daha sistemli və oxunaqlı formada təqdim olunur. Bölmələr, kateqoriyalar, vizual kartlar, cədvəllər və əlaqəli keçidlər istifadəçilərə mürəkkəb məlumatları daha rahat mənimsəməyə və lazımi informasiyanı qısa müddətdə əldə etməyə kömək edir. Bu yanaşma saytın həm funksionallığını, həm də istifadəçi təcrübəsini əhəmiyyətli dərəcədə yaxşılaşdırır."
        ],
        "en": [
            "The website employs modern information architecture principles to present content in a clear, organized, and user-friendly format. Categories, sections, visual cards, tables, and contextual links help users understand complex information more easily and locate relevant content with minimal effort. This significantly enhances both usability and knowledge discovery."
        ],
    },
    {
        "az_title": "Alimlər kataloqunun təkmilləşdirilmiş idarə edilməsi",
        "en_title": "Advanced Scientist Directory Management",
        "az": [
            "Təkmilləşdirilmiş Alimlər kataloqu güclü axtarış, filtrləmə, sıralama və qruplaşdırma imkanları ilə istifadəçilərə azərbaycanlı alimlərin qlobal şəbəkəsini daha effektiv araşdırmağa imkan verir. Alimlər ölkə, fəaliyyət sahəsi, elmi istiqamət, təşkilat və digər meyarlar üzrə müəyyən edilə və kateqoriyalara ayrıla bilər. Bu funksionallıq şəbəkələşmə, əməkdaşlıq, bilik mübadiləsi və ölkələr və elmi sahələr üzrə yeni elmi tərəfdaşlıqların inkişafına kömək edir."
        ],
        "en": [
            "The enhanced Scientist Directory incorporates powerful search, filtering, sorting, and grouping capabilities that enable users to explore the global network of Azerbaijani scientists more effectively. Scientists can be identified and categorized by country, field of expertise, academic discipline, institution, and other criteria. This functionality facilitates networking, collaboration, knowledge exchange, and the development of new scientific partnerships across countries and disciplines."
        ],
    },
    {
        "az_title": "İnteraktiv məlumat araşdırması",
        "en_title": "Interactive Data Exploration",
        "az": [
            "Yeni interaktiv məlumat araşdırması funksiyaları istifadəçilərə təkmilləşdirilmiş filtrləmə, kateqoriyalaşdırma və naviqasiya vasitələri ilə məlumatları dinamik şəkildə nəzərdən keçirməyə və təhlil etməyə imkan verir. Bu imkanlar məzmunu kəşf etmək, əlaqələri müəyyən etmək və platforma üzrə müvafiq resurslara çıxış üçün daha fərdiləşdirilmiş və səmərəli yol təqdim edir."
        ],
        "en": [
            "The newly introduced interactive data exploration features allow users to dynamically browse and analyze information through advanced filtering, categorization, and navigation tools. These capabilities provide a more personalized and efficient way to discover content, identify connections, and access relevant resources throughout the platform."
        ],
    },
    {
        "az_title": "Genişlənə bilən məlumat idarəetməsi",
        "en_title": "Scalable Information Management",
        "az": [
            "Platforma davamlı böyümə və genişlənməni dəstəkləmək üçün nəzərdə tutulub. Genişlənə bilən arxitekturası yeni məzmun, alimlər, layihələr, tədbirlər, nəşrlər və digər resursların səmərəli şəkildə əlavə edilməsinə və təşkilinə imkan verir, eyni zamanda istifadəçilər üçün yüksək performans, ardıcıllıq və rahat naviqasiyanı qoruyur."
        ],
        "en": [
            "The platform has been designed to support continuous growth and expansion. Its scalable architecture allows new content, scientists, projects, events, publications, and other resources to be added and organized efficiently while maintaining high performance, consistency, and ease of navigation for users."
        ],
    },
    {
        "az_title": "Mətnin Səsləndirilməsi (Text-to-Speech) Funksionallığı",
        "en_title": "Text-to-Speech (Listening) Functionality",
        "az": [
            "Mətnin səsləndirilməsi funksionallığı istifadəçilərə sayt məzmununu oxumaq əvəzinə dinləmək imkanı yaradır və əlçatanlığı artırır. Bu funksiya vasitəsilə məqalələr, elanlar, profillər və digər mətnlər bir kliklə səsləndirilə bilər. Xüsusilə audio məzmuna üstünlük verən, görmə məhdudiyyəti olan və ya eyni vaxtda digər işlərlə məşğul olarkən məlumat əldə etmək istəyən istifadəçilər üçün faydalıdır. Nəzərə almaq lazımdır ki, bu funksionallığın işləməsi brauzer və cihaz dəstəyindən asılıdır və bəzi brauzerlərdə və ya mobil cihazlarda mövcud olmaya bilər."
        ],
        "en": [
            "The Text-to-Speech functionality allows users to listen to website content instead of reading it, improving accessibility and convenience. With this feature, visitors can hear articles, announcements, profiles, and other textual content directly from the website. It is especially useful for users who prefer audio content, have visual limitations, or access the site while multitasking. Please note that this functionality depends on browser and device support and may not be available on some browsers or mobile devices."
        ],
    },
    {
        "az_title": "QR Kod Funksionallığı",
        "en_title": "QR Code Functionality",
        "az": [
            "QR kod funksionallığı istifadəçilərin saytın məlumatlarına, resurslarına və xidmətlərinə daha sürətli və rahat çıxışını təmin etmək məqsədilə tətbiq edilmişdir. İstifadəçilər smartfon və ya planşet vasitəsilə QR kodu skan etməklə veb-səhifələrə, sənədlərə, tədbir məlumatlarına, üzvlük müraciəti formalarına, sponsorluq və ianə bölmələrinə, eləcə də digər əlaqəli resurslara dərhal keçid əldə edə bilərlər. Bu funksionallıq məlumat mübadiləsini sadələşdirir, əlçatanlığı artırır və xüsusilə forumlar, görüşlər, təqdimatlar və digər tədbirlər zamanı rəqəmsal resurslara operativ çıxışı təmin etməklə istifadəçi təcrübəsini yaxşılaşdırır."
        ],
        "en": [
            "The QR Code functionality has been introduced to provide users with faster and more convenient access to information, resources, and services available on the website. By simply scanning a QR code with a smartphone or tablet, users can instantly navigate to specific web pages, documents, event information, membership applications, sponsorship opportunities, donation pages, or other relevant content without the need to manually enter web addresses. This functionality enhances accessibility, simplifies information sharing, and improves the overall user experience, particularly during conferences, meetings, presentations, and promotional activities where quick access to digital resources is essential."
        ],
    },
]

OVERVIEW_SECTIONS: list[dict] = [
    {
        "az_title": "Yenilənmiş dizayn",
        "en_title": "Redesigned appearance",
        "az": [
            "Saytın görünüşü tamamilə yenidən işlənmişdir. Müasir naviqasiya paneli, açılan menyular, kart formatında hazırlanmış funksional səhifələr və yeni elektron poçt ünvanı — info@daab-waas.com — tətbiq edilmişdir.",
            "Səhifələrin başlıq (header) və ətək (footer) nahiyələri yeni dizaynla hazırlanmışdır. Ətək nahiyəsində (ən soldakı paneldə) saytın QR kodu yerləşdirilmişdir. Bu kod vasitəsilə sayt istənilən cihazdan, smartfon və ya planşetdən, birbaşa açıla bilər.",
            "Cədvəllərdə sücgəc (filter), sıralama (sorting), qruplaşdirma (groupping) və s. kimi funksiyalar da əlavə edilmişdir.",
        ],
        "en": [
            "The site’s appearance has been completely reworked. A modern navigation bar, dropdown menus, functional pages in card layout, and the new email address info@daab-waas.com have been introduced.",
            "Page headers and footers use the new design. The site QR code is placed in the footer (left panel), allowing the site to be opened directly from any device, smartphone, or tablet.",
            "Tables also include filter, sorting, grouping, and similar functions.",
        ],
    },
    {
        "az_title": "İkidillilik",
        "en_title": "Bilingual site",
        "az": [
            "Veb-sayt Azərbaycan və ingilis dillərində tam ikidilli formatda hazırlanmışdır. Bu, saytın həm diaspor ictimaiyyəti, həm də Azərbaycan dilini bilməyən beynəlxalq istifadəçilər üçün daha əlçatan olmasını təmin edir."
        ],
        "en": [
            "The website is fully bilingual in Azerbaijani and English, making it more accessible both to the diaspora community and to international users who do not speak Azerbaijani."
        ],
    },
    {
        "az_title": "Sponsorluq və tərəfdaşlıq",
        "en_title": "Sponsorship and partnership",
        "az": [
            "Potensial sponsorlar və institusional tərəfdaşlar üçün ayrıca bölmə yaradılmışdır. Bu bölmədə müəssisələrin, fondların və təşkilatların DAAB-ın missiyasını necə dəstəkləyə biləcəyi ətraflı şəkildə izah olunur. Buraya forumların, təhsil proqramlarının, elmi layihələrin və digər təşəbbüslərin maliyyələşdirilməsi, eləcə də birgə tərəfdaşlıq imkanları daxildir."
        ],
        "en": [
            "A dedicated section has been created for potential sponsors and institutional partners. It explains in detail how companies, foundations, and organizations can support DAAB’s mission, including funding forums, education programs, scientific projects, and other initiatives, as well as opportunities for joint partnership."
        ],
    },
    {
        "az_title": "Sayt üzrə qlobal axtarış",
        "en_title": "Site-wide global search",
        "az": [
            "İstifadəçilər saytın istənilən səhifəsindən açar söz yazaraq bütün məzmun üzrə axtarış apara bilərlər. Axtarış üzv profilləri, xəbərlər, tədbirlər, kataloqlar və digər bölmələri əhatə edir. Nəticələr səhifə yenilənmədən dərhal ekranda göstərilir."
        ],
        "en": [
            "Users can search across all site content by entering keywords from any page. Search covers member profiles, news, events, catalogues, and other sections. Results appear on screen immediately without reloading the page."
        ],
    },
    {
        "az_title": "QR kodlar",
        "en_title": "QR codes",
        "az": [
            "Hər bir alimin profil səhifəsində QR kod yerləşdirilmişdir. Bu kod vasitəsilə alimin akademik profili istənilən cihazdan, smartfon və ya planşetdən, birbaşa açıla bilər."
        ],
        "en": [
            "Each scientist’s profile page includes a QR code. Scanning it opens the academic profile directly on any device, smartphone, or tablet."
        ],
    },
    {
        "az_title": "Səsləndirmə funksiyası",
        "en_title": "Listen (text-to-speech)",
        "az": [
            "Alimlərin profil mətnlərini oxumaqla yanaşı, dinləmək də mümkündür. İngilis dilində bu funksiya bütün əsas brauzerlərdə işləməlidir. Azərbaycan dili üçün isə səsləndirmə funksiyasının işləməsi istifadəçinin brauzerində və ya mbil cihazında Azərbaycan dili üzrə text-to-speech — TTS — dəstəyinin olub-olmamasından asılıdır."
        ],
        "en": [
            "Scientist profile texts can be read or listened to. In English this function should work in all major browsers. For Azerbaijani, whether text-to-speech works depends on TTS support for Azerbaijani in the user’s browser or mobile device."
        ],
    },
]


SUMMARY_AZ = (
    "Bu sənəd DAAB həmkarlarına yenilənmiş https://daab-waas.com veb-saytının təqdimatını ehtiva edir: "
    "saytın yeni struktur və dizaynı, Azərbaycan və ingilis dillərində tam ikidillilik, onlayn üzvlük "
    "müraciəti, sponsorluq, ianə və tərəfdaşlıq səhifələri, sayt üzrə qlobal axtarış, filtrləmə, "
    "sıralama və qruplaşdırma (o cümlədən alimlər kataloqu), təkmilləşdirilmiş naviqasiya, mobil "
    "uyğun görünüş, mətnin səsləndirilməsi və QR kod funksiyaları qısa izah olunur; məqsəd isə "
    "həmkarları yeni platformanı nəzərdən keçirməyə və mümkün düzəlişlər barədə geribə bildirim "
    "verməyə dəvət etməkdir."
)

SUMMARY_EN = (
    "This document introduces colleagues to the renewed https://daab-waas.com website, summarizing "
    "its new structure and design, full Azerbaijani–English bilingual support, and key capabilities—"
    "including online membership applications, sponsorship and donation pages, site-wide search, "
    "filtering, sorting and grouping (notably in the scientist directory), improved navigation, "
    "responsive layout, text-to-speech, and QR codes—and invites colleagues to review the site and "
    "share feedback for further improvements."
)


LOCALE_META: dict[str, dict[str, str | list[str]]] = {
    "az": {
        "title": DOC_TITLE_AZ,
        "subtitle": "Həmkarlara yeni sayt haqqında məlumat  |  https://daab-waas.com",
        "header": f"DAAB — {DOC_TITLE_AZ}",
        "summary_heading": "Xülasə",
        "summary": SUMMARY_AZ,
        "toc_heading": "Mündəricat",
        "toc_note": (
            "Microsoft Word-da cədvəli seçib F9 düyməsini basın və ya sağ klik → "
            "Update Field → Update entire table."
        ),
        "intro_heading": "Giriş məktubu",
        "features_heading": "Yeni saytın əsas dəyişiklikləri və imkanları",
        "overview_heading": "Qısa icmal",
        "title_key": "az_title",
        "body_key": "az",
        "intro": INTRO_AZ,
    },
    "en": {
        "title": DOC_TITLE_EN,
        "subtitle": "Information for colleagues  |  https://daab-waas.com",
        "header": f"WAAS — {DOC_TITLE_EN}",
        "summary_heading": "Summary",
        "summary": SUMMARY_EN,
        "toc_heading": "Table of contents",
        "toc_note": (
            "In Microsoft Word, select the table below and press F9, "
            "or right-click → Update Field → Update entire table."
        ),
        "intro_heading": "Cover letter",
        "features_heading": "Key changes and capabilities of the new site",
        "overview_heading": "Brief overview",
        "title_key": "en_title",
        "body_key": "en",
        "intro": INTRO_EN,
    },
}


def add_locale_toc_page(document: Document, meta: dict[str, str | list[str]]) -> None:
    document.add_heading(str(meta["toc_heading"]), level=1)
    note = document.add_paragraph()
    note.add_run(str(meta["toc_note"]))
    note.paragraph_format.space_after = Pt(12)
    toc_p = document.add_paragraph()
    add_table_of_contents(toc_p)
    document.add_page_break()


def build_locale_document(lang: str) -> Document:
    meta = LOCALE_META[lang]
    title_key = str(meta["title_key"])
    body_key = str(meta["body_key"])
    intro = meta["intro"]
    assert isinstance(intro, list)

    document = Document()
    setup_styles(document)
    setup_header_footer(document, str(meta["header"]))
    add_title_page(document, str(meta["title"]), str(meta["subtitle"]))
    document.add_heading(str(meta["summary_heading"]), level=1)
    summary_p = document.add_paragraph(str(meta["summary"]))
    summary_p.paragraph_format.space_after = Pt(12)
    document.add_page_break()
    add_locale_toc_page(document, meta)

    document.add_heading(str(meta["intro_heading"]), level=1)
    add_body_paragraphs(document, intro)
    document.add_page_break()

    document.add_heading(str(meta["features_heading"]), level=1)
    for section in FEATURE_SECTIONS:
        add_mono_section(document, section[title_key], section[body_key])

    document.add_page_break()
    document.add_heading(str(meta["overview_heading"]), level=1)
    for section in OVERVIEW_SECTIONS:
        add_mono_section(document, section[title_key], section[body_key])

    set_update_fields_on_open(document)
    return document


def save_document(document: Document, path: Path) -> bool:
    try:
        document.save(str(path))
        return True
    except PermissionError:
        return False


def main() -> int:
    outputs = {"az": OUT_AZ, "en": OUT_EN}
    saved: list[str] = []
    for lang, out_path in outputs.items():
        document = build_locale_document(lang)
        if save_document(document, out_path):
            print(f"Wrote: {out_path.relative_to(ROOT)}")
            saved.append(lang)
        else:
            print(f"Could not save {out_path.name} — close it in Word.", file=sys.stderr)
    if not saved:
        return 1
    print(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} ({', '.join(saved)})")
    return 0 if len(saved) == len(outputs) else 1


if __name__ == "__main__":
    raise SystemExit(main())
