#!/usr/bin/env python3
"""Write documents/DAAB-Cleanup-and-Quality-Review-2026-07.docx."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

try:
    from _paths import ROOT
except ImportError:
    from helpers._paths import ROOT  # type: ignore

OUT = ROOT / "documents" / "DAAB-Cleanup-and-Quality-Review-2026-07.docx"


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    r = run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Right-click → Update Field → Update entire table"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r.append(begin)
    r.append(instr)
    r.append(sep)
    r.append(text)
    r.append(end)


def issue_table(doc: Document, rows: list[tuple[str, str, str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Where", "Problem", "Why it matters", "How to correct", "Priority"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def main() -> int:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str, *, bold: bool = False) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold

    def bullets(items: list[str]) -> None:
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("DAAB / WAAS Website — Cleanup & Quality Review", 0)
    sub = doc.add_paragraph()
    sub.add_run(
        "Comprehensive post-change audit of HTML, CSS, JavaScript, assets, configuration, navigation, "
        "structure, multilingual consistency, accessibility, performance, and UX."
    ).italic = True

    meta = doc.add_paragraph()
    meta.add_run(f"Date: {date.today().isoformat()}\n").bold = True
    meta.add_run("Scope: Repository source (az/, en/, css/, js/, i18n/, images/, helpers/, Deployment package).\n")
    meta.add_run(
        "Automated baseline: helpers/_validate_site.py — OK (90 pages, 6415 refs); "
        "helpers/_deploy_preflight.py — OK after recent cache-version fixes.\n"
    )
    meta.add_run(
        "Method: automated validators + targeted codebase audits (nav/i18n/forms/CSS/JS/assets). "
        "This is not a full visual QA of every breakpoint in a browser."
    )

    note = doc.add_paragraph()
    note.add_run(
        "Note: Broken local paths were not found by the site validator. Findings below focus on quality, "
        "maintainability, consistency, accessibility, performance, and residual production risks."
    ).bold = True

    contents = doc.add_paragraph()
    contents.add_run("Contents").bold = True
    contents.runs[0].font.size = Pt(14)
    toc_note = doc.add_paragraph()
    toc_note.add_run("In Word: right-click TOC → Update Field → Update entire table.").italic = True
    toc_p = doc.add_paragraph()
    add_toc_field(toc_p)
    doc.add_page_break()

    # ------------------------------------------------------------------
    h("1. Executive summary")
    p(
        "The site is in a deployable structural state: bilingual route pairs exist, local asset references "
        "validate, scientists catalogue wiring is clean, and page shells are largely consistent. After recent "
        "legal/footer/nav work, the highest residual risks are (1) membership-form server-side privacy/honeypot "
        "enforcement, (2) CSS/token and cache-bust drift, (3) legal naming/breadcrumb inconsistencies, and "
        "(4) duplicated sticky/footer styling debt."
    )
    p("Counts by priority in this report:", bold=True)
    bullets(
        [
            "High: 6 issues (fix privacy, token drift, cookie-banner cache, routes cache, sticky duplication cluster counted as one program item)",
            "Medium: ~18 issues (legal labels, breadcrumbs, a11y, AZ/EN form values, z-index, unused-manifest drift)",
            "Low: ~15 issues (casing, legacy redirects, vendor cache bust, optional cleanups)",
        ]
    )
    p(
        "Recommended approach: fix High items before or immediately after go-live; schedule Medium in a cleanup "
        "sprint; treat Low as backlog polish."
    )

    # ------------------------------------------------------------------
    h("2. What already looks healthy")
    bullets(
        [
            "No broken local HTML/CSS/JS/image path references detected by _validate_site.py",
            "43 AZ/EN routed page pairs present; sitemap includes legal pages",
            "Scientists list/profiles: 83 cards, catalog sync OK, name-order OK",
            "Page shells: lang, data-daab-page-id, nav mount, footer legal links generally consistent",
            "Inline style= attributes are rare/absent on main pages (good hygiene)",
            "Legal pages exist AZ+EN with shared builder (helpers/_build_legal_pages.py)",
            "Cookie analytics path is consent-gated (GA loads only after consent)",
            "Deployment exclusions correctly omit helpers/ and documents/",
        ]
    )

    # ------------------------------------------------------------------
    h("3. Links, paths, and configuration")
    issue_table(
        doc,
        [
            (
                "Whole site (validator)",
                "No broken local internal paths found in current tree",
                "Confirms packaging/refs are coherent for deploy",
                "Keep running _validate_site.py / _deploy_preflight.py before each upload",
                "Info",
            ),
            (
                "js/daab-i18n.js vs i18n/routes.json",
                "Hardcoded routes.json?v=11 while JSON documentary version is 13",
                "Browsers may keep stale route maps after content changes",
                "Bump ?v= whenever routes.json changes; prefer one version map for i18n JSON",
                "High",
            ),
            (
                "js/daab-i18n.js vs js/daab-search.js",
                "search-index.json loaded as ?v=19 in i18n but fallback ?v=6 in search.js",
                "Fallback path can serve stale search index",
                "Align fallback to same version as daab-i18n.js or remove fallback fetch",
                "Medium",
            ),
            (
                "i18n/routes.json legal navOrder",
                "navOrder: legal-notice=1, privacy=2, cookies=3, terms=4 — conflicts with nav/footer order privacy→terms→cookies→legal-notice",
                "Any UI that sorts by navOrder will show wrong Legal order",
                "Set privacy=1, terms=2, cookies=3, legal-notice=4 to match nav/footer",
                "Medium",
            ),
            (
                "helpers/_deploy_assets.py",
                "Still describes daab-tokens.css / daab-site-background.css as @import; they are inlined into daab-common.css; DYNAMIC_JS omits cookie-consent",
                "Deploy/audit tooling can miss required cookie assets or mis-document sources of truth",
                "Update manifest: inlined tokens/bg; list daab-cookie-consent.js + daab-cookie-banner.css as dynamic",
                "Medium",
            ),
            (
                "az|en/membership.html",
                "Legacy redirect pages on disk but not in routes.json",
                "Orphans confuse audits; OK if intentional aliases",
                "Document as legacyRedirect in routes or keep with clear noindex comment in audit docs",
                "Low",
            ),
        ],
    )

    # ------------------------------------------------------------------
    h("4. Navigation, headers, footers, breadcrumbs")
    issue_table(
        doc,
        [
            (
                "i18n/ui.json vs footer vs legal H1 (EN)",
                "Nav: “Privacy Notice” / “Legal notice”; page/footer: “Privacy notice” / “Legal notice (Imprint)”; Terms/Cookie Title Case vs sentence case",
                "Users see different names for the same page; weakens trust and i18n consistency",
                "Pick one EN string per page; sync ui.json labelKeys, footer snippets, H1, <title>, cookie-banner link text",
                "Medium",
            ),
            (
                "js/daab-breadcrumbs.js + routes navParent legal-pages",
                "Legal pages breadcrumb is Home → page only (no About / Legal parent). FALLBACK_ROUTES omit legal pages",
                "Weaker wayfinding; offline/fallback breadcrumb incomplete",
                "Teach breadcrumbs nested legal-pages group or set navParent about; add four legal routes to FALLBACK_ROUTES",
                "Medium",
            ),
            (
                "az|en/application.html vs nav membershipJoin",
                "H1 “Üzvlüyə müraciət” / “Membership Application” vs breadcrumb/nav “Bizə qoşulun” / “Join us”",
                "Inconsistent page identity in chrome vs content",
                "Add dedicated title key or align H1 with nav label",
                "Medium",
            ),
            (
                "Forum 2024 pages vs Forum 2026",
                "Forum 2024 uses static .forum-breadcrumbs; omits daab-breadcrumbs.js (2026 uses dynamic)",
                "Inconsistent crumb UX across forum years",
                "Document as intentional or migrate 2024 to shared breadcrumb helper",
                "Low",
            ),
            (
                "AZ legal labels",
                "AZ nav/footer/H1 already aligned (Məxfilik bildirişi, İstifadə şərtləri, Kuki siyasəti, Hüquqi rekvizitlər)",
                "Positive finding — keep AZ as the consistency model for EN cleanup",
                "No AZ rename required for legal set",
                "Info",
            ),
        ],
    )

    # ------------------------------------------------------------------
    h("5. HTML structure and reusable components")
    issue_table(
        doc,
        [
            (
                "az|en/scientists/eldar-ahadov-esse.html and eldar-ahadov-poetik-dastanlar.html",
                "Two <h1> elements on the same page",
                "Accessibility / SEO heading outline incorrect",
                "Keep one H1 (hero); demote the second to H2",
                "Medium",
            ),
            (
                "az|en/{privacy,cookies,terms,legal-notice}.html",
                "Tiny redundant <style>html{scroll-behavior:auto!important}</style> while daab-legal-page.css already sets scroll-behavior",
                "Noise in generated HTML; two sources of truth",
                "Remove from _build_legal_pages.py shell if CSS is sufficient",
                "Low",
            ),
            (
                "Legal pages cookie script loading",
                "Legal HTML loads daab-cookie-consent.js directly AND daab-analytics.js may inject it again",
                "Usually guarded, but pattern differs from ~78 other pages",
                "Either omit static tag (analytics inject) or give static tag id=daab-cookie-consent-script",
                "Low",
            ),
            (
                "Scientist media / Eldar standalone pages",
                "Several page ids lack page-subtitles.json entries",
                "Subtitle injector may no-op; uneven hero treatment",
                "Add subtitles or explicitly exclude those page ids",
                "Low",
            ),
        ],
    )

    # ------------------------------------------------------------------
    h("6. CSS, styling consistency, responsiveness")
    issue_table(
        doc,
        [
            (
                "css/daab-tokens.css vs inlined :root in css/daab-common.css",
                "Token file has --z-sidebar-sticky / --z-sidebar-mobile; inlined common :root does not. Risk of token drift after edits",
                "Sticky/sidebar stacking can regress; maintainers edit the wrong file",
                "Re-run inline sync whenever tokens change; treat common :root as generated",
                "High",
            ),
            (
                "css/daab-sticky-sidebar.css vs copies in daab-legal-page.css, daab-charter-page.css, daab-membership-application.css",
                "Near-duplicate sticky TOC/sidebar rule blocks with !important",
                "Hard to keep mobile/desktop sticky behaviour consistent; high maintenance cost",
                "Centralize sticky behaviour in daab-sticky-sidebar.css; import/use from all TOC pages",
                "High",
            ),
            (
                "css/daab-common.css + daab-nav-mega.css + daab-sticky-chrome.css + daab-mobile.css",
                "Heavy !important usage (common ~437, nav-mega ~238, sticky-chrome ~74)",
                "Layout fights across breakpoints; harder responsive fixes",
                "Reduce !important gradually; prefer chrome-ready state + CSS variables",
                "Medium",
            ),
            (
                "css/daab-tokens.css --z-breadcrumbs: 100 vs sticky-chrome fallback 9998",
                "Token scale disagrees with chrome stacking intent",
                "Breadcrumbs can sit under other UI in edge cases",
                "Set --z-breadcrumbs to chrome scale and re-sync common",
                "Medium",
            ),
            (
                "css/daab-cookie-banner.css z-index 10050 vs --z-back-to-top 10050",
                "Same stacking level for banner and back-to-top",
                "Controls can obscure each other on mobile",
                "Assign distinct z-index tokens (e.g. banner 10060)",
                "Medium",
            ),
            (
                "css/daab-scientists-list-page.css",
                "Unscoped .hero / table / thead rules mixed with later page-id-scoped rules",
                "Leakage risk if file is reused; fragile specificity",
                "Prefix remaining bare selectors with html[data-daab-page-id=\"scientists-list\"]",
                "Medium",
            ),
            (
                "Footer styles in common vs mission/foundation/hub/executive-board/flyer/work-done CSS",
                "Footer typography/layout overridden in multiple page CSS files with !important",
                "Footer appearance drifts page-to-page",
                "Keep structure in common; scope page tweaks under data-daab-page-id only when necessary",
                "Medium",
            ),
            (
                "css/daab-forum-book.css, css/daab-membership-page.css",
                "Not linked by live pages; correctly deploy-ignored",
                "Dead weight in repo if unused forever",
                "Keep if build-only; otherwise archive or delete after confirming no generator needs them",
                "Low",
            ),
            (
                "Responsive layouts (desktop/tablet/mobile)",
                "No automated visual regression suite; sticky/nav !important conflicts are the main risk area",
                "Real device bugs may remain after CSS changes",
                "Manual smoke: ≥1181px, ~900px, ≤640px — nav mega, Legal flyout, footer wrap, scientists filters, application form, cookie banner",
                "Medium",
            ),
        ],
    )

    # ------------------------------------------------------------------
    h("7. JavaScript, unused code, performance")
    issue_table(
        doc,
        [
            (
                "js/daab-cookie-consent.js → css/daab-cookie-banner.css?v=1",
                "Banner CSS version hardcoded and not in STYLE_VERSIONS",
                "Stale banner CSS after updates; inconsistent with site-wide cache policy",
                "Add daab-cookie-banner.css to STYLE_VERSIONS; load that version from one place",
                "High",
            ),
            (
                "js/vendor/jspdf.umd.min.js (~356KB) + html2canvas.min.js (~194KB)",
                "Large vendor pair for membership flyer PDF; lazy-loaded (good) but no ?v=",
                "Slow first PDF on weak networks; hard cache invalidation",
                "Keep lazy load; add ?v= on vendor URLs when updating libraries",
                "Medium",
            ),
            (
                "js/video-gallery-data.json",
                "Not referenced by live video_gallery.html (inline data used); deploy-ignored",
                "Obsolete data file can confuse maintainers",
                "Delete or keep only if a helper still regenerates from it — document clearly",
                "Low",
            ),
            (
                "helpers/_tmp_*.py",
                "Temporary inspection/bump scripts present under helpers/",
                "Clutter; risk of accidental reuse",
                "Delete or move to _archive after one-off tasks",
                "Low",
            ),
            (
                "Script fan-out on scientists/legal pages",
                "Profiles/list ~23 scripts; legal ~18 deferred scripts",
                "Many HTTP requests; acceptable for static hosting but not optimal",
                "Later: optional shell bundle; not urgent for launch",
                "Low",
            ),
            (
                "css/daab-forum-content.css (~65KB), daab-activities-page.css (~53KB)",
                "Largest stylesheets with high !important density",
                "Parse/apply cost and maintenance burden",
                "Run CSS usage audit; split or trim dead rules over time",
                "Low",
            ),
        ],
    )

    # ------------------------------------------------------------------
    h("8. Assets (images, fonts, icons)")
    issue_table(
        doc,
        [
            (
                "Deployment images strategy",
                "_build_deployment_folder.py skips images by default (preserve host/Deployment images)",
                "New/changed scientist photos may not upload if you only sync HTML/CSS/JS",
                "When photos change, upload images/ explicitly or build with --include-images",
                "High",
            ),
            (
                "Sampled static <img> on key pages",
                "Alt attributes generally present; profile renderers set name-based alt; decorative QR alt=\"\" inside labeled links",
                "Baseline a11y for images is acceptable on sampled pages",
                "Continue requiring alt in any new markup/helpers",
                "Info",
            ),
            (
                "Forum lightbox images",
                "Lightbox may use empty alt while visible",
                "Screen-reader users miss image context",
                "Update alt when lightbox src changes",
                "Low",
            ),
            (
                "Fonts (css/daab-fonts.css)",
                "Shared font pipeline in place; no broken local font refs reported by validator",
                "Performance depends on host compression/caching",
                "Ensure production sends long-cache headers for fonts/css/js",
                "Low",
            ),
        ],
    )

    # ------------------------------------------------------------------
    h("9. Forms, CTAs, and contact behaviour")
    issue_table(
        doc,
        [
            (
                "mail-application-send.php + az|en/application.html",
                "Privacy checkbox required in HTML/JS only; PHP never validates privacyconfirm / privacy_confirm; consent not written into outbound mail",
                "Applications can be accepted without privacy acknowledgment if JS bypassed; weak audit trail",
                "Reject unless privacy confirmed; include confirmation line in email body",
                "High",
            ),
            (
                "js/daab-membership-application.js honeypot #website vs PHP",
                "Honeypot enforced client-side only; server ignores website field",
                "Spam bots can POST directly",
                "Silent-drop or reject when website is non-empty server-side",
                "Medium",
            ),
            (
                "az/application.html vs en/application.html science field values",
                "AZ-only values biologiya/ekologiya/folklor vs EN biology/ecology/folklore",
                "Broken filtering/reporting if values are compared across locales",
                "Use one canonical value set (prefer EN keys) in both forms; show localized labels",
                "Medium",
            ),
            (
                "az|en/application.html privacy labels",
                "Two label elements point at for=\"privacyconfirm\"",
                "Ambiguous accessible name for the control",
                "Use a single label (or aria-labelledby with one id)",
                "Medium",
            ),
            (
                "Host PHP mail()",
                "Form depends on server mail configuration (az/mail.php, en/mail.php, mail-application-send.php)",
                "Membership CTA fails silently in UX if mail() blocked",
                "Post-deploy: send test AZ+EN applications; confirm inbox delivery",
                "High",
            ),
            (
                "application.html id=\"h1\" on a checkbox",
                "Confusing id next to real headings",
                "Fragile selectors / a11y confusion",
                "Rename to sci-philology (or similar)",
                "Low",
            ),
        ],
    )

    # ------------------------------------------------------------------
    h("10. Accessibility")
    issue_table(
        doc,
        [
            (
                "js/daab-cookie-consent.js banner",
                "role=dialog without aria-modal, initial focus, or focus trap",
                "Keyboard/screen-reader users can struggle with first-visit consent UI",
                "Add aria-modal=true, move focus into dialog, trap/restore focus, Esc to Essential-only or Preferences",
                "Medium",
            ),
            (
                "Eldar dual H1 pages (see §5)",
                "Multiple H1s",
                "Heading outline fails WCAG best practice",
                "Single H1 per page",
                "Medium",
            ),
            (
                "Contrast / keyboard nav across mega-menu",
                "Not fully instrumented in this audit; mega-menu uses complex hover/focus CSS with !important",
                "Risk of focus styles or mobile menu regressions",
                "Manual keyboard pass: Tab through primary nav, Legal nested panel, mobile hamburger, cookie banner",
                "Medium",
            ),
            (
                "Link label casing in cookie banner EN",
                "“Cookie Policy” / “Privacy Notice” vs page H1 sentence case",
                "Minor consistency/a11y naming mismatch",
                "Match page H1 strings",
                "Low",
            ),
        ],
    )

    # ------------------------------------------------------------------
    h("11. Multilingual consistency")
    issue_table(
        doc,
        [
            (
                "EN legal naming (ui.json vs pages)",
                "Privacy Notice vs Privacy notice; Legal notice vs Legal notice (Imprint); Title Case policies vs sentence-case H1s",
                "Nav/footer/document titles disagree",
                "Canonicalize EN strings; update ui.json, footer, builder H1, docx filenames/notes",
                "Medium",
            ),
            (
                "documents/Legal/01-Privacy-Policy.docx",
                "Filename still says Policy while site says Notice",
                "Internal docs diverge from public site",
                "Rename or add cover note: public name is Privacy Notice",
                "Low",
            ),
            (
                "AZ legal set",
                "Aligned across nav/footer/H1",
                "Good model",
                "Keep AZ stable when changing EN",
                "Info",
            ),
            (
                "Application science field values AZ≠EN",
                "Different option values for same disciplines",
                "Data integrity across languages",
                "Canonical values + localized labels",
                "Medium",
            ),
        ],
    )

    # ------------------------------------------------------------------
    h("12. Content / naming / obsolete materials")
    issue_table(
        doc,
        [
            (
                "Git working tree",
                "Large pending change set: new legal pages untracked; many documents/ deletions; scientist photo mods",
                "Hard to know what will be committed vs uploaded; risk of incomplete deploy of images",
                "Commit logical groups; ensure Deployment rebuild before upload; don’t ship documents/",
                "Medium",
            ),
            (
                "templates/ (e.g. Forum 2026.html with spaces)",
                "Template naming inconsistent; not part of deploy root",
                "Confuses contributors",
                "Normalize template names if still used by builders",
                "Low",
            ),
            (
                "Legal compliance Word drafts in documents/Legal/",
                "Internal only — correctly excluded from deploy",
                "Must not be uploaded to public web root",
                "Keep exclusion; never copy into Deployment/",
                "Info",
            ),
        ],
    )

    # ------------------------------------------------------------------
    h("13. Recommended cleanup sequence")
    p("Phase A — before or immediately after production upload", bold=True)
    bullets(
        [
            "Server-side privacy + honeypot checks in mail-application-send.php; test AZ/EN mail delivery",
            "Bump routes.json cache ?v= in daab-i18n.js; add cookie-banner.css to STYLE_VERSIONS",
            "Re-sync daab-tokens.css into daab-common.css; fix z-index tokens",
            "If photos changed: upload images/ (or --include-images)",
            "Manual smoke: nav Legal flyout, footer links, cookie Essential-only, application submit, mobile menu",
        ]
    )
    p("Phase B — cleanup sprint (1–2 weeks)", bold=True)
    bullets(
        [
            "Unify EN legal labels; fix routes navOrder; improve legal breadcrumbs",
            "Deduplicate sticky-sidebar CSS; scope scientists-list bare selectors",
            "Fix Eldar dual H1; application privacy label; sci field canonical values",
            "Cookie banner focus management",
            "Update _deploy_assets.py dynamic/inlined asset lists",
        ]
    )
    p("Phase C — backlog polish", bold=True)
    bullets(
        [
            "Reduce !important density; optional script bundling",
            "Remove obsolete video-gallery-data.json / _tmp helpers if unused",
            "Vendor ?v=; accessibility statement (from legal compliance plan)",
            "Visual QA matrix for tablet breakpoints",
        ]
    )

    # ------------------------------------------------------------------
    h("14. Suggested Definition of Done for “cleanup complete”")
    bullets(
        [
            "python helpers/_validate_site.py and _deploy_preflight.py both OK with 0 WARN from artifact audit",
            "All High items in this report closed or explicitly deferred with owner/date",
            "EN legal names identical in nav, footer, H1, titles",
            "Membership POST without privacyconfirm is rejected by PHP",
            "Cookie banner keyboard-accessible; GA still blocked until consent",
            "Deployment package rebuilt; smoke test signed off on desktop + mobile",
        ]
    )

    h("15. Document control")
    bullets(
        [
            "Document ID: DAAB-Cleanup-and-Quality-Review-2026-07",
            "Location: documents/",
            "Related: documents/Legal/05-Legal-Compliance-Gap-Review.docx (legal content gaps)",
            "Regenerate: python helpers/_write_cleanup_quality_report_docx.py",
            "Status: Internal quality audit — actionable backlog",
        ]
    )

    end = doc.add_paragraph()
    end.add_run(
        "End of report. Automated link validation found no broken local paths; prioritize High form/cache/token "
        "items, then Medium consistency and a11y work."
    ).italic = True

    doc.save(OUT)
    print(f"Wrote {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
