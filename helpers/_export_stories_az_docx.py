#!/usr/bin/env python3
"""Export forum stories to a publication-ready Azerbaijani Word document.

Source text: forum_2024/Hekayələr.docx (Azerbaijani translation of
forum_2024/Рассказы о Форуме.docx). Images from images/forum/.
"""
from __future__ import annotations

import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
if str(ROOT / "helpers") not in sys.path:
    sys.path.insert(0, str(ROOT / "helpers"))

from _build_stories_from_docx import (  # noqa: E402
    DOCX as SOURCE_DOCX,
    SECTION_IMAGES,
    is_pull_quote,
    parse_docx,
)
from daab_docx_export import (  # noqa: E402
    BLUE,
    FONT_BODY,
    FONT_HEADING,
    GOLD,
    INK,
    MUTED,
    NAVY,
    QUOTE_FILL,
    add_page_number_field,
    add_table_of_contents,
    set_update_fields_on_open,
    setup_styles,
    shade_paragraph,
)

OUT_DOCX = ROOT / "forum_2024" / "Forumla_bagli_hekayeler.docx"
OUT_COPY = ROOT / "documents" / "docx" / "Forumla-bagli-hekayeler.docx"
IMAGES = ROOT / "images" / "forum"

ORG_AZ_LINE1 = "DÜNYA AZƏRBAYCANLI ALIMLƏR BİRLİYİ"
ORG_AZ_SHORT = "DAAB / WAAS"
DOC_TITLE = "Forumla bağlı hekayələr"
DOC_SUBTITLE = "Eldar Əhədovun ədəbi yazıları"
HEADER_TITLE = "DAAB — Forumla bağlı hekayələr"


def setup_header_footer_az(document: Document) -> None:
    section = document.sections[0]
    section.different_first_page_header_footer = True
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    header = section.header
    header._element.clear()
    hp = header.add_paragraph()
    run = hp.add_run(HEADER_TITLE)
    run.font.name = FONT_BODY
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    section.first_page_header._element.clear()
    section.first_page_header.add_paragraph()

    footer = section.footer
    footer._element.clear()
    fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run(f"{ORG_AZ_SHORT}   |   Səhifə ")
    r1.font.size = Pt(9)
    r1.font.color.rgb = MUTED
    add_page_number_field(fp, align_center=False)
    month_year = datetime.now().strftime("%B %Y")
    r2 = fp.add_run(f"   |   {month_year}")
    r2.font.size = Pt(9)
    r2.font.color.rgb = MUTED

    first_footer = section.first_page_footer
    first_footer._element.clear()
    fp_first = first_footer.add_paragraph()
    fp_first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = fp_first.add_run("daab-waas.com")
    rf.font.size = Pt(9)
    rf.font.color.rgb = GOLD


def add_title_page_az(document: Document, author_line: str) -> None:
    for _ in range(3):
        document.add_paragraph()

    org = document.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ro = org.add_run(ORG_AZ_LINE1)
    ro.font.name = FONT_BODY
    ro.font.size = Pt(13)
    ro.font.color.rgb = MUTED
    ro.font.small_caps = True

    document.add_paragraph()

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title_p.add_run(DOC_TITLE)
    rt.font.name = FONT_HEADING
    rt.font.size = Pt(26)
    rt.font.bold = True
    rt.font.color.rgb = NAVY

    rule = document.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rule.add_run("━" * 32)
    rr.font.color.rgb = GOLD

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(DOC_SUBTITLE)
    rs.font.size = Pt(13)
    rs.font.italic = True
    rs.font.color.rgb = BLUE

    if author_line:
        lead = document.add_paragraph()
        lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = lead.add_run(author_line)
        rl.font.size = Pt(11)
        rl.font.bold = True
        rl.font.color.rgb = INK

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(24)
    rm = meta.add_run(
        "Xaricdə Yaşayan Azərbaycanlı Alimlərin I Forumu\n"
        f"Hazırlanma tarixi: {datetime.now().strftime('%d %B %Y')}"
    )
    rm.font.size = Pt(10)
    rm.font.color.rgb = MUTED

    document.add_page_break()


def add_toc_page_az(document: Document) -> None:
    document.add_heading("Mündəricat", level=1)
    note = document.add_paragraph()
    note.paragraph_format.space_after = Pt(10)
    rn = note.add_run(
        "Microsoft Word-də cədvələ klik edin və F9 düyməsini basın, "
        "və ya sağ klik → Sahəni yenilə → Bütün cədvəli yenilə."
    )
    rn.font.size = Pt(10)
    rn.font.italic = True
    rn.font.color.rgb = MUTED
    toc_p = document.add_paragraph()
    add_table_of_contents(toc_p)
    document.add_page_break()


def add_body_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(1.0)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(11)
    run.font.color.rgb = INK


def add_blockquote_az(document: Document, quote: str, cite: str = "") -> None:
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.right_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    shade_paragraph(p, QUOTE_FILL)
    rq = p.add_run(quote)
    rq.font.name = FONT_BODY
    rq.font.size = Pt(11)
    rq.font.italic = True
    rq.font.color.rgb = INK
    if cite:
        cp = document.add_paragraph()
        cp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cp.paragraph_format.right_indent = Cm(0.6)
        cp.paragraph_format.space_after = Pt(10)
        rc = cp.add_run(cite)
        rc.font.name = FONT_BODY
        rc.font.size = Pt(10)
        rc.font.bold = True
        rc.font.color.rgb = BLUE


def add_story_image(document: Document, section_id: str, alt: str) -> None:
    filename = SECTION_IMAGES.get(section_id)
    if not filename:
        return
    path = IMAGES / filename
    if not path.is_file():
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run()
    width = Cm(10.5) if section_id == "veten-hissleri" else Cm(14.5)
    run.add_picture(str(path), width=width)


def paragraphs_to_docx(document: Document, paragraphs: list[str]) -> None:
    i = 0
    while i < len(paragraphs):
        line = paragraphs[i]
        nxt = paragraphs[i + 1] if i + 1 < len(paragraphs) else None
        if is_pull_quote(line, nxt):
            quote_text = line.strip().strip('"').strip("\u201c\u201d").strip()
            cite = (nxt or "").strip()
            add_blockquote_az(document, quote_text, cite)
            i += 2
            continue
        if line.startswith("— ") and i > 0 and paragraphs[i - 1].startswith('"'):
            i += 1
            continue
        add_body_paragraph(document, line)
        i += 1


def add_intro_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = NAVY


def export_stories_docx(data: dict) -> Path:
    document = Document()
    setup_styles(document)
    setup_header_footer_az(document)
    add_title_page_az(document, data.get("title", ""))
    add_toc_page_az(document)

    if data.get("title"):
        add_intro_paragraph(document, data["title"])
        document.add_paragraph()

    for index, section in enumerate(data["sections"]):
        if index > 0:
            document.add_page_break()
        document.add_heading(section["title"], level=1)
        add_story_image(document, section["id"], section["title"])
        paragraphs_to_docx(document, section["paragraphs"])

    set_update_fields_on_open(document)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUT_DOCX))

    OUT_COPY.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUT_COPY))
    return OUT_DOCX


def main() -> None:
    if not SOURCE_DOCX.is_file():
        raise SystemExit(f"Missing source: {SOURCE_DOCX.relative_to(ROOT)}")
    data = parse_docx(Document(str(SOURCE_DOCX)))
    out = export_stories_docx(data)
    print(f"Wrote {out.relative_to(ROOT)} ({len(data['sections'])} stories)")
    print(f"Copy: {OUT_COPY.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
