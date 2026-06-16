#!/usr/bin/env python3
"""Build Forum 2024 session organization pages from Word source.

Outputs:
  - documents/preview/forum_sessions_organization.html (review copy)
  - az/forum/2024/sessions_organization.html
  - en/forum/2024/sessions_organization.html
"""
from __future__ import annotations

import html
import re
import subprocess
import sys
import zipfile
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from _paths import ROOT
from _site_wide_cleanup import SCRIPT_VERSIONS, STYLE_VERSIONS
from _sync_scientists_from_docx import ALIASES, person_key
from forum_en_common import FORUM_FOOTER_EN
from forum_en_sessions import (
    META_DESCRIPTION_EN,
    SUMMARY_EN,
    SUMMARY_TITLE_EN,
    UI_EN,
    SessionsTranslator,
)
from forum_en_sessions_structured import localize_structured_section
from scientists_profiles_core import load_profiles

DOCX = (
    ROOT
    / "documents"
    / "docx"
    / "DAAB Forum - Sessiyaların Təşkili (AZ) 14.06.2026.docx"
)
OUT_DIR = ROOT / "documents" / "preview"
OUT_HTML = OUT_DIR / "forum_sessions_organization.html"
OUT_IMAGES = OUT_DIR / "images" / "forum-sessions"
LIVE_IMAGES = ROOT / "images" / "forum-sessions"
AZ_OUT = ROOT / "az" / "forum" / "2024" / "sessions_organization.html"
EN_OUT = ROOT / "en" / "forum" / "2024" / "sessions_organization.html"
PHOTOS_DIR = ROOT / "images" / "scientists-photos"
NAV_SOURCE = ROOT / "az" / "forum" / "2024" / "program.html"
PREVIEW_ASSET = "../../"
LIVE_ASSET = "../../../"
PAGE_ID = "forum-sessions-organization"
PHOTO_HEADER = "Foto"

SESSION_PERSON_ALIASES = {
    **ALIASES,
    "dinarə abbasova": "dinara abbasova",
    "agamalı mamedov": "agamalı məmmədov",
    "nizami mamedov": "nizami məmmədov",
    "mehdi ismayilov": "mehdi gəncəli ismayilov",
    "qafar mehdiyev": "qafar mehdiyev caxmaxcı",
    "qafar": "qafar mehdiyev caxmaxcı",
    "qərib murshydov": "qərib mursudov",
    "reza moridi": "rıza moridi",
    "süleyman allahverdiyev": "suleyman allahverdiyev",
    "makbulə sabziyeva": "məqbulə səbziyeva",
    "sevinc məmmədova": "sevinj məmmədova",
    "sevinc esgerova": "sevinc esgerova",
    "sevinc əsgərova": "sevinc esgerova",
    "həsən babazadə": "hasan babazada",
}

SESSION_PHOTO_OVERRIDES = {
    "hasan babazada": "hasan-babazada.png",
    "həsən babazadə": "hasan-babazada.png",
}

TITLE = "Sessiyaların təşkili və yönləndirilməsi"
SUBTITLE = "10 sentyabr 2024 — strateji sessiyaların QARIŞIQ və İXTİSAS qrupları üzrə təşkili"
SUMMARY_ARIA = "Günün icmalı"
SUMMARY_TITLE = "Günün icmalı"
SUMMARY = (
    "Bu səhifədə Forumun 10 sentyabr proqramı çərçivəsində səhər saatlarında qarışıq qruplar, "
    "nahardan sonra isə ixtisas qrupları üzrə beyin fırtınası prinsipi əsasında təşkil olunmuş "
    "sessiyalar haqqında ətraflı məlumat verilir. Əlavə olaraq, iştirakçıların masalar ətrafında "
    "yerləşdirilməsi qaydası, sessiyaların təşkili prinsipləri, moderatorların rolu və vəzifələri "
    "barədə izahlı məlumat təqdim olunur."
)

SIDEBAR_SCRIPT = """
<script>
(function () {
  const links = Array.from(document.querySelectorAll('.timeline-list a[href^="#"]'));
  const ids = links.map(a => a.getAttribute('href').slice(1));
  const cards = ids.map(id => document.getElementById(id)).filter(Boolean);
  const eventsWidget = document.querySelector('.sidebar-widget');
  const eventsToggle = document.querySelector('.events-menu-toggle');
  const mobileQuery = window.matchMedia('(max-width: 1060px)');

  function activate(link) {
    links.forEach(a => a.classList.remove('tl-active'));
    if (link) link.classList.add('tl-active');
  }

  function closeEventsMenu() {
    if (!eventsWidget || !eventsToggle) return;
    eventsWidget.classList.remove('events-open');
    eventsToggle.setAttribute('aria-expanded', 'false');
  }

  function toggleEventsMenu() {
    if (!eventsWidget || !eventsToggle) return;
    const open = eventsWidget.classList.toggle('events-open');
    eventsToggle.setAttribute('aria-expanded', open ? 'true' : 'false');
  }

  function jumpToTarget(event) {
    const link = event.currentTarget;
    const id = link.getAttribute('href').slice(1);
    const target = document.getElementById(id);
    if (!target) return;
    event.preventDefault();
    activate(link);
    const Pos = window.DAAB_LANG_POSITION;
    if (Pos && Pos.scrollToAnchor) {
      Pos.scrollToAnchor(id, false);
    } else {
      target.scrollIntoView({ block: 'start', behavior: 'auto' });
    }
    history.pushState(null, '', link.getAttribute('href'));
    if (mobileQuery.matches) closeEventsMenu();
  }

  function onScroll() {
    const mid = window.scrollY + window.innerHeight * 0.35;
    let active = null;
    for (let i = cards.length - 1; i >= 0; i--) {
      if (cards[i] && cards[i].offsetTop <= mid) {
        active = i;
        break;
      }
    }
    activate(active !== null ? links[ids.indexOf(cards[active].id)] : null);
  }

  links.forEach(link => link.addEventListener('click', jumpToTarget));
  if (eventsToggle) {
    eventsToggle.addEventListener('click', event => {
      event.stopPropagation();
      toggleEventsMenu();
    });
  }
  document.addEventListener('click', event => {
    if (!mobileQuery.matches || !eventsWidget || !eventsWidget.classList.contains('events-open')) return;
    if (eventsWidget.contains(event.target)) return;
    closeEventsMenu();
  });
  document.addEventListener('keydown', event => {
    if (event.key === 'Escape') closeEventsMenu();
  });
  window.addEventListener('scroll', onScroll, { passive: true });
  onScroll();
})();
</script>"""


def esc(text: str) -> str:
    return html.escape(text, quote=True)


def norm(text: str) -> str:
    return text.replace("\r", "").replace("\xa0", " ").strip()


def slugify(text: str) -> str:
    s = text.lower()
    s = s.replace("ı", "i").replace("ə", "e").replace("ö", "o").replace("ü", "u")
    s = s.replace("ğ", "g").replace("ş", "s").replace("ç", "c").replace("İ", "i")
    s = re.sub(r"[^a-z0-9]+", "-", s).strip("-")
    return s or "section"


def find_heading_index(
    paras,
    *needles: str,
    start: int = 0,
    heading_only: bool = False,
) -> int:
    targets = {norm(n).upper() for n in needles}
    for i in range(start, len(paras)):
        text = norm(paras[i].text).upper()
        if not text:
            continue
        style = paras[i].style.name if paras[i].style else ""
        if heading_only and not style.startswith("Heading"):
            continue
        if text in targets:
            return i
        if not heading_only and any(needle in text for needle in targets):
            return i
    return -1


def collect_footnote_bullets(paras, start: int) -> list[str]:
    items: list[str] = []
    for i in range(start, len(paras)):
        text = norm(paras[i].text)
        style = paras[i].style.name if paras[i].style else ""
        if style == "Caption" and text.startswith("Şəkil"):
            break
        if not text:
            continue
        if style in {"footnote text", "List Paragraph"}:
            items.append(text)
    return items


def parse_structured_section(
    paras,
    start_idx: int,
    end_idx: int,
    *,
    nested_lists: bool,
) -> dict:
    title = norm(paras[start_idx].text)
    blocks: list[dict] = []
    current_list: list[dict] | None = None
    current_item: dict | None = None

    def flush_list() -> None:
        nonlocal current_list, current_item
        if current_list:
            blocks.append({"type": "list", "items": current_list})
        current_list = None
        current_item = None

    for i in range(start_idx + 1, end_idx):
        text = norm(paras[i].text)
        if not text:
            continue
        style = paras[i].style.name if paras[i].style else ""

        if style.startswith("Heading"):
            flush_list()
            if "Heading 2" in style or "Heading 3" in style:
                level = 3
            else:
                level = 4
            blocks.append({"type": "heading", "level": level, "text": text})
            continue

        if style == "Normal":
            flush_list()
            blocks.append({"type": "paragraph", "text": text})
            continue

        if style == "List Paragraph":
            if current_list is None:
                current_list = []
            current_item = {"text": text, "children": []}
            current_list.append(current_item)
            continue

        if style == "Plain Text":
            if nested_lists and current_item is not None:
                current_item["children"].append(text)
            else:
                flush_list()
                blocks.append({"type": "paragraph", "text": text})
            continue

        flush_list()
        blocks.append({"type": "paragraph", "text": text})

    flush_list()
    return {"title": title, "blocks": blocks}


def render_list_items(items: list[dict]) -> str:
    lis: list[str] = []
    for item in items:
        child_html = ""
        children = item.get("children") or []
        if children:
            sub = "".join(f"<li>{esc(child)}</li>" for child in children)
            child_html = f'<ul class="content-list">{sub}</ul>'
        lis.append(f"<li>{esc(item['text'])}{child_html}</li>")
    return f'<ul class="content-list">{"".join(lis)}</ul>'


def render_structured_block(block: dict) -> str:
    block_type = block.get("type")
    if block_type == "paragraph":
        return f'<p class="card-text">{esc(block["text"])}</p>'
    if block_type == "heading":
        level = block.get("level", 3)
        tag = "h3" if level <= 3 else "h4"
        css_class = "sessions-subsection-title" if level <= 3 else "sessions-topic-title"
        return f'<{tag} class="{css_class}">{esc(block["text"])}</{tag}>'
    if block_type == "list":
        return render_list_items(block.get("items", []))
    return ""


def session_person_key(name: str) -> str:
    key = person_key(normalize_person_name(name))
    return SESSION_PERSON_ALIASES.get(key, key)


def normalize_person_name(name: str) -> str:
    return name.strip().rstrip(",").strip()


class ParticipantPhotoIndex:
    """Resolve participant names from session tables to scientist profile photos."""

    def __init__(self, asset_prefix: str) -> None:
        self.asset_prefix = asset_prefix
        self.profiles_by_key: dict[str, dict] = {}
        for profile in load_profiles():
            self.profiles_by_key[session_person_key(profile["name"])] = profile
            self.profiles_by_key[person_key(profile["name"])] = profile
            if profile.get("name_en"):
                self.profiles_by_key[session_person_key(profile["name_en"])] = profile
                self.profiles_by_key[person_key(profile["name_en"])] = profile
        self.photo_files = {
            path.stem: path.name
            for path in PHOTOS_DIR.glob("*")
            if path.suffix.lower() in {".png", ".jpg", ".jpeg"}
        }

    def resolve(self, name: str) -> str | None:
        key = session_person_key(name)
        override = SESSION_PHOTO_OVERRIDES.get(key)
        if override and (PHOTOS_DIR / override).is_file():
            return override
        profile = self.profiles_by_key.get(key)
        if profile:
            photo = (profile.get("photo") or "").strip()
            if photo and (PHOTOS_DIR / photo).is_file():
                return photo
        best_stem = None
        best_score = 0.0
        for stem in self.photo_files:
            score = SequenceMatcher(None, key, stem.replace("-", " ")).ratio()
            if score > best_score:
                best_score = score
                best_stem = stem
        if best_stem and best_score >= 0.72:
            return self.photo_files[best_stem]
        return None

    def cell_html(self, name: str, *, alt_name: str | None = None) -> str:
        clean = normalize_person_name(name)
        alt = alt_name if alt_name is not None else clean
        photo = self.resolve(clean)
        if photo:
            src = f"{self.asset_prefix}images/scientists-photos/_thumbs/{Path(photo).stem}.jpg"
            return (
                f'<td class="sessions-photo-cell">'
                f'<span class="sessions-photo-frame">'
                f'<img class="sessions-participant-photo" src="{esc(src)}" '
                f'alt="{esc(alt)}" width="44" height="44" loading="lazy" decoding="async">'
                f"</span></td>"
            )
        initials = "".join(word[0] for word in clean.split()[:2] if word).upper()[:2] or "?"
        return (
            f'<td class="sessions-photo-cell sessions-photo-cell--empty">'
            f'<span class="sessions-photo-frame sessions-photo-frame--empty">'
            f'<span class="sessions-photo-placeholder" aria-hidden="true">{esc(initials)}</span>'
            f"</span></td>"
        )


def extract_para_image(para: Paragraph) -> str | None:
    for blip in para._element.findall(".//" + qn("a:blip")):
        rel_id = blip.get(qn("r:embed"))
        if rel_id and para.part.rels.get(rel_id):
            return para.part.rels[rel_id].target_ref
    return None


def export_images(doc: Document) -> dict[str, str]:
    LIVE_IMAGES.mkdir(parents=True, exist_ok=True)
    OUT_IMAGES.mkdir(parents=True, exist_ok=True)
    mapping: dict[str, str] = {}
    with zipfile.ZipFile(DOCX) as zf:
        for name in zf.namelist():
            if not name.startswith("word/media/"):
                continue
            src_name = Path(name).name
            out_name = src_name
            data = zf.read(name)
            (LIVE_IMAGES / out_name).write_bytes(data)
            (OUT_IMAGES / out_name).write_bytes(data)
            mapping[src_name] = f"images/forum-sessions/{out_name}"

    figure_paths: list[str] = []
    for para in doc.paragraphs:
        rel = extract_para_image(para)
        if rel:
            media_name = Path(rel).name
            if media_name in mapping:
                figure_paths.append(mapping[media_name])

    figures: dict[str, str] = {}
    if figure_paths:
        figures["fig1"] = figure_paths[0]
    if len(figure_paths) > 1:
        figures["fig2"] = figure_paths[1]
    return figures


def participant_headers(headers: list[str]) -> list[str]:
    if headers and headers[0] == "№" and PHOTO_HEADER not in headers:
        return [headers[0], PHOTO_HEADER, *headers[1:]]
    return headers


MODERATOR_SPECIALTY_FIXES = {
    "bəxtiyar siracov": "Kompüter elmləri / Dr.",
    "səadət kərimi": "Riyaziyyat və mexanika / Prof. Dr.",
    "teymur rzayev": "Rəssam / Prof. Dr.",
}


def moderator_specialty(name: str, specialty: str) -> str:
    fixed = MODERATOR_SPECIALTY_FIXES.get(session_person_key(name))
    return fixed if fixed else specialty


def moderator_photo_headers(
    headers: list[str], name_header: str
) -> tuple[list[str], int, bool]:
    """Return display headers with row number and photo first."""
    clean_headers = [header for header in headers if header != PHOTO_HEADER]
    target = name_header.strip().rstrip(",")
    name_idx = next(
        i
        for i, header in enumerate(clean_headers)
        if header.strip().rstrip(",") == target
    )
    has_source_num = bool(clean_headers) and clean_headers[0] == "№"
    if has_source_num:
        display = [clean_headers[0], PHOTO_HEADER, *clean_headers[1:]]
    else:
        display = ["№", PHOTO_HEADER, *clean_headers]
    return display, name_idx, has_source_num


def moderator_cell_display(
    *,
    cell: str,
    header: str,
    name_idx: int,
    cell_idx: int,
    moderator_name: str,
    translator: SessionsTranslator | None = None,
) -> str:
    if translator:
        return translator.cell(header, cell, moderator_name)
    if cell_idx == name_idx:
        return normalize_person_name(cell)
    if header.strip().rstrip(",") == "İxtisas":
        return moderator_specialty(moderator_name, cell)
    return cell


def table_html(
    table,
    *,
    headers: list[str] | None = None,
    skip_header_row: bool = False,
    photo_index: ParticipantPhotoIndex | None = None,
    photo_name_header: str | None = None,
    translator: SessionsTranslator | None = None,
) -> str:
    rows = table.rows
    if not rows:
        return ""
    start = 1 if skip_header_row else 0
    if headers is None and not skip_header_row:
        headers = [norm(c.text) for c in rows[0].cells]
        start = 1
    headers = list(headers or [])
    with_moderator_photos = (
        photo_index is not None
        and photo_name_header
        and any(h.strip().rstrip(",") == photo_name_header.strip().rstrip(",") for h in headers)
    )
    with_participant_photos = (
        photo_index is not None
        and headers
        and headers[0] == "№"
        and not with_moderator_photos
    )
    moderator_name_idx = -1
    has_source_row_num = False
    moderator_data_headers: list[str] | None = None
    if with_participant_photos:
        headers = participant_headers(headers)
    elif with_moderator_photos:
        moderator_data_headers = [header for header in headers if header != PHOTO_HEADER]
        headers, moderator_name_idx, has_source_row_num = moderator_photo_headers(
            moderator_data_headers, photo_name_header or ""
        )
    photo_header = translator.photo_header() if translator else PHOTO_HEADER
    thead = ""
    if headers:
        th_parts: list[str] = []
        for header in headers:
            if header == PHOTO_HEADER:
                th_parts.append(
                    f'<th scope="col" class="sessions-photo-col">{esc(photo_header)}</th>'
                )
            else:
                display_header = translator.header(header) if translator else header
                th_parts.append(f"<th>{esc(display_header)}</th>")
        thead = f"<thead><tr>{''.join(th_parts)}</tr></thead>"
    body_rows = []
    if with_participant_photos:
        data_headers = [headers[0], *headers[2:]]
    elif with_moderator_photos:
        data_headers = moderator_data_headers or []
    else:
        data_headers = headers
    moderator_row_num = 0
    for row in rows[start:]:
        cells = [norm(c.text) for c in row.cells]
        if not any(cells):
            continue
        if with_participant_photos and not re.fullmatch(r"\d+", cells[0] or ""):
            continue
        while len(cells) < len(data_headers or cells):
            cells.append("")
        if with_participant_photos and photo_index:
            name = cells[1] if len(cells) > 1 else ""
            alt_name = translator.name(name) if translator else None
            tds = f"<td>{esc(cells[0])}</td>{photo_index.cell_html(name, alt_name=alt_name)}"
            for i, cell in enumerate(cells[1:], start=1):
                header = data_headers[i] if i < len(data_headers) else ""
                display = (
                    translator.cell(header, cell, name)
                    if translator
                    else cell
                )
                tds += f"<td>{esc(display)}</td>"
        elif with_moderator_photos and photo_index:
            name = cells[moderator_name_idx] if len(cells) > moderator_name_idx else ""
            alt_name = translator.name(name) if translator else None
            if has_source_row_num:
                row_num = cells[0]
            else:
                moderator_row_num += 1
                row_num = str(moderator_row_num)
            tds = f"<td>{esc(row_num)}</td>{photo_index.cell_html(name, alt_name=alt_name)}"
            for i, cell in enumerate(cells[: len(data_headers)]):
                if has_source_row_num and i == 0:
                    continue
                header = data_headers[i]
                display = moderator_cell_display(
                    cell=cell,
                    header=header,
                    name_idx=moderator_name_idx,
                    cell_idx=i,
                    moderator_name=name,
                    translator=translator,
                )
                tds += f"<td>{esc(display)}</td>"
        else:
            tds = "".join(f"<td>{esc(c)}</td>" for c in cells[: len(data_headers or cells)])
        body_rows.append(f"<tr>{tds}</tr>")
    if not body_rows:
        return ""
    table_class = "program-table program-table--wide"
    if with_participant_photos or with_moderator_photos:
        table_class += " program-table--with-photos"
    if with_moderator_photos:
        table_class += " program-table--moderators"
    return (
        f'<div class="program-table-wrap"><table class="{table_class}">'
        f"{thead}<tbody>{''.join(body_rows)}</tbody></table></div>"
    )


def parse_intro(doc: Document, figures: dict[str, str]) -> list[dict]:
    paras = doc.paragraphs
    sections: list[dict] = []

    def collect_until(end_idx: int, start: int) -> list[str]:
        items: list[str] = []
        for i in range(start, end_idx):
            text = norm(paras[i].text)
            if not text:
                continue
            style = paras[i].style.name if paras[i].style else ""
            if "List" in style or style in {"footnote text", "Normal"}:
                items.append(text)
        return items

    principles = collect_until(11, 1)
    sections.append(
        {
            "id": "teskil-prinsipleri",
            "title": "Təşkil prinsipləri",
            "paragraphs": principles[:3],
            "bullets": principles[3:],
        }
    )

    sections.append(
        {
            "id": "qarisiq-muzakire",
            "title": "QARIŞIQ qruplar — müzakirə mövzuları",
            "lead": "QARIŞIQ qruplar",
            "bullets": collect_footnote_bullets(paras, 12),
            "figure": figures.get("fig1"),
            "caption": "Şəkil 1. QARIŞIQ qrupların fəaliyyəti",
        }
    )

    start_strategi = find_heading_index(
        paras,
        "STRATEJİ PLANLAŞDIRMA İSTİQAMƏTLƏRİ",
        heading_only=True,
    )
    start_elm = find_heading_index(
        paras,
        "ELM SAHƏLƏRİ ÜZRƏ TÖVSİYƏ OLUNAN MÜZAKİRƏ MÖVZULARI",
        heading_only=True,
    )
    if start_strategi < 0 or start_elm < 0:
        raise ValueError("Structured planning sections not found in Word source.")

    strategi = parse_structured_section(
        paras, start_strategi, start_elm, nested_lists=False
    )
    sections.append(
        {
            "id": "strategi-planlasdirma",
            "title": strategi["title"],
            "blocks": strategi["blocks"],
        }
    )

    sections.append(
        {
            "id": "ixtisas-muzakire",
            "title": "İXTİSAS qrupları — müzakirə mövzuları",
            "lead": "İXTİSAS qrupları",
            "bullets": collect_footnote_bullets(paras, 27),
            "figure": figures.get("fig2"),
            "caption": "Şəkil 2. İXTİSAS qruplarının fəaliyyəti",
        }
    )

    elm = parse_structured_section(
        paras, start_elm, len(paras), nested_lists=True
    )
    sections.append(
        {
            "id": "elm-saheleri-tovsiyeler",
            "title": elm["title"],
            "blocks": elm["blocks"],
        }
    )

    return sections


def render_intro_card(section: dict, asset: str) -> str:
    parts: list[str] = []
    if section.get("blocks"):
        parts.extend(render_structured_block(block) for block in section["blocks"])
    else:
        for p in section.get("paragraphs", []):
            parts.append(f'<p class="card-text">{esc(p)}</p>')
        if section.get("lead"):
            parts.append(f'<p class="card-text"><strong>{esc(section["lead"])}</strong></p>')
        if section.get("example_lead"):
            parts.append(f'<p class="card-text">{esc(section["example_lead"])}</p>')
        bullets = section.get("bullets", [])
        if bullets:
            lis = "".join(f"<li>{esc(b)}</li>" for b in bullets)
            parts.append(f'<ul class="content-list">{lis}</ul>')
        if section.get("figure"):
            figure_src = f"{asset}{section['figure']}"
            parts.append(
                f'<figure class="sessions-figure">'
                f'<img src="{esc(figure_src)}" alt="{esc(section.get("caption", ""))}">'
                f'<figcaption>{esc(section.get("caption", ""))}</figcaption>'
                f"</figure>"
            )
    card_class = "news-card sessions-intro-card"
    if section.get("blocks"):
        card_class += " sessions-detail-card"
    return f"""
<article class="{card_class}" id="{esc(section["id"])}">
<div class="card-header"><h2 class="card-title">{esc(section["title"])}</h2></div>
<div class="card-body">{"".join(parts)}</div>
</article>"""


def parse_qarisiq_tables(
    doc: Document,
    photo_index: ParticipantPhotoIndex,
    translator: SessionsTranslator | None = None,
) -> tuple[str, list[dict]]:
    moderators = table_html(
        doc.tables[0],
        headers=["Moderator", "İxtisas", "Ölkə", "Masalar"],
        skip_header_row=True,
        photo_index=photo_index,
        photo_name_header="Moderator",
        translator=translator,
    )
    masas: list[dict] = []
    masa_re = re.compile(r"^MASA\s+(\d+)$", re.I)
    table_idx = 1
    for para in doc.paragraphs:
        m = masa_re.match(norm(para.text))
        if not m or table_idx >= 14:
            continue
        n = m.group(1)
        masas.append(
            {
                "id": f"qarisiq-masa-{n}",
                "title": f"MASA {n}",
                "table_html": table_html(
                    doc.tables[table_idx],
                    headers=["№", "Ad, soyad", "İxtisas", "Ölkə"],
                    photo_index=photo_index,
                    translator=translator,
                ),
            }
        )
        table_idx += 1
    return moderators, masas


def parse_ixtisas_tables(
    doc: Document,
    photo_index: ParticipantPhotoIndex,
    translator: SessionsTranslator | None = None,
) -> tuple[str, str, list[dict]]:
    mod_simple = table_html(
        doc.tables[14],
        headers=["Moderatorlar", "Masalar"],
        skip_header_row=True,
        photo_index=photo_index,
        photo_name_header="Moderatorlar",
        translator=translator,
    )
    mod_detail = table_html(
        doc.tables[15],
        headers=["№", "Elm sahəsi", "Moderator", "Masalar"],
        photo_index=photo_index,
        photo_name_header="Moderator",
        translator=translator,
    )
    groups: list[dict] = []
    current_field = ""
    field_slug = ""
    table_idx = 16
    in_ixtisas = False
    masa_re = re.compile(r"^MASA\s+(\d+)$", re.I)

    for para in doc.paragraphs:
        text = norm(para.text)
        if not text:
            continue
        style = para.style.name if para.style else ""
        if style == "Heading 1" and "İXTİSAS QRUPLARI" in text.upper():
            in_ixtisas = True
            continue
        if not in_ixtisas:
            continue
        if style == "Heading 2" and not masa_re.match(text):
            current_field = text
            field_slug = slugify(current_field)
            continue
        m = masa_re.match(text)
        if m and table_idx <= 30:
            n = m.group(1)
            masa_id = f"ixtisas-{field_slug}-masa-{n}" if field_slug else f"ixtisas-masa-{n}"
            groups.append(
                {
                    "id": masa_id,
                    "field": current_field,
                    "field_slug": field_slug,
                    "title": f"MASA {n}",
                    "table_html": table_html(
                        doc.tables[table_idx],
                        headers=["№", "Ad, soyad", "İxtisas", "Ölkə"],
                        photo_index=photo_index,
                        translator=translator,
                    ),
                }
            )
            table_idx += 1
    return mod_simple, mod_detail, groups


def render_qarisiq_section(
    moderators: str,
    masas: list[dict],
    translator: SessionsTranslator | None = None,
) -> str:
    masa_blocks = []
    for masa in masas:
        title = masa["title"]
        if translator:
            title = translator.table_title(title)
        masa_blocks.append(
            f'<section class="sessions-masa-block" id="{esc(masa["id"])}">'
            f'<h3 class="sessions-masa-title">{esc(title)}</h3>'
            f'{masa["table_html"]}'
            f"</section>"
        )
    section_title = (
        UI_EN["qarisiq_tables_title"]
        if translator
        else "QARIŞIQ qruplar üzrə masalar"
    )
    lead = UI_EN["moderators_lead"] if translator else "Sessiyaların yönəldiciləri və moderatorlar"
    return f"""
<article class="news-card sessions-tables-card" id="qarisiq-qruplar">
<div class="card-header"><h2 class="card-title">{esc(section_title)}</h2></div>
<div class="card-body">
<p class="card-text">{esc(lead)}</p>
{moderators}
{"".join(masa_blocks)}
</div>
</article>"""


def render_ixtisas_section(
    mod_simple: str,
    mod_detail: str,
    groups: list[dict],
    translator: SessionsTranslator | None = None,
) -> str:
    mod_block = mod_detail
    if mod_simple:
        mod_block = mod_simple + mod_detail
    current_field = ""
    blocks: list[str] = []
    for group in groups:
        if group["field"] and group["field"] != current_field:
            current_field = group["field"]
            fid = group.get("field_slug") or slugify(current_field)
            field_label = (
                translator.field_group(current_field) if translator else current_field
            )
            blocks.append(
                f'<h3 class="sessions-field-title" id="ixtisas-{esc(fid)}">{esc(field_label)}</h3>'
            )
        title = group["title"]
        if translator:
            title = translator.table_title(title)
        blocks.append(
            f'<section class="sessions-masa-block" id="{esc(group["id"])}">'
            f'<h4 class="sessions-masa-title">{esc(title)}</h4>'
            f'{group["table_html"]}'
            f"</section>"
        )
    section_title = (
        UI_EN["ixtisas_tables_title"]
        if translator
        else "İXTİSAS qrupları üzrə masalar"
    )
    lead = UI_EN["moderators_lead"] if translator else "Sessiyaların yönəldiciləri və moderatorlar"
    return f"""
<article class="news-card sessions-tables-card" id="ixtisas-qruplari">
<div class="card-header"><h2 class="card-title">{esc(section_title)}</h2></div>
<div class="card-body">
<p class="card-text">{esc(lead)}</p>
{mod_block}
{"".join(blocks)}
</div>
</article>"""


def build_toc(
    intro: list[dict],
    masas_q: list[dict],
    groups_i: list[dict],
    translator: SessionsTranslator | None = None,
) -> str:
    if translator:
        items = [f'<li><a href="#teskil-prinsipleri">{esc(UI_EN["toc_principles"])}</a></li>']
        items.append(f'<li><a href="#qarisiq-muzakire">{esc(UI_EN["toc_mixed"])}</a></li>')
        items.append(
            f'<li><a href="#strategi-planlasdirma">{esc(UI_EN["toc_strategic"])}</a></li>'
        )
        items.append(f'<li><a href="#ixtisas-muzakire">{esc(UI_EN["toc_specialty"])}</a></li>')
        items.append(
            f'<li><a href="#elm-saheleri-tovsiyeler">{esc(UI_EN["toc_topics"])}</a></li>'
        )
        items.append(
            f'<li class="toc-section-label"><span>{esc(UI_EN["toc_mixed_label"])}</span></li>'
        )
    else:
        items = ['<li><a href="#teskil-prinsipleri">Təşkil prinsipləri</a></li>']
        items.append('<li><a href="#qarisiq-muzakire">QARIŞIQ qruplar</a></li>')
        items.append(
            '<li><a href="#strategi-planlasdirma">Strateji istiqamətlər</a></li>'
        )
        items.append('<li><a href="#ixtisas-muzakire">İXTİSAS qrupları</a></li>')
        items.append(
            '<li><a href="#elm-saheleri-tovsiyeler">Elm sahələri üzrə mövzular</a></li>'
        )
        items.append('<li class="toc-section-label"><span>QARIŞIQ qruplar</span></li>')
    for masa in masas_q:
        title = masa["title"]
        if translator:
            title = translator.table_title(title)
        items.append(
            f'<li class="toc-qarisiq-sub"><a href="#{esc(masa["id"])}">{esc(title)}</a></li>'
        )
    if translator:
        items.append(
            f'<li class="toc-section-label"><span>{esc(UI_EN["toc_specialty_label"])}</span></li>'
        )
    else:
        items.append('<li class="toc-section-label"><span>İXTİSAS qrupları</span></li>')
    current_field = ""
    for group in groups_i:
        if group["field"] and group["field"] != current_field:
            current_field = group["field"]
            fid = group.get("field_slug") or slugify(current_field)
            label = translator.field_group(current_field) if translator else current_field
            if len(label) > 42:
                label = label[:39] + "…"
            items.append(
                f'<li class="toc-ixtisas-group"><a href="#ixtisas-{esc(fid)}">{esc(label)}</a></li>'
            )
        sub_title = group["title"]
        if translator:
            sub_title = translator.table_title(sub_title)
        items.append(
            f'<li class="toc-sub"><a href="#{esc(group["id"])}">{esc(sub_title)}</a></li>'
        )
    return f'<ul class="timeline-list" id="sessionsTOC">{"".join(items)}</ul>'


@dataclass(frozen=True)
class PageTarget:
    mode: str
    asset: str
    lang: str
    out_path: Path

    @property
    def is_preview(self) -> bool:
        return self.mode == "preview"


def nav_strip_placeholder(lang: str, asset: str) -> str:
    home_href = "../../index.html"
    logo_src = f"{asset}images/daab-logo.png"
    if lang == "en":
        return f"""<nav aria-label="Main navigation" class="nav-strip"><div class="nav-inner"><button class="mobile-menu-toggle" type="button" aria-label="Open menu" aria-expanded="false" aria-controls="primaryNavMenu"><span></span><span></span><span></span></button><div class="page-logo"><a title="Home page" aria-label="WAAS home" href="{home_href}"><img src="{logo_src}" class="nav-brand-logo" alt="WAAS Logo"></a></div><a aria-label="WAAS home" class="nav-brand" href="{home_href}"><span class="nav-brand-text">World Association of<br class="mobile-hidden-break">Azerbaijani Scientists</span></a><div class="nav-menu" id="primaryNavMenu" data-daab-nav-placeholder="1"><div class="nav-divider"></div></div><div class="nav-actions" role="group"></div></div></nav>"""
    return f"""<nav aria-label="Əsas naviqasiya" class="nav-strip"><div class="nav-inner"><button class="mobile-menu-toggle" type="button" aria-label="Menyunu aç" aria-expanded="false" aria-controls="primaryNavMenu"><span></span><span></span><span></span></button><div class="page-logo"><a title="Ana səhifə" aria-label="DAAB ana səhifə" href="{home_href}"><img src="{logo_src}" class="nav-brand-logo" alt="DAAB Logo"></a></div><a aria-label="DAAB ana səhifə" class="nav-brand" href="{home_href}"><span class="nav-brand-text">Dünya Azərbaycanlı<br class="mobile-hidden-break">Alimlər Birliyi</span></a><div class="nav-menu" id="primaryNavMenu" data-daab-nav-placeholder="1"><div class="nav-divider"></div></div><div class="nav-actions" role="group"></div></div></nav>"""


def load_preview_nav_html() -> str:
    text = NAV_SOURCE.read_text(encoding="utf-8")
    m = re.search(r"<nav[^>]*>.*?</nav>", text, re.DOTALL)
    if not m:
        return ""
    nav = m.group(0)
    nav = nav.replace('src="../../../images/', 'src="../../images/')
    nav = nav.replace('href="../2026/', 'href="../../az/forum/2026/')
    nav = re.sub(r'href="\.\./\.\./(?!az/)', 'href="../../az/', nav)
    nav = nav.replace('href="../../az/az/', 'href="../../az/')

    def fix_href(match: re.Match[str]) -> str:
        path = match.group(1)
        if path.startswith("http") or path.startswith("#") or path.startswith("../"):
            return match.group(0)
        return f'href="../../az/forum/2024/{path}"'

    nav = re.sub(r'href="(?!https?:|#|\.\./)([^"]+)"', fix_href, nav)
    nav = nav.replace('aria-label="Əsas naviqasiya"', 'aria-label="Əsas naviqasiya (ön baxış)"')
    return nav


def breadcrumbs_html(target: PageTarget) -> str:
    if target.is_preview:
        return """<div class="breadcrumbs forum-breadcrumbs" role="navigation" aria-label="Səhifə yolu">
<a href="../../az/index.html">Ana səhifə</a><span aria-hidden="true">›</span>
<a href="../../az/forum/2024/index.html">Forum 2024</a><span aria-hidden="true">›</span>
<span class="forum-breadcrumbs-current" aria-current="page">Sessiyaların təşkili (ön baxış)</span>
</div>"""
    if target.lang == "en":
        return """<div class="breadcrumbs forum-breadcrumbs" role="navigation" aria-label="Breadcrumb">
<a href="../../index.html">Home</a><span aria-hidden="true">›</span>
<a href="index.html">Forum 2024</a><span aria-hidden="true">›</span>
<span class="forum-breadcrumbs-current" aria-current="page">Sessions</span>
</div>"""
    return """<div class="breadcrumbs forum-breadcrumbs" role="navigation" aria-label="Səhifə yolu"><a href="../../index.html">Ana səhifə</a><span aria-hidden="true">›</span><a href="index.html">Forum 2024</a><span aria-hidden="true">›</span><span class="forum-breadcrumbs-current" aria-current="page">Sessiyaların təşkili</span></div>"""


def footer_html(lang: str) -> str:
    if lang == "en":
        return FORUM_FOOTER_EN
    return """<footer class="footer-pro">
<div class="footer-inner">
<div class="footer-brand"><h3>Dünya Azərbaycanlı Alimlər Birliyi</h3></div>
<div class="footer-grid">
<div class="footer-col"><h4 class="footer-title">Əlaqə</h4><div class="footer-item">✉ <a href="mailto:info@daab-waas.com">info@daab-waas.com</a></div><div class="footer-item">☎ <span>+90 555 147 46 74</span></div><div class="footer-item">🌐 <a href="https://daab-waas.com" rel="noopener noreferrer" target="_blank">daab-waas.com</a></div></div>
<div class="footer-col"><h4 class="footer-title">Ünvan</h4><p class="footer-address">Feneryolu Mahallesi<br/>Gazi Muhtar Paşa Sokak No:44<br/>Kadıköy, İstanbul, Türkiyə</p></div>
<div class="footer-col"><h4 class="footer-title">Rəhbərlik</h4><p class="footer-leader"><strong>Prof. Dr. Məsud Əfəndiyev</strong><br/>DAAB İdarə Heyətinin Sədri<br/>Germany — James D. Murray Distinguished Professor</p></div>
</div>
</div>
<div class="footer-bottom">© 2026 DAAB — Bütün hüquqlar qorunur</div>
</footer>"""


def build_html(
    target: PageTarget,
    intro: list[dict],
    moderators: str,
    masas_q: list[dict],
    mod_s: str,
    mod_d: str,
    groups_i: list[dict],
    translator: SessionsTranslator | None = None,
) -> str:
    st = STYLE_VERSIONS
    sv = SCRIPT_VERSIONS
    asset = target.asset
    nav = load_preview_nav_html() if target.is_preview else nav_strip_placeholder(target.lang, asset)
    page_intro = translator.localize_intro(intro) if translator else intro
    if translator:
        page_intro = [
            localize_structured_section(section) if section.get("blocks") else section
            for section in page_intro
        ]
    intro_html = "".join(render_intro_card(s, asset) for s in page_intro)
    toc = build_toc(intro, masas_q, groups_i, translator)
    meta_description = META_DESCRIPTION_EN if translator else (
        "Forum 2024 — 10 sentyabr strateji sessiyalarının QARIŞIQ və İXTİSAS qrupları üzrə "
        "təşkili, moderatorlar və masa iştirakçıları."
    )
    panel_title = SUMMARY_TITLE_EN if translator else SUMMARY_TITLE
    panel_copy = SUMMARY_EN if translator else SUMMARY

    if target.is_preview:
        title = "Sessiyaların təşkili — DAAB (ön baxış)"
        robots = '<meta name="robots" content="noindex, nofollow"/>'
        page_id = PAGE_ID
        body_class = "forum-sessions-preview-page"
        preview_banner = """<div class="forum-sessions-preview-banner" role="status">
  <strong>Yalnız ön baxış</strong> — canlı səhifə:
  <code>az/forum/2024/sessions_organization.html</code>
</div>"""
        sessions_css = 'href="daab-forum-sessions-preview.css?v=4"'
        skip = '<a class="skip" href="#content">Məzmuna keç</a>'
        hero_h1 = "Sessiyaların <span>təşkili</span>"
        hero_sub = SUBTITLE
        summary_aria = SUMMARY_ARIA
        toc_label = "📋 Mündəricat"
        toc_toggle = 'aria-label="Mündəricat menyusunu aç"'
        activities_v = st.get("daab-activities-layout.css", 16)
        forum_content_v = st.get("daab-forum-content.css", 36)
    elif target.lang == "en":
        title = "Sessions — WAAS"
        robots = ""
        page_id = PAGE_ID
        body_class = ""
        preview_banner = ""
        sessions_css = f'href="{asset}css/daab-forum-sessions.css?v={st.get("daab-forum-sessions.css", 1)}"'
        skip = '<a class="skip" href="#content">Skip to content</a>'
        hero_h1 = "Sessions <span>organization</span>"
        hero_sub = "10 September 2024 — strategic sessions in mixed and discipline-specific groups"
        summary_aria = "Day summary"
        toc_label = "📋 Contents"
        toc_toggle = 'aria-label="Open contents menu"'
        activities_v = st.get("daab-activities-layout.css", 17)
        forum_content_v = st.get("daab-forum-content.css", 37)
    else:
        title = "Sessiyaların təşkili — DAAB"
        robots = ""
        page_id = PAGE_ID
        body_class = ""
        preview_banner = ""
        sessions_css = f'href="{asset}css/daab-forum-sessions.css?v={st.get("daab-forum-sessions.css", 1)}"'
        skip = '<a class="skip" href="#content">Məzmuna keç</a>'
        hero_h1 = "Sessiyaların <span>təşkili</span>"
        hero_sub = SUBTITLE
        summary_aria = SUMMARY_ARIA
        toc_label = "📋 Mündəricat"
        toc_toggle = 'aria-label="Mündəricat menyusunu aç"'
        activities_v = st.get("daab-activities-layout.css", 17)
        forum_content_v = st.get("daab-forum-content.css", 37)

    return f"""<!DOCTYPE html>
<html lang="{target.lang}" data-daab-lang="{target.lang}" data-daab-asset-root="{asset}" data-daab-page-id="{page_id}" data-daab-nav-mount="1">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0, viewport-fit=cover"/>
<title>{esc(title)}</title>
<meta name="description" content="{esc(meta_description)}"/>
{robots}
<link rel="preconnect" href="https://fonts.googleapis.com"/>
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin/>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&family=Playfair+Display:wght@700;800&display=swap" rel="stylesheet"/>
<link href="{asset}css/daab-common.css?v={st["daab-common.css"]}" rel="stylesheet"/>
<link href="{asset}css/daab-perf.css?v={st.get("daab-perf.css", 1)}" rel="stylesheet"/>
<link href="{asset}css/daab-mobile.css?v={st["daab-mobile.css"]}" rel="stylesheet"/>
<link href="{asset}css/daab-sticky-chrome.css?v={st.get("daab-sticky-chrome.css", 1)}" rel="stylesheet"/>
<link href="{asset}css/daab-search.css?v={st["daab-search.css"]}" rel="stylesheet"/>
<link href="{asset}css/daab-back-to-top.css?v=2" rel="stylesheet"/>
<link href="{asset}css/daab-lang.css?v={st["daab-lang.css"]}" rel="stylesheet"/>
<link href="{asset}css/daab-nav-mega.css?v={st["daab-nav-mega.css"]}" rel="stylesheet"/>
<link href="{asset}css/daab-hero-summary.css?v={st["daab-hero-summary.css"]}" rel="stylesheet"/>
<link href="{asset}css/daab-sidebar-widget.css?v={st.get("daab-sidebar-widget.css", 6)}" rel="stylesheet"/>
<link href="{asset}css/daab-activities-layout.css?v={activities_v}" rel="stylesheet"/>
<link href="{asset}css/daab-forum-content.css?v={forum_content_v}" rel="stylesheet"/>
<link rel="stylesheet" {sessions_css}/>
<script src="{asset}js/daab-mobile.js?v={sv["daab-mobile.js"]}" defer></script>
<script src="{asset}js/daab-perf.js?v={sv.get("daab-perf.js", 1)}" defer></script>
<script src="{asset}js/daab-sticky-chrome.js?v={sv.get("daab-sticky-chrome.js", 3)}" defer></script>
<script src="{asset}js/daab-back-to-top.js?v={sv["daab-back-to-top.js"]}" defer></script>
<script src="{asset}js/daab-i18n.js?v={sv["daab-i18n.js"]}" defer></script>
<script src="{asset}js/daab-lang-position.js?v={sv["daab-lang-position.js"]}" defer></script>
<script src="{asset}js/daab-nav.js?v={sv["daab-nav.js"]}" defer></script>
<script src="{asset}js/daab-primary-nav.js?v={sv["daab-primary-nav.js"]}" defer></script>
<script src="{asset}js/daab-shell.js?v={sv["daab-shell.js"]}" defer></script>
<script src="{asset}js/daab-search.js?v={sv["daab-search.js"]}" defer></script>
</head>
<body class="{body_class}">
{preview_banner}
{skip}
{nav}
{breadcrumbs_html(target)}
<header class="page-hero">
<div class="hero-wrap shell">
<section class="hero-copy">
<h1>{hero_h1}</h1>
<p class="page-hero-subtitle" id="page-hero-subtitle" role="doc-subtitle">{esc(hero_sub)}</p>
</section>
<aside aria-label="{esc(summary_aria)}" class="hero-panel">
<div class="panel-card">
<h2 class="panel-title">{esc(panel_title)}</h2>
<div class="panel-copy">
<p class="panel-copy-lead">{esc(panel_copy)}</p>
</div>
</div>
</aside>
</div>
</header>
<div class="content-wrap">
<aside class="sidebar">
<div class="sidebar-widget">
<div class="widget-head">
<span>{toc_label}</span>
<button type="button" class="events-menu-toggle" aria-controls="sessionsTOC" aria-expanded="false" {toc_toggle}><span></span><span></span><span></span></button>
</div>
<div class="widget-body">
{toc}
</div>
</div>
</aside>
<main class="news-feed main" id="content">
{intro_html}
{render_qarisiq_section(moderators, masas_q, translator)}
{render_ixtisas_section(mod_s, mod_d, groups_i, translator)}
</main>
</div>
{footer_html(target.lang)}
<script src="{asset}js/daab-sidebar-timeline.js?v={sv.get('daab-sidebar-timeline.js', 3)}" defer></script>
</body>
</html>"""


def main() -> None:
    if not DOCX.exists():
        raise SystemExit(f"Missing source: {DOCX}")
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    AZ_OUT.parent.mkdir(parents=True, exist_ok=True)
    EN_OUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(DOCX))
    figures = export_images(doc)
    intro = parse_intro(doc, figures)

    targets = [
        PageTarget("preview", PREVIEW_ASSET, "az", OUT_HTML),
        PageTarget("live_az", LIVE_ASSET, "az", AZ_OUT),
        PageTarget("live_en", LIVE_ASSET, "en", EN_OUT),
    ]

    stats: dict[str, int] = {}
    for target in targets:
        photo_index = ParticipantPhotoIndex(target.asset)
        translator = None
        if target.lang == "en":
            translator = SessionsTranslator(photo_index.profiles_by_key, session_person_key)
        moderators, masas_q = parse_qarisiq_tables(doc, photo_index, translator)
        mod_s, mod_d, groups_i = parse_ixtisas_tables(doc, photo_index, translator)
        html_out = build_html(
            target,
            intro,
            moderators,
            masas_q,
            mod_s,
            mod_d,
            groups_i,
            translator,
        )
        target.out_path.write_text(html_out, encoding="utf-8", newline="\n")
        stats[target.mode] = html_out.count("sessions-participant-photo")

    embed = ROOT / "helpers/_embed_static_nav.py"
    subprocess.run([sys.executable, str(embed)], check=True, cwd=ROOT)

    print(f"Wrote {OUT_HTML.relative_to(ROOT)}")
    print(f"Wrote {AZ_OUT.relative_to(ROOT)}")
    print(f"Wrote {EN_OUT.relative_to(ROOT)}")
    print(f"  Intro sections: {len(intro)}")
    print(f"  Figures: {list(figures.keys())}")
    print(f"  Participant photos (live AZ): {stats.get('live_az', 0)}")


if __name__ == "__main__":
    main()
