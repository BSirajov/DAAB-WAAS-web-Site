#!/usr/bin/env python3
"""Export Forum 2024 Strategic Implementation Framework to a formatted Word document."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

try:
    from _paths import ROOT
except ImportError:
    ROOT = Path(__file__).resolve().parents[1]

MD_SOURCE = ROOT / "documents" / "DAAB-Forum-2024-Strategic-Implementation-Framework.md"
DOCX_OUT = ROOT / "documents" / "DAAB-Forum-2024-Strategic-Implementation-Framework.docx"

# WAAS / site palette
NAVY = RGBColor(0x09, 0x4D, 0x78)
BLUE = RGBColor(0x00, 0x69, 0xB4)
GOLD = RGBColor(0xC9, 0xA2, 0x27)
INK = RGBColor(0x1A, 0x2E, 0x3D)
MUTED = RGBColor(0x34, 0x5D, 0x76)
TABLE_HEADER = "094D78"
TABLE_ALT = "F3F9FD"
FONT_BODY = "Calibri"
FONT_HEADING = "Calibri Light"


def set_update_fields_on_open(document: Document) -> None:
    settings = document.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


def add_page_number_field(paragraph, align_center: bool = True) -> None:
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)

    run4 = paragraph.add_run("1")
    run5 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)


def add_table_of_contents(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)

    run4 = paragraph.add_run("Right-click here and choose Update Field → Update entire table.")
    run4.font.italic = True
    run4.font.color.rgb = MUTED
    run4.font.size = Pt(10)

    run5 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)


def setup_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    for level, size, color in [
        ("Heading 1", 18, NAVY),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 12, NAVY),
    ]:
        h = document.styles[level]
        h.font.name = FONT_HEADING
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = color
        h.paragraph_format.space_before = Pt(14 if level == "Heading 1" else 10)
        h.paragraph_format.space_after = Pt(8)
        h.paragraph_format.keep_with_next = True


def shade_cell(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_run(paragraph, text: str, bold: bool = False, italic: bool = False) -> None:
    """Parse **bold** and *italic* within a text segment."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:
            r = paragraph.add_run(part)
            if bold:
                r.bold = True
            if italic:
                r.italic = True


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    cells = [c.strip() for c in line.split("|")]
    return cells


def is_table_separator(line: str) -> bool:
    s = line.strip().replace("|", "").replace(":", "").replace("-", "").strip()
    return not s


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.autofit = True

    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.rows[ri].cells[ci]
            text = row[ci] if ci < len(row) else ""
            text = text.replace("<br>", "\n").replace("<br/>", "\n")
            cell.text = ""
            p = cell.paragraphs[0]
            for j, segment in enumerate(text.split("\n")):
                if j:
                    p.add_run("\n")
                add_formatted_run(p, segment)
            if ri == 0:
                shade_cell(cell, TABLE_HEADER)
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)
            else:
                for run in p.runs:
                    run.font.size = Pt(10)
                if ri % 2 == 0:
                    shade_cell(cell, TABLE_ALT)


def setup_header_footer(document: Document, title_short: str) -> None:
    section = document.sections[0]
    section.different_first_page_header_footer = True
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default header (pages 2+)
    header = section.header
    header._element.clear()
    hp = header.add_paragraph()
    run_left = hp.add_run("WAAS — Forum 2024 Strategic Implementation Framework")
    run_left.font.size = Pt(9)
    run_left.font.color.rgb = MUTED
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Blank first-page header
    first_header = section.first_page_header
    first_header._element.clear()
    first_header.add_paragraph()

    # Default footer with page number
    footer = section.footer
    footer._element.clear()
    fp2 = footer.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp2.add_run("World Association of Azerbaijani Scientists   |   Page ")
    r1.font.size = Pt(9)
    r1.font.color.rgb = MUTED
    add_page_number_field(fp2, align_center=False)
    r2 = fp2.add_run("   |   May 2026   |   Partner shareable")
    r2.font.size = Pt(9)
    r2.font.color.rgb = MUTED

    # First page header/footer minimal
    first_footer = section.first_page_footer
    first_footer._element.clear()
    fp_first = first_footer.add_paragraph()
    fp_first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = fp_first.add_run("daab-waas.com")
    rf.font.size = Pt(9)
    rf.font.color.rgb = GOLD


def add_title_page(document: Document) -> None:
    for _ in range(4):
        document.add_paragraph()

    t = document.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("WORLD ASSOCIATION OF")
    r.font.size = Pt(14)
    r.font.color.rgb = MUTED
    r.font.small_caps = True

    t2 = document.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("AZERBAIJANI SCIENTISTS")
    r2.font.size = Pt(14)
    r2.font.color.rgb = MUTED
    r2.font.small_caps = True

    document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title.add_run("Forum 2024")
    rt.font.name = FONT_HEADING
    rt.font.size = Pt(28)
    rt.font.bold = True
    rt.font.color.rgb = NAVY

    title2 = document.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt2 = title2.add_run("Strategic Implementation Framework")
    rt2.font.name = FONT_HEADING
    rt2.font.size = Pt(22)
    rt2.font.bold = True
    rt2.font.color.rgb = BLUE

    # Gold rule
    rule = document.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rule.add_run("━" * 42)
    rr.font.color.rgb = GOLD

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(
        "From roadmap themes to implementable projects,\n"
        "action plans, and measurable national impact"
    )
    rs.font.size = Pt(12)
    rs.font.italic = True
    rs.font.color.rgb = MUTED

    document.add_paragraph()
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for label, value in [
        ("Document type", "Institutional strategic implementation plan"),
        ("Version", "1.0 — May 2026"),
        ("Status", "Working framework for leadership review"),
        ("Classification", "Internal / partner shareable"),
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(f"{label}: ")
        rl.font.size = Pt(10)
        rl.font.bold = True
        rl.font.color.rgb = NAVY
        rv = p.add_run(value)
        rv.font.size = Pt(10)
        rv.font.color.rgb = INK

    document.add_page_break()


def add_toc_page(document: Document) -> None:
    h = document.add_heading("Table of contents", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note = document.add_paragraph()
    add_formatted_run(
        note,
        "After opening in Microsoft Word, click the table below, press **F9**, "
        "or right-click → **Update Field** → **Update entire table** to refresh page numbers.",
    )
    note.paragraph_format.space_after = Pt(12)
    toc_p = document.add_paragraph()
    add_table_of_contents(toc_p)
    document.add_page_break()


def convert_markdown_body(document: Document, text: str) -> None:
    lines = text.splitlines()
    i = 0
    table_buffer: list[list[str]] = []
    in_code = False
    code_lines: list[str] = []

    # Skip YAML-like title block at top (handled by title page)
    while i < len(lines) and not lines[i].startswith("## Executive"):
        if lines[i].startswith("## "):
            break
        i += 1

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = document.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = MUTED
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            if table_buffer and is_table_separator(stripped):
                i += 1
                continue
            row = parse_table_row(stripped)
            if row:
                table_buffer.append(row)
            i += 1
            continue
        elif table_buffer:
            add_table(document, table_buffer)
            table_buffer = []

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            add_formatted_run(p, stripped[2:].strip())
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = document.add_paragraph(style="List Number")
            add_formatted_run(p, re.sub(r"^\d+\.\s*", "", stripped))
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = document.add_paragraph()
            add_formatted_run(p, stripped.strip("*"))
            p.runs[-1].italic = True if p.runs else None
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("*End of document*"):
            i += 1
            continue

        p = document.add_paragraph()
        add_formatted_run(p, stripped)
        i += 1

    if table_buffer:
        add_table(document, table_buffer)


def build_document() -> Path:
    if not MD_SOURCE.exists():
        raise FileNotFoundError(MD_SOURCE)

    md_text = MD_SOURCE.read_text(encoding="utf-8")

    document = Document()
    setup_styles(document)
    setup_header_footer(document, "Forum 2024 Implementation")
    add_title_page(document)
    add_toc_page(document)
    convert_markdown_body(document, md_text)
    set_update_fields_on_open(document)

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(DOCX_OUT))
    return DOCX_OUT


def main() -> None:
    out = build_document()
    print(f"Wrote {out.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
