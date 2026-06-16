#!/usr/bin/env python3
"""Build standalone preview HTML from the inventions DOCX (not live site)."""
from __future__ import annotations

import html
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from _paths import ROOT

DOCX = (
    ROOT
    / "documents"
    / "docx"
    / "Most_Influential_Scientific_Inventions_and_Innovations (with references).docx"
)
OUT_DIR = ROOT / "documents" / "preview"
OUT_HTML = OUT_DIR / "major_scientific_inventions_expanded.html"
OUT_IMAGES = OUT_DIR / "images"
CARD_DATA = OUT_DIR / "inventions-card-data.json"
CARD_OVERRIDES = OUT_DIR / "inventions-card-overrides.json"
NAV_SOURCE = ROOT / "en" / "major_scientific_inventions.html"

SECTION_LABELS = {
    "HISTORICAL CONTEXT": "Historical context",
    "SCIENTIFIC / TECHNOLOGICAL SIGNIFICANCE": "Scientific significance",
    "IMPACT ON HUMANITY": "Impact on humanity",
    "MULTIMEDIA & FURTHER RESOURCES": "Further resources",
}

SKIP_H1 = {"contents", "overview by category"}

URL_RE = re.compile(r"https?://[^\s\]\)>,]+")


def slugify(text: str) -> str:
    s = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")
    return s or "entry"


def esc(text: str) -> str:
    return html.escape(text, quote=True)


def linkify_text(text: str) -> str:
    parts: list[str] = []
    last = 0
    for match in URL_RE.finditer(text):
        parts.append(esc(text[last : match.start()]))
        url = match.group(0).rstrip(".,;)]")
        parts.append(
            f'<a class="resource-link" href="{esc(url)}" target="_blank" '
            f'rel="noopener noreferrer">{esc(url)}</a>'
        )
        last = match.end()
    parts.append(esc(text[last:]))
    return "".join(parts)


def parse_resource_line(line: str) -> tuple[str, str]:
    line = line.lstrip("▸ ").strip()
    if " — " in line:
        parts = [p.strip() for p in line.split(" — ")]
        if len(parts) >= 2:
            return parts[0], parts[-1]
    return line, ""


def extract_para_url(para: Paragraph) -> str:
    for hyperlink in para._element.iter(qn("w:hyperlink")):
        rel_id = hyperlink.get(qn("r:id"))
        if rel_id and para.part.rels.get(rel_id):
            return para.part.rels[rel_id].target_ref
    return ""


def extract_para_image(para: Paragraph) -> str:
    for blip in para._element.findall(".//" + qn("a:blip")):
        rel_id = blip.get(qn("r:embed"))
        if rel_id and para.part.rels.get(rel_id):
            return para.part.rels[rel_id].target_ref
    return ""


def parse_resource_item(para: Paragraph) -> dict[str, str]:
    text = para.text.strip()
    title, source = parse_resource_line(text)
    return {"title": title, "source": source, "url": extract_para_url(para)}


def parse_docx() -> dict:
    doc = Document(DOCX)
    data: dict = {
        "title": "",
        "subtitle": "",
        "tagline": "",
        "executive_summary": "",
        "categories": [],
        "entries": [],
        "overview_table": [],
        "top20_table": [],
        "conclusion": {"title": "", "paragraphs": []},
        "references": {"title": "", "intro": "", "groups": []},
    }

    paras = doc.paragraphs
    if paras:
        data["title"] = paras[0].text.strip()
    if len(paras) > 1:
        data["subtitle"] = paras[1].text.strip()
    if len(paras) > 2:
        data["tagline"] = paras[2].text.strip()
    for para in paras[3:8]:
        t = para.text.strip()
        if t.upper().startswith("EXECUTIVE SUMMARY"):
            data["executive_summary"] = re.sub(
                r"^EXECUTIVE SUMMARY\s*",
                "",
                t,
                flags=re.IGNORECASE,
            ).strip()
            break

    current_cat: dict | None = None
    current_entry: dict | None = None
    current_section: str | None = None
    current_ref_group: dict | None = None
    in_conclusion = False
    in_references = False

    def flush_entry() -> None:
        nonlocal current_entry
        if current_entry and current_cat:
            current_cat["entries"].append(current_entry)
        current_entry = None

    for para in paras:
        text = para.text.strip()
        style = para.style.name if para.style else "Normal"

        if (
            current_entry is not None
            and not in_conclusion
            and not in_references
            and style != "Heading 2"
        ):
            image_ref = extract_para_image(para)
            if image_ref:
                current_entry["image"] = image_ref
                continue

        if not text:
            continue

        if style == "Heading 1":
            key = text.lower().lstrip("0123456789. ").strip()
            if key in SKIP_H1:
                current_entry = None
                current_section = None
                in_conclusion = False
                continue
            if key.startswith("concise top 20"):
                current_cat = None
                current_entry = None
                in_conclusion = False
                continue
            if key == "concluding note":
                in_conclusion = True
                in_references = False
                current_cat = None
                current_entry = None
                current_ref_group = None
                data["conclusion"]["title"] = text
                continue
            if key == "references":
                in_references = True
                in_conclusion = False
                current_cat = None
                current_entry = None
                current_ref_group = None
                data["references"]["title"] = text
                continue
            in_conclusion = False
            in_references = False
            current_ref_group = None
            flush_entry()
            current_cat = {"title": text, "slug": slugify(text), "entries": []}
            data["categories"].append(current_cat)
            current_section = None
            continue

        if in_conclusion:
            data["conclusion"]["paragraphs"].append(text)
            continue

        if in_references:
            if style == "Heading 3":
                current_ref_group = {
                    "title": text,
                    "slug": slugify(text),
                    "items": [],
                }
                data["references"]["groups"].append(current_ref_group)
                continue
            if style == "Normal":
                if current_ref_group is None and not data["references"]["intro"]:
                    data["references"]["intro"] = text
                elif current_ref_group is not None:
                    current_ref_group["items"].append(text)
            continue

        if style == "Heading 2":
            flush_entry()
            current_entry = {
                "title": text,
                "slug": slugify(text),
                "meta": "",
                "image": "",
                "image_src": "",
                "sections": {},
                "category": current_cat["title"] if current_cat else "",
            }
            data["entries"].append(current_entry)
            current_section = None
            continue

        if current_entry is None:
            continue

        if text in SECTION_LABELS:
            current_section = text
            current_entry["sections"].setdefault(current_section, [])
            continue

        if current_section:
            if current_section == "MULTIMEDIA & FURTHER RESOURCES" and text.startswith("▸"):
                current_entry["sections"][current_section].append(parse_resource_item(para))
            else:
                current_entry["sections"][current_section].append(text)
        elif not current_entry["meta"] and text.lower().startswith("period:"):
            current_entry["meta"] = text

    flush_entry()

    for cat_index, cat in enumerate(data["categories"], start=1):
        cat_num = _cat_number(cat["title"]) or str(cat_index)
        for entry_index, entry in enumerate(cat["entries"], start=1):
            entry["number"] = f"{cat_num}.{entry_index}"

    for table in doc.tables:
        rows = [[c.text.strip() for c in row.cells] for row in table.rows]
        if not rows:
            continue
        header = [h.lower() for h in rows[0]]
        if "category" in header[0] and "items" in " ".join(header):
            data["overview_table"] = rows
        elif "rank" in header[0] and "breakthrough" in " ".join(header):
            data["top20_table"] = rows

    return data


def meta_key_figures(meta: str) -> str:
    if "|" not in meta:
        return ""
    tail = meta.split("|", 1)[1].strip()
    return re.sub(r"^Key figure\(s\):\s*", "", tail, flags=re.IGNORECASE).strip()


def load_card_data() -> dict[str, dict]:
    if not CARD_DATA.exists():
        return {}
    return json.loads(CARD_DATA.read_text(encoding="utf-8"))


def load_card_overrides() -> dict[str, dict]:
    if not CARD_OVERRIDES.exists():
        return {}
    return json.loads(CARD_OVERRIDES.read_text(encoding="utf-8"))


def card_key_figures_clean(text: str) -> bool:
    if not text:
        return False
    if re.search(r"KEY\s*FACT|ikebdi|CHk\s*EY|MAcasmerica|ventional|ssachusetts", text, re.IGNORECASE):
        return False
    return len(text) <= 140


def attach_card_content(
    data: dict, cards: dict[str, dict], overrides: dict[str, dict] | None = None
) -> None:
    overrides = overrides or {}
    for entry in data["entries"]:
        card = cards.get(entry["slug"], {})
        override = overrides.get(entry["slug"], {})
        meta_figures = meta_key_figures(entry.get("meta", ""))
        ocr_figures = card.get("key_figures", "")
        if override.get("key_figures"):
            ocr_figures = override["key_figures"]
        entry["card"] = card
        entry["card_key_figures"] = (
            ocr_figures if card_key_figures_clean(ocr_figures) else meta_figures
        )
        entry["card_summary"] = override.get("summary") or card.get("summary", "")
        entry["card_facts"] = override.get("key_facts") or card.get("key_facts", [])
        entry["card_icon"] = override.get("icon") or card.get("icon", "")


def render_entry_visual(entry: dict) -> str:
    icon = entry.get("card_icon", "")
    if not icon and not entry.get("card_summary") and not entry.get("card_facts"):
        return ""

    parts = ['<div class="inventions-entry-visual">']
    if icon:
        parts.append(
            '<figure class="inventions-entry-icon">'
            f'<img src="{esc(icon)}" alt="{esc("Illustration: " + entry["title"])}" '
            'loading="lazy" decoding="async"/>'
            f'<figcaption class="inventions-entry-icon-caption">'
            f'Illustration: {esc(entry["title"])}</figcaption>'
            "</figure>"
        )

    copy_bits: list[str] = []
    if entry.get("card_key_figures"):
        copy_bits.append(
            '<p class="inventions-entry-visual-figures">'
            f'<strong>Key figure(s):</strong> {esc(entry["card_key_figures"])}</p>'
        )
    if entry.get("card_summary"):
        copy_bits.append(
            f'<p class="inventions-entry-visual-summary">{esc(entry["card_summary"])}</p>'
        )
    if entry.get("card_facts"):
        items = "".join(f"<li>{esc(fact)}</li>" for fact in entry["card_facts"])
        copy_bits.append(
            '<div class="inventions-key-facts">'
            "<h4>Key facts</h4>"
            f"<ul>{items}</ul>"
            "</div>"
        )

    if copy_bits:
        parts.append('<div class="inventions-entry-visual-copy">' + "".join(copy_bits) + "</div>")
    parts.append("</div>")
    return "\n".join(parts)


def export_entry_images(data: dict) -> int:
    OUT_IMAGES.mkdir(parents=True, exist_ok=True)
    exported = 0
    with zipfile.ZipFile(DOCX) as archive:
        available = set(archive.namelist())
        for entry in data["entries"]:
            media_path = entry.get("image", "")
            if not media_path:
                continue
            zip_path = f"word/{media_path}"
            if zip_path not in available:
                continue
            ext = Path(media_path).suffix or ".png"
            dest = OUT_IMAGES / f"{entry['slug']}{ext}"
            dest.write_bytes(archive.read(zip_path))
            entry["image_src"] = f"images/{dest.name}"
            exported += 1
    return exported


def load_nav_html() -> str:
    text = NAV_SOURCE.read_text(encoding="utf-8")
    m = re.search(r"<nav[^>]*>.*?</nav>", text, re.DOTALL)
    if not m:
        return ""
    nav = m.group(0)
    nav = nav.replace('href="', 'href="../../en/')
    nav = nav.replace('src="../images/', 'src="../../images/')
    nav = nav.replace('aria-label="Main navigation"', 'aria-label="Main navigation (preview)"')
    return nav


def render_resources(items: list[dict[str, str] | str]) -> str:
    lis: list[str] = []
    for item in items:
        if isinstance(item, str):
            title, source = parse_resource_line(item)
            url = ""
        else:
            title = item.get("title", "")
            source = item.get("source", "")
            url = item.get("url", "")
        if url:
            label = (
                f'<a class="resource-link" href="{esc(url)}" target="_blank" '
                f'rel="noopener noreferrer">{esc(title)}</a>'
            )
        else:
            label = f'<span class="resource-title">{esc(title)}</span>'
        if source:
            lis.append(
                f"<li>{label} <span class=\"resource-source\">— {esc(source)}</span></li>"
            )
        else:
            lis.append(f"<li>{label}</li>")
    return f'<ul class="inventions-resources">{"".join(lis)}</ul>'


def render_entry(entry: dict) -> str:
    number = entry.get("number", "")
    card_bits = " ".join(
        [
            entry.get("card_key_figures", ""),
            entry.get("card_summary", ""),
            " ".join(entry.get("card_facts", [])),
        ]
    )
    search_blob = f"{number} {entry['title']} {entry.get('meta', '')} {card_bits}".strip()
    if number:
        title_html = (
            '<h2 class="inventions-entry-title">'
            f'<span class="inventions-entry-num" aria-hidden="true">{esc(number)}</span>'
            f'<span class="inventions-entry-name">{esc(entry["title"])}</span></h2>'
        )
    else:
        title_html = (
            '<h2 class="inventions-entry-title">'
            f'<span class="inventions-entry-name">{esc(entry["title"])}</span></h2>'
        )
    parts = [
        f'<article class="inventions-entry" id="{esc(entry["slug"])}" '
        f'data-search="{esc(search_blob)}">',
        title_html,
    ]
    visual_html = render_entry_visual(entry)
    if visual_html:
        parts.append(visual_html)
    if entry.get("meta"):
        parts.append(f'<p class="inventions-entry-meta">{esc(entry["meta"])}</p>')
    for key, label in SECTION_LABELS.items():
        lines = entry["sections"].get(key, [])
        if not lines:
            continue
        parts.append('<div class="inventions-entry-section">')
        parts.append(f"<h3>{esc(label)}</h3>")
        if key == "MULTIMEDIA & FURTHER RESOURCES":
            parts.append(render_resources(lines))
        else:
            for line in lines:
                parts.append(f"<p>{esc(line)}</p>")
        parts.append("</div>")
    parts.append("</article>")
    return "\n".join(parts)


def _cat_number(title: str) -> str:
    m = re.match(r"(\d+)", title.strip())
    return m.group(1) if m else ""


def _cat_label(title: str) -> str:
    return re.sub(r"^\d+\.\s*", "", title).strip()


def render_toc(data: dict) -> str:
    items = ['<ul class="timeline-list" id="inventionsTocList">']
    for cat in data["categories"]:
        if not cat["entries"]:
            continue
        num = _cat_number(cat["title"])
        label = _cat_label(cat["title"])
        items.append(
            f'<li class="inventions-toc-cat-row" data-toc-cat="{esc(cat["slug"])}">'
            f'<span class="tl-date">§{esc(num)}</span>'
            f'<a href="#{esc(cat["slug"])}">{esc(label)}</a></li>'
        )
        for entry in cat["entries"]:
            number = entry.get("number", "")
            items.append(
                f'<li class="inventions-toc-entry" data-toc-entry="{esc(entry["slug"])}" '
                f'data-toc-cat="{esc(cat["slug"])}" '
                f'data-search="{esc((number + " " if number else "") + entry["title"])}">'
                f'<span class="tl-date">{esc(number) if number else "·"}</span>'
                f'<a href="#{esc(entry["slug"])}">{esc(entry["title"])}</a></li>'
            )
    items.append(
        '<li class="inventions-toc-extra"><span class="tl-date">⊕</span>'
        '<a href="#overview-by-category">Overview by category</a></li>'
    )
    if data["conclusion"]["title"]:
        items.append(
            f'<li class="inventions-toc-extra"><span class="tl-date">⊕</span>'
            f'<a href="#concluding-note">{esc(data["conclusion"]["title"])}</a></li>'
        )
    if data["references"]["title"]:
        items.append(
            f'<li class="inventions-toc-extra"><span class="tl-date">⊕</span>'
            f'<a href="#references">{esc(data["references"]["title"])}</a></li>'
        )
        for group in data["references"]["groups"]:
            label = re.sub(r"^[A-J]\.\s*", "", group["title"]).strip()
            items.append(
                f'<li class="inventions-toc-ref-row" data-search="{esc(group["title"])}">'
                f'<span class="tl-date">↳</span>'
                f'<a href="#{esc(group["slug"])}">{esc(label)}</a></li>'
            )
    items.append("</ul>")
    return "\n".join(items)


def render_references(data: dict) -> str:
    refs = data["references"]
    if not refs["title"] and not refs["groups"]:
        return ""

    parts = [
        '<section class="inventions-references" id="references">',
        f'<h2>{esc(refs["title"] or "References")}</h2>',
    ]
    if refs.get("intro"):
        parts.append(f'<p class="inventions-references-intro">{esc(refs["intro"])}</p>')

    for group in refs["groups"]:
        parts.append(
            f'<div class="inventions-references-group" id="{esc(group["slug"])}">'
            f'<h3>{esc(group["title"])}</h3>'
            '<ol class="inventions-bibliography">'
        )
        for item in group["items"]:
            parts.append(f"<li>{linkify_text(item)}</li>")
        parts.append("</ol></div>")

    parts.append("</section>")
    return "\n".join(parts)


def render_table(rows: list[list[str]], caption: str) -> str:
    if not rows:
        return ""
    head, body = rows[0], rows[1:]
    th = "".join(f"<th>{esc(c)}</th>" for c in head)
    trs = []
    for row in body:
        tds = "".join(f"<td>{esc(c)}</td>" for c in row)
        trs.append(f"<tr>{tds}</tr>")
    return f"""
<div class="inventions-table-wrap">
  <h2>{esc(caption)}</h2>
  <table class="inventions-data-table">
    <thead><tr>{th}</tr></thead>
    <tbody>{"".join(trs)}</tbody>
  </table>
</div>"""


def build_html(data: dict) -> str:
    nav = load_nav_html()
    categories_html = []
    for cat in data["categories"]:
        if not cat["entries"]:
            continue
        entries_html = "\n".join(render_entry(e) for e in cat["entries"])
        categories_html.append(
            f'<section class="inventions-category" id="{esc(cat["slug"])}" '
            f'data-category="{esc(cat["title"])}">'
            f'<h2 class="inventions-category-head">{esc(cat["title"])}</h2>'
            f"{entries_html}</section>"
        )

    conclusion_html = ""
    if data["conclusion"]["paragraphs"]:
        ps = "".join(f"<p>{esc(p)}</p>" for p in data["conclusion"]["paragraphs"])
        conclusion_html = f"""
<section class="inventions-conclusion" id="concluding-note">
  <h2>{esc(data["conclusion"]["title"] or "Concluding note")}</h2>
  {ps}
</section>"""

    exec_summary = esc(data["executive_summary"] or data["tagline"])

    return f"""<!DOCTYPE html>
<html lang="en" data-daab-lang="en" data-daab-asset-root="../../" data-daab-page-id="major-scientific-inventions">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0, viewport-fit=cover"/>
<title>WAAS — Major Scientific Inventions (Preview)</title>
<meta name="description" content="Sixty landmark inventions, discoveries, and innovations that shaped human history — expanded reference with bibliography preview."/>
<meta name="robots" content="noindex, nofollow"/>
<link href="../../css/daab-fonts.css?v=1" rel="stylesheet"/>
<link href="../../css/daab-common.css?v=67" rel="stylesheet"/>
<link href="../../css/daab-perf.css?v=1" rel="stylesheet"/>
<link href="../../css/daab-mobile.css?v=13" rel="stylesheet"/>
<link href="../../css/daab-sticky-chrome.css?v=1" rel="stylesheet"/>
<link href="../../css/daab-search.css?v=4" rel="stylesheet"/>
<link href="../../css/daab-back-to-top.css?v=2" rel="stylesheet"/>
<link href="../../css/daab-lang.css?v=12" rel="stylesheet"/>
<link href="../../css/daab-nav-mega.css?v=69" rel="stylesheet"/>
<link href="../../css/daab-hero-summary.css?v=12" rel="stylesheet"/>
<link href="../../css/daab-sidebar-widget.css?v=6" rel="stylesheet"/>
<link href="../../css/daab-charter-page.css?v=9" rel="stylesheet"/>
<link href="daab-inventions-preview.css?v=14" rel="stylesheet"/>
<script src="../../js/daab-mobile.js?v=6" defer></script>
<script src="../../js/daab-perf.js?v=1" defer></script>
<script src="../../js/daab-sticky-chrome.js?v=3" defer></script>
<script src="../../js/daab-back-to-top.js?v=3" defer></script>
<script src="../../js/daab-i18n.js?v=28" defer></script>
<script src="../../js/daab-lang-position.js?v=7" defer></script>
<script src="../../js/daab-design-tokens.js?v=1" defer></script>
<script src="../../js/daab-nav.js?v=31" defer></script>
<script src="../../js/daab-primary-nav.js?v=46" defer></script>
<script src="../../js/daab-breadcrumbs.js?v=30" defer></script>
<script src="../../js/daab-shell.js?v=12" defer></script>
<script src="daab-inventions-preview.js?v=5" defer></script>
</head>
<body class="charter-page inventions-preview-page">
<div class="inventions-preview-banner" role="status">
  <strong>Preview only</strong> — not linked in site navigation. Review before publishing to <code>en/major_scientific_inventions.html</code>.
</div>
<a class="skip" href="#content">Skip to content</a>
{nav}
<header class="hero">
<div class="hero-wrap shell inventions-hero-wrap">
<section class="hero-copy">
<h1>Major scientific <span>inventions</span></h1>
<p class="page-hero-subtitle" id="page-hero-subtitle" role="doc-subtitle">{esc(data["tagline"] or data["subtitle"] or "Discoveries and innovations in human history")}</p>
</section>
<aside aria-label="Executive summary" class="hero-panel">
<div class="inventions-summary-card">
<h2 class="panel-title">Executive summary</h2>
<p class="hero-text panel-copy-lead">{exec_summary}</p>
</div>
</aside>
</div>
</header>
<main class="main" id="content">
<div class="inventions-toolbar" aria-label="Search inventions">
<div class="inventions-search-wrap">
<svg fill="none" height="15" stroke="currentColor" stroke-width="2" viewBox="0 0 24 24" width="15" aria-hidden="true"><circle cx="11" cy="11" r="8"></circle><line x1="21" x2="16.65" y1="21" y2="16.65"></line></svg>
<input type="search" id="inventionsSearch" placeholder="Search inventions, periods, or key figures…" aria-label="Search inventions"/>
</div>
</div>
<div class="charter-layout inventions-layout">
<aside class="charter-sidebar toc-card" aria-label="Article navigation">
<div class="sidebar-widget events-open" id="inventionsArticlesWidget">
<div class="widget-head">
<span>📚 Articles</span>
<button type="button" class="events-menu-toggle" aria-expanded="true" aria-controls="inventionsTocList" aria-label="Toggle article menu"><span></span><span></span><span></span></button>
</div>
<div class="widget-body">
{render_toc(data)}
</div>
</div>
</aside>
<div class="charter-stack inventions-stack">
{"".join(categories_html)}
<div id="overview-by-category">{render_table(data["overview_table"], "Overview by category")}</div>
{conclusion_html}
{render_references(data)}
</div>
</div>
</main>
<footer class="footer-pro">
<div class="footer-inner">
<div class="footer-brand"><h3>World Association of Azerbaijani Scientists</h3></div>
<div class="footer-grid">
<div class="footer-col"><div class="footer-title">Contact</div><div class="footer-item">✉ <a href="mailto:info@daab-waas.com">info@daab-waas.com</a></div><div class="footer-item">☎ <span>+90 555 147 46 74</span></div><div class="footer-item">🌐 <a href="https://daab-waas.com" target="_blank" rel="noopener">daab-waas.com</a></div></div>
<div class="footer-col"><div class="footer-title">Address</div><p class="footer-address">Feneryolu Mahallesi<br/>Gazi Muhtar Paşa Sokak No:44<br/>Kadıköy, Istanbul, Türkiye</p></div>
<div class="footer-col"><div class="footer-title">Leadership</div><p class="footer-leader"><strong>Prof. Dr. Messoud Efendiyev</strong><br/>Chair of the WAAS Executive Board</p></div>
</div>
</div>
<div class="footer-bottom">© 2026 DAAB / WAAS — All Rights Reserved</div>
</footer>
</body>
</html>
"""


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = parse_docx()
    attach_card_content(data, load_card_data(), load_card_overrides())
    image_count = export_entry_images(data)
    html_out = build_html(data)
    OUT_HTML.write_text(html_out, encoding="utf-8", newline="\n")
    print(f"Wrote {OUT_HTML.relative_to(ROOT)}")
    print(f"  Categories: {len(data['categories'])}")
    print(f"  Entries: {len(data['entries'])}")
    print(f"  Images: {image_count}")


if __name__ == "__main__":
    main()
