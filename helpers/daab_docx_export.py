"""
DAAB / WAAS — Markdown to publication-ready Word (.docx) export utilities.
"""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# WAAS / site palette
NAVY = RGBColor(0x09, 0x4D, 0x78)
BLUE = RGBColor(0x00, 0x69, 0xB4)
GOLD = RGBColor(0xC9, 0xA2, 0x27)
INK = RGBColor(0x1A, 0x2E, 0x3D)
MUTED = RGBColor(0x34, 0x5D, 0x76)
TABLE_HEADER = "094D78"
TABLE_ALT = "F3F9FD"
CODE_FILL = "F4F7FA"
QUOTE_FILL = "EEF8FF"
FONT_BODY = "Calibri"
FONT_HEADING = "Calibri Light"
ORG_NAME_LINE1 = "WORLD ASSOCIATION OF"
ORG_NAME_LINE2 = "AZERBAIJANI SCIENTISTS"
ORG_SHORT = "WAAS / DAAB"
SITE_LABEL = "daab-waas.com"

LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def set_update_fields_on_open(document: Document) -> None:
    settings = document.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


def add_page_number_field(paragraph, align_center: bool = True) -> None:
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)

    run4 = paragraph.add_run("1")
    run5 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)


def add_table_of_contents(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)

    run4 = paragraph.add_run("Open in Word, then press F9 or Update Field to refresh.")
    run4.font.italic = True
    run4.font.color.rgb = MUTED
    run4.font.size = Pt(10)

    run5 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)


def setup_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    for level, size, color in [
        ("Heading 1", 18, NAVY),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 12, NAVY),
        ("Heading 4", 11, BLUE),
    ]:
        h = document.styles[level]
        h.font.name = FONT_HEADING
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = color
        h.paragraph_format.space_before = Pt(14 if level == "Heading 1" else 10)
        h.paragraph_format.space_after = Pt(8)
        h.paragraph_format.keep_with_next = True

    try:
        quote = document.styles["Quote"]
        quote.font.name = FONT_BODY
        quote.font.size = Pt(11)
        quote.font.italic = True
        quote.font.color.rgb = MUTED
        quote.paragraph_format.left_indent = Cm(1)
        quote.paragraph_format.space_before = Pt(6)
        quote.paragraph_format.space_after = Pt(6)
    except KeyError:
        pass


def shade_paragraph(paragraph, fill_hex: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    p_pr.append(shading)


def shade_cell(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0069B4")
    r_pr.append(color)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_BODY)
    r_fonts.set(qn("w:hAnsi"), FONT_BODY)
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    r_pr.append(sz)
    new_run.append(r_pr)

    text_el = OxmlElement("w:t")
    text_el.text = text
    text_el.set(qn("xml:space"), "preserve")
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_formatted_run(paragraph, text: str, bold: bool = False, italic: bool = False) -> None:
    """Parse [links](url), **bold**, *italic*, and `code` within a text segment."""
    pos = 0
    for link in LINK_RE.finditer(text):
        if link.start() > pos:
            _add_inline_styles(paragraph, text[pos : link.start()], bold=bold, italic=italic)
        add_hyperlink(paragraph, link.group(1), link.group(2))
        pos = link.end()
    if pos < len(text):
        _add_inline_styles(paragraph, text[pos:], bold=bold, italic=italic)


def _add_inline_styles(paragraph, text: str, bold: bool = False, italic: bool = False) -> None:
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = INK
        else:
            r = paragraph.add_run(part)
            if bold:
                r.bold = True
            if italic:
                r.italic = True


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_table_separator(line: str) -> bool:
    s = line.strip().replace("|", "").replace(":", "").replace("-", "").strip()
    return not s


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.autofit = True

    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.rows[ri].cells[ci]
            text = row[ci] if ci < len(row) else ""
            text = text.replace("<br>", "\n").replace("<br/>", "\n")
            cell.text = ""
            p = cell.paragraphs[0]
            for j, segment in enumerate(text.split("\n")):
                if j:
                    p.add_run("\n")
                add_formatted_run(p, segment)
            if ri == 0:
                shade_cell(cell, TABLE_HEADER)
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)
            else:
                for run in p.runs:
                    run.font.size = Pt(10)
                if ri % 2 == 0:
                    shade_cell(cell, TABLE_ALT)


def setup_header_footer(document: Document, header_title: str) -> None:
    section = document.sections[0]
    section.different_first_page_header_footer = True
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    header = section.header
    header._element.clear()
    hp = header.add_paragraph()
    run_left = hp.add_run(header_title[:120])
    run_left.font.size = Pt(9)
    run_left.font.color.rgb = MUTED
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    section.first_page_header._element.clear()
    section.first_page_header.add_paragraph()

    footer = section.footer
    footer._element.clear()
    fp2 = footer.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp2.add_run(f"{ORG_SHORT}   |   Page ")
    r1.font.size = Pt(9)
    r1.font.color.rgb = MUTED
    add_page_number_field(fp2, align_center=False)
    r2 = fp2.add_run(f"   |   {datetime.now().strftime('%B %Y')}")
    r2.font.size = Pt(9)
    r2.font.color.rgb = MUTED

    first_footer = section.first_page_footer
    first_footer._element.clear()
    fp_first = first_footer.add_paragraph()
    fp_first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = fp_first.add_run(SITE_LABEL)
    rf.font.size = Pt(9)
    rf.font.color.rgb = GOLD


def extract_title_info(md_text: str, stem: str) -> tuple[str, str, str]:
    """Return (title, subtitle, header_short)."""
    title = stem.replace("-", " ").replace("_", " ")
    subtitle = ""
    for line in md_text.splitlines():
        s = line.strip()
        if s.startswith("# "):
            title = s[2:].strip()
            break
    lines = md_text.splitlines()
    for idx, line in enumerate(lines):
        if line.strip().startswith("# "):
            for nxt in lines[idx + 1 : idx + 6]:
                s = nxt.strip()
                if not s or s == "---":
                    continue
                if s.startswith("#"):
                    break
                subtitle = s
                break
            break
    header_short = title if len(title) <= 70 else title[:67] + "…"
    return title, subtitle, header_short


def add_title_page(document: Document, title: str, subtitle: str = "") -> None:
    for _ in range(4):
        document.add_paragraph()

    t = document.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(ORG_NAME_LINE1)
    r.font.size = Pt(14)
    r.font.color.rgb = MUTED
    r.font.small_caps = True

    t2 = document.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(ORG_NAME_LINE2)
    r2.font.size = Pt(14)
    r2.font.color.rgb = MUTED
    r2.font.small_caps = True

    document.add_paragraph()

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title_p.add_run(title)
    rt.font.name = FONT_HEADING
    rt.font.size = Pt(24 if len(title) < 60 else 20)
    rt.font.bold = True
    rt.font.color.rgb = NAVY

    rule = document.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rule.add_run("━" * 36)
    rr.font.color.rgb = GOLD

    if subtitle:
        sub = document.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rs = sub.add_run(subtitle)
        rs.font.size = Pt(12)
        rs.font.italic = True
        rs.font.color.rgb = MUTED

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = p.add_run("Generated: ")
    rl.font.size = Pt(10)
    rl.font.bold = True
    rl.font.color.rgb = NAVY
    rv = p.add_run(datetime.now().strftime("%d %B %Y"))
    rv.font.size = Pt(10)
    rv.font.color.rgb = INK

    document.add_page_break()


def add_toc_page(document: Document) -> None:
    document.add_heading("Table of contents", level=1)
    note = document.add_paragraph()
    add_formatted_run(
        note,
        "In Microsoft Word, click the table below and press **F9**, "
        "or right-click → **Update Field** → **Update entire table**.",
    )
    note.paragraph_format.space_after = Pt(12)
    toc_p = document.add_paragraph()
    add_table_of_contents(toc_p)
    document.add_page_break()


def add_code_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    shade_paragraph(p, CODE_FILL)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = INK


def add_blockquote(document: Document, text: str) -> None:
    try:
        p = document.add_paragraph(style="Quote")
    except KeyError:
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
    shade_paragraph(p, QUOTE_FILL)
    add_formatted_run(p, text)


def add_image(document: Document, image_path: Path, alt: str = "") -> None:
    if not image_path.is_file():
        p = document.add_paragraph()
        add_formatted_run(p, f"[Image not found: {alt or image_path.name}]")
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(str(image_path), width=Cm(14))
    except Exception:
        add_formatted_run(p, f"[Could not embed image: {image_path.name}]")
        return
    if alt:
        cap = document.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(alt)
        cr.font.size = Pt(9)
        cr.font.italic = True
        cr.font.color.rgb = MUTED


def convert_markdown_body(
    document: Document,
    text: str,
    md_path: Path,
    *,
    skip_first_h1: bool = True,
) -> None:
    lines = text.splitlines()
    i = 0
    table_buffer: list[list[str]] = []
    in_code = False
    code_lines: list[str] = []
    first_h1_skipped = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img_match:
            if table_buffer:
                add_table(document, table_buffer)
                table_buffer = []
            alt, src = img_match.group(1), img_match.group(2)
            img_path = (md_path.parent / src).resolve()
            add_image(document, img_path, alt)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            if table_buffer and is_table_separator(stripped):
                i += 1
                continue
            row = parse_table_row(stripped)
            if row:
                table_buffer.append(row)
            i += 1
            continue
        if table_buffer:
            add_table(document, table_buffer)
            table_buffer = []

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("#### "):
            document.add_heading(stripped[5:].strip(), level=4)
            i += 1
            continue

        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("# "):
            if skip_first_h1 and not first_h1_skipped:
                first_h1_skipped = True
                i += 1
                continue
            document.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith(">"):
            quote_text = stripped.lstrip(">").strip()
            i += 1
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_text += " " + lines[i].strip().lstrip(">").strip()
                i += 1
            add_blockquote(document, quote_text)
            continue

        bullet_match = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            try:
                p = document.add_paragraph(style=style)
            except KeyError:
                p = document.add_paragraph(style="List Bullet")
                if indent >= 2:
                    p.paragraph_format.left_indent = Cm(1)
            add_formatted_run(p, bullet_match.group(2).strip())
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = document.add_paragraph(style="List Number")
            add_formatted_run(p, re.sub(r"^\d+\.\s*", "", stripped))
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = document.add_paragraph()
            add_formatted_run(p, stripped.strip("*"))
            i += 1
            continue

        p = document.add_paragraph()
        add_formatted_run(p, stripped)
        i += 1

    if table_buffer:
        add_table(document, table_buffer)


def export_markdown_to_docx(md_path: Path, docx_path: Path) -> Path:
    """Convert one Markdown file to a formatted Word document."""
    md_text = md_path.read_text(encoding="utf-8")
    title, subtitle, header_short = extract_title_info(md_text, md_path.stem)

    document = Document()
    setup_styles(document)
    setup_header_footer(document, f"{ORG_SHORT} — {header_short}")
    add_title_page(document, title, subtitle)
    add_toc_page(document)
    convert_markdown_body(document, md_text, md_path, skip_first_h1=True)
    set_update_fields_on_open(document)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(docx_path))
    return docx_path
