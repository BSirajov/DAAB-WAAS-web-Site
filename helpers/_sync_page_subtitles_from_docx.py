"""Sync i18n/page-subtitles.json from documents/docx/Subtitles by page.docx."""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "documents" / "docx" / "Subtitles by page.docx"
OUT = ROOT / "i18n" / "page-subtitles.json"

# Table "Page" label -> routes.json page id
PAGE_IDS = {
    "home": "home",
    "foundation": "foundation",
    "mission & values": "mission",
    "mission and values": "mission",
    "activities": "activities",
    "presentations": "forum-2024-presentations",
    "official addresses": "forum-official",
    "programme": "forum-program",
    "program": "forum-program",
    "impressions": "forum-impressions",
    "photos gallery": "forum-photos-gallery",
    "scientists directory": "scientists-list",
    "scientists profiles": "scientists-profiles",
    "executive board": "executive-board",
    "charter": "charter",
    "why become a member": "membership-value",
    "membership terms": "membership",
    "application": "membership-application",
    "roadmap": "forum-roadmap",
    "related stories": "forum-bagli-hekayeler",
    "cooperation": "forum-cooperation",
}


def norm_label(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip().lower())


def page_id_from_label(label: str) -> str | None:
    key = norm_label(label)
    if key in PAGE_IDS:
        return PAGE_IDS[key]
    for pattern, pid in PAGE_IDS.items():
        if key == pattern or key.startswith(pattern):
            return pid
    return None


def parse_docx(path: Path) -> dict[str, dict[str, str]]:
    doc = Document(path)
    if not doc.tables:
        raise SystemExit(f"No table in {path}")
    table = doc.tables[0]
    pages: dict[str, dict[str, str]] = {}
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 3:
            continue
        label, en, az = cells[0], cells[1], cells[2]
        pid = page_id_from_label(label)
        if not pid:
            raise SystemExit(f"Unknown page label in docx: {label!r}")
        pages[pid] = {"en": en, "az": az}
    return pages


def main() -> None:
    if not DOCX.is_file():
        raise SystemExit(f"Missing {DOCX}")
    pages = parse_docx(DOCX)
    payload = {
        "version": 2,
        "source": "documents/docx/Subtitles by page.docx",
        "pages": pages,
    }
    OUT.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"Wrote {len(pages)} page subtitles to {OUT.relative_to(ROOT)}")
    import importlib.util

    spec = importlib.util.spec_from_file_location(
        "embed", ROOT / "helpers" / "_embed_page_subtitles.py"
    )
    mod = importlib.util.module_from_spec(spec)
    assert spec.loader
    spec.loader.exec_module(mod)
    mod.main()


if __name__ == "__main__":
    main()
