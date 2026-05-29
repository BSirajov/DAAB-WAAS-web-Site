#!/usr/bin/env python3
"""Export scientist profile cards to publication-ready Word catalogues (AZ + EN).

Source: i18n/scientists-profiles.json, images/scientists-photos/, images/qr/{az,en}/.
Layout: floating portrait (square wrap) + top-right QR; profile text flows continuously.
Includes title page, table of contents, headers, footers, and page numbers.
"""
from __future__ import annotations

import html as html_lib
import sys
from datetime import datetime
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
if str(ROOT / "helpers") not in sys.path:
    sys.path.insert(0, str(ROOT / "helpers"))

from daab_docx_export import (  # noqa: E402
    FONT_BODY,
    FONT_HEADING,
    GOLD,
    INK,
    MUTED,
    NAVY,
    add_floating_picture,
    add_hyperlink,
    add_page_number_field,
    add_paragraph_box_border,
    add_table_of_contents,
    set_update_fields_on_open,
    setup_styles,
    shade_paragraph,
)
from i18n_person_names_en import az_upper_name_latin  # noqa: E402
from scientists_profiles_core import (  # noqa: E402
    CRED_LABEL,
    META_LABELS,
    az_upper_name,
    load_profiles,
    profile_bio,
    profile_country,
    profile_deep_link,
    profile_field,
    profile_name,
    profile_title,
    slug_from_photo,
)

OUT_AZ = ROOT / "documents" / "docx" / "Scientists-Profiles-Catalogue-AZ.docx"
OUT_EN = ROOT / "documents" / "docx" / "Scientists-Profiles-Catalogue-EN.docx"

PHOTOS_DIR = ROOT / "images" / "scientists-photos"
QR_DIR = ROOT / "images" / "qr"

AVATAR_W = Cm(3.92)
AVATAR_H = Cm(4.65)
QR_SIZE = Cm(2.11)
FLOAT_DIST = Emu(114300)
META_FILL = "F3F9FD"
HEADING_BLUE = RGBColor(0x00, 0x69, 0xB4)
TITLE_INK = RGBColor(0x10, 0x20, 0x33)
LEAD_INK = RGBColor(0x1A, 0x3D, 0x5C)

DOC_META = {
    "az": {
        "org_line": "DÜNYA AZƏRBAYCANLI ALIMLƏR BİRLİYİ",
        "org_short": "DAAB / WAAS",
        "title": "Alimlər profilləri kataloqu",
        "subtitle": "Sayt kartları ilə eyni struktur — foto, QR kod, ad, vəzifə, ixtisas, bioqrafiya",
        "header": "DAAB — Alimlər profilləri (AZ)",
        "toc_title": "Mündəricat",
        "toc_note": (
            "Microsoft Word-də cədvələ klik edin və F9 düyməsini basın, "
            "və ya sağ klik → Sahəni yenilə → Bütün cədvəli yenilə."
        ),
        "section_title": "Alim profilləri",
        "intro": (
            "Bu sənəd saytın «Alimlər profilləri» səhifəsindəki kartların tam surətidir. "
            "Portret solda yerləşir və profil mətni foto ətrafında davamlı axınla sarılır; "
            "QR kod yuxarı sağ küncdədir."
        ),
        "page_label": "Səhifə",
        "profiles_label": "profil",
        "prepared": "Hazırlanma tarixi",
        "qr_caption": "Profil linki",
        "listen_label": "Dinlə — qısa məlumat",
        "site": "daab-waas.com",
    },
    "en": {
        "org_line": "WORLD ASSOCIATION OF AZERBAIJANI SCIENTISTS",
        "org_short": "WAAS / DAAB",
        "title": "Scientists profiles catalogue",
        "subtitle": "Same structure as site cards — photo, QR code, name, title, field, biography",
        "header": "WAAS — Scientists profiles (EN)",
        "toc_title": "Table of contents",
        "toc_note": (
            "Click the table in Microsoft Word and press F9, "
            "or right-click → Update Field → Update entire table."
        ),
        "section_title": "Scientist profiles",
        "intro": (
            "This document reproduces the Scientists profiles page cards in full. "
            "The portrait sits on the left with profile text flowing continuously around it; "
            "the QR code appears in the top-right corner."
        ),
        "page_label": "Page",
        "profiles_label": "profiles",
        "prepared": "Prepared",
        "qr_caption": "Profile link",
        "listen_label": "Listen — summary line",
        "site": "daab-waas.com",
    },
}


def _run_style(run, *, size=11, bold=False, italic=False, color: RGBColor | None = None, name=FONT_BODY) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _para_spacing(paragraph, *, before=0, after=4, line=1.35) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    pf.keep_with_next = False
    pf.page_break_before = False
    pf.widow_control = True


def _name_heading(profile: dict, lang: str) -> str:
    name_display = profile_name(profile, lang)
    if lang == "az" and profile.get("name_heading_az"):
        return profile["name_heading_az"]
    if lang == "en" and profile.get("name_heading_en"):
        return profile["name_heading_en"]
    if lang == "en":
        return az_upper_name_latin(name_display)
    return az_upper_name(name_display)


def _element_text(el: Tag) -> str:
    return html_lib.unescape(el.get_text(" ", strip=True))


def _listen_lead(profile: dict, lang: str) -> str:
    key = "listen_lead_az" if lang == "az" else "listen_lead_en"
    return (profile.get(key) or "").strip()


def _add_text_paragraph(document: Document, text: str, *, size=10, bold=False, color=INK, after=4, justify=True) -> None:
    p = document.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    _run_style(run, size=size, bold=bold, color=color)
    _para_spacing(p, after=after)


def _add_meta_block(document: Document, profile: dict, lang: str) -> None:
    labels = META_LABELS[lang]
    field = profile_field(profile, lang)
    email = (profile.get("email") or "").strip()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    shade_paragraph(p, META_FILL)
    add_paragraph_box_border(p)

    r1 = p.add_run(f"{labels['field']} ")
    _run_style(r1, size=10, bold=True, color=MUTED)
    r1b = p.add_run(field)
    _run_style(r1b, size=10, bold=True, color=HEADING_BLUE)
    p.add_run("\n")
    r2 = p.add_run(f"{labels['email']} ")
    _run_style(r2, size=10, bold=True, color=MUTED)
    if email:
        add_hyperlink(p, email, f"mailto:{email}")
    else:
        r2b = p.add_run("—")
        _run_style(r2b, size=10, color=MUTED)
    _para_spacing(p, before=2, after=6, line=1.35)


def _add_bio_divider(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D4E6F2")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_bio_blocks(document: Document, bio_html: str) -> None:
    if not bio_html.strip():
        _add_text_paragraph(document, "—", color=MUTED)
        return

    soup = BeautifulSoup(bio_html, "html.parser")
    count = 0
    _add_bio_divider(document)

    def add_para(text: str, *, lead: bool = False, section: bool = False, bullet: bool = False) -> None:
        nonlocal count
        text = text.strip()
        if not text:
            return
        p = document.add_paragraph(style="List Bullet" if bullet else None)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        if section:
            _run_style(run, size=10, bold=True, color=HEADING_BLUE, name=FONT_HEADING)
            _para_spacing(p, before=6 if count else 2, after=3, line=1.3)
        elif lead:
            _run_style(run, size=10, color=LEAD_INK)
            _para_spacing(p, before=0, after=4, line=1.45)
        else:
            _run_style(run, size=10, color=MUTED)
            _para_spacing(p, before=0, after=4, line=1.5)
        count += 1

    for node in soup.children:
        if isinstance(node, NavigableString) or not isinstance(node, Tag):
            continue
        classes = node.get("class") or []
        if node.name == "p":
            cls_set = set(classes)
            if "bio-section-title" in cls_set:
                add_para(_element_text(node), section=True)
            elif "bio-lead" in cls_set:
                add_para(_element_text(node), lead=True)
            else:
                add_para(_element_text(node))
        elif node.name == "ul" and "bullets" in classes:
            for li in node.find_all("li", recursive=False):
                add_para(_element_text(li), bullet=True)
        elif node.name == "div" and "awards-block" in classes:
            for ul in node.find_all("ul", class_="awards-list"):
                for li in ul.find_all("li", recursive=False):
                    add_para(_element_text(li), bullet=True)

    if count == 0:
        add_para(_element_text(soup))


def _add_profile_card(document: Document, profile: dict, lang: str, toc_label: str) -> None:
    meta = DOC_META[lang]
    photo_name = (profile.get("photo") or "").strip()
    slug = slug_from_photo(photo_name)
    photo_path = PHOTOS_DIR / photo_name
    qr_path = QR_DIR / lang / f"{slug}.png"
    bio_html = profile_bio(profile, lang)
    profile_url = profile_deep_link(lang, slug)

    heading = document.add_heading(toc_label, level=2)
    _para_spacing(heading, before=10, after=4, line=1.2)
    heading.paragraph_format.keep_with_next = False

    if qr_path.is_file():
        qr_p = document.add_paragraph()
        qr_p.paragraph_format.space_after = Pt(0)
        add_floating_picture(qr_p, qr_path, QR_SIZE, QR_SIZE, side="right", dist_emu=FLOAT_DIST)
        cap = document.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _para_spacing(cap, before=0, after=2)
        add_hyperlink(cap, meta["qr_caption"], profile_url)

    if photo_path.is_file():
        photo_p = document.add_paragraph()
        photo_p.paragraph_format.space_after = Pt(2)
        add_floating_picture(photo_p, photo_path, AVATAR_W, AVATAR_H, side="left", dist_emu=FLOAT_DIST)
    else:
        _add_text_paragraph(document, "[photo missing]", size=9, color=MUTED, after=2)

    country = profile_country(profile, lang)
    title = profile_title(profile, lang)
    listen_lead = _listen_lead(profile, lang)

    _add_text_paragraph(document, country, size=9, bold=True, color=MUTED, after=3)

    if title:
        _add_text_paragraph(document, title, size=10, bold=True, color=TITLE_INK, after=5)

    _add_meta_block(document, profile, lang)

    if listen_lead:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rl = p.add_run(f"{meta['listen_label']}: ")
        _run_style(rl, size=9, bold=True, color=MUTED)
        rt = p.add_run(listen_lead)
        _run_style(rt, size=10, bold=True, color=LEAD_INK)
        _para_spacing(p, before=2, after=4, line=1.45)

    _add_bio_blocks(document, bio_html)

    end = document.add_paragraph()
    end.paragraph_format.space_before = Pt(8)
    end.paragraph_format.space_after = Pt(12)
    p_pr = end._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B8D4E8")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _setup_header_footer(document: Document, lang: str) -> None:
    meta = DOC_META[lang]
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.different_first_page_header_footer = True

    header = section.header
    header._element.clear()
    hp = header.add_paragraph()
    run = hp.add_run(meta["header"])
    _run_style(run, size=9, color=MUTED)

    section.first_page_header._element.clear()

    footer = section.footer
    footer._element.clear()
    fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run(f"{meta['org_short']}   |   {meta['page_label']} ")
    _run_style(r1, size=9, color=MUTED)
    add_page_number_field(fp, align_center=False)
    month_year = datetime.now().strftime("%B %Y")
    r2 = fp.add_run(f"   |   {month_year}   |   {meta['site']}")
    _run_style(r2, size=9, color=MUTED)

    first_footer = section.first_page_footer
    first_footer._element.clear()
    fp_first = first_footer.add_paragraph()
    fp_first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = fp_first.add_run(meta["site"])
    _run_style(rf, size=9, color=GOLD)


def _add_title_page(document: Document, lang: str, profile_count: int) -> None:
    meta = DOC_META[lang]
    for _ in range(2):
        document.add_paragraph()

    org = document.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ro = org.add_run(meta["org_line"])
    _run_style(ro, size=13, color=MUTED)
    ro.font.small_caps = True

    document.add_paragraph()

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title_p.add_run(meta["title"])
    _run_style(rt, size=26, bold=True, color=NAVY, name=FONT_HEADING)

    rule = document.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rule.add_run("━" * 32)
    _run_style(rr, size=11, color=GOLD)

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(meta["subtitle"])
    _run_style(rs, size=12, italic=True, color=HEADING_BLUE)

    count_p = document.add_paragraph()
    count_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    count_p.paragraph_format.space_before = Pt(18)
    rc = count_p.add_run(f"{profile_count} {meta['profiles_label']}")
    _run_style(rc, size=11, bold=True, color=INK)

    date_p = document.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_before = Pt(20)
    rd = date_p.add_run(f"{meta['prepared']}: {datetime.now().strftime('%d %B %Y')}")
    _run_style(rd, size=10, color=MUTED)

    date_p.paragraph_format.page_break_after = True


def _add_toc_page(document: Document, lang: str) -> None:
    meta = DOC_META[lang]
    h = document.add_heading(meta["toc_title"], level=1)
    _para_spacing(h, after=8)

    note = document.add_paragraph()
    rn = note.add_run(meta["toc_note"])
    _run_style(rn, size=10, italic=True, color=MUTED)
    _para_spacing(note, after=8)

    toc_p = document.add_paragraph()
    add_table_of_contents(toc_p)
    toc_p.paragraph_format.page_break_after = True


def _configure_export_styles(document: Document) -> None:
    setup_styles(document)
    for level in ("Heading 1", "Heading 2"):
        h = document.styles[level]
        h.paragraph_format.keep_with_next = False
        h.paragraph_format.page_break_before = False


def _add_profiles_section(document: Document, profiles: list[dict], lang: str) -> None:
    meta = DOC_META[lang]
    h = document.add_heading(meta["section_title"], level=1)
    _para_spacing(h, after=8)

    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = intro.add_run(meta["intro"])
    _run_style(run, size=11, color=INK)
    _para_spacing(intro, after=12)

    for profile in profiles:
        say = profile.get("say") or 0
        name_heading = _name_heading(profile, lang)
        degree = (profile.get("degree") or "").strip()
        cred = CRED_LABEL.get(degree, degree)
        toc_label = f"{say}. {name_heading}"
        if cred:
            toc_label = f"{say}. {name_heading} ({cred})"
        _add_profile_card(document, profile, lang, toc_label)


def export_catalogue(lang: str, out_path: Path) -> Path:
    profiles = load_profiles()
    document = Document()
    _configure_export_styles(document)
    _setup_header_footer(document, lang)
    _add_title_page(document, lang, len(profiles))
    _add_toc_page(document, lang)
    _add_profiles_section(document, profiles, lang)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    set_update_fields_on_open(document)
    tmp_path = out_path.with_name(f"_{out_path.name}")
    document.save(str(tmp_path))
    try:
        tmp_path.replace(out_path)
    except PermissionError:
        print(f"Warning: {out_path.name} is open — saved as {tmp_path.name}")
        return tmp_path
    if tmp_path.is_file():
        tmp_path.unlink(missing_ok=True)
    return out_path


def main() -> None:
    az_path = export_catalogue("az", OUT_AZ)
    en_path = export_catalogue("en", OUT_EN)
    n = len(load_profiles())
    print(f"Exported {n} AZ profile cards -> {az_path.relative_to(ROOT)}")
    print(f"Exported {n} EN profile cards -> {en_path.relative_to(ROOT)}")
    print("Open in Word and press F9 to refresh the table of contents.")


if __name__ == "__main__":
    main()
