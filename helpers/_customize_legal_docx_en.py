#!/usr/bin/env python3
"""Customise documents/Legal/*.docx for WAAS (English) using live site facts."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from _paths import ROOT

LEGAL = ROOT / "documents" / "Legal"
UPDATED = "27 July 2026"
YEAR = "2026"

ORG_TR = "Dünya Azerbaycanlı Alimler Derneği"
ORG_EN = "World Association of Azerbaijani Scientists"
BRAND = "DAAB / WAAS"
SITE = "https://daab-waas.com"
EMAIL = "info@daab-waas.com"
PHONE = "+90 555 147 46 74"
ADDRESS_LINES = (
    "Feneryolu Mahallesi",
    "Gazi Muhtar Paşa Sokak No:44",
    "Kadıköy, Istanbul, Türkiye",
)
ADDRESS = ", ".join(ADDRESS_LINES)
CHAIR = "Prof. Dr. Messoud Efendiyev, Chair of the WAAS Executive Board"


def set_run_font(run, *, bold=False, size=11):
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, bold=True, size=18)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, bold=False, size=11)


def add_h(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, bold=True, size=12)


def add_p(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, size=11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_run_font(r, bold=True, size=10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(val)
            set_run_font(r, size=10)


def contact_block(doc: Document) -> None:
    add_p(doc, f"{ORG_TR}")
    add_p(doc, f"{ORG_EN} ({BRAND})")
    add_p(doc, ADDRESS)
    add_p(doc, f"Email: {EMAIL}")
    add_p(doc, f"Phone: {PHONE}")
    add_p(doc, f"Website: {SITE}")
    add_p(doc, f"Leadership: {CHAIR}")


def write_privacy() -> Path:
    doc = Document()
    add_title(doc, "Privacy Notice (KVKK)")
    add_subtitle(doc, f"{ORG_EN} ({BRAND})")
    add_subtitle(doc, f"Last updated: {UPDATED}")
    add_p(
        doc,
        "This English document is aligned with the live Privacy notice on "
        f"{SITE}/en/privacy.html. It is an information notice under Türkiye’s "
        "Personal Data Protection Law No. 6698 (KVKK), Article 10. It is not legal "
        "advice and should be reviewed by counsel when needed.",
    )

    add_h(doc, "1. Data controller")
    add_p(
        doc,
        f"The data controller is {ORG_TR}, a Turkish association (dernek), "
        f"operating publicly as {ORG_EN} ({BRAND}).",
    )
    contact_block(doc)

    add_h(doc, "2. Personal data we process")
    add_p(doc, "Depending on your interaction with the Site, we may process:")
    add_bullets(
        doc,
        [
            "Identity: name, surname, academic degree or title",
            "Contact: email, phone, country and city",
            "Professional: university, workplace, scientific field, contribution notes",
            "Membership application / CV: online form fields; CV and photo sent separately by email",
            "Scientist / board profile content: biography, photograph, publications and related materials published with permission",
            "Technical: IP address, browser type, page visits (if analytics consent is given)",
            "Cookies and local storage: essential cookies/storage and, with consent, analytics — see the Cookie Policy",
        ],
    )

    add_h(doc, "3. Purposes of processing")
    add_bullets(
        doc,
        [
            "Operate and secure the website",
            "Receive, assess and respond to membership applications",
            "Respond to donation and sponsorship enquiries",
            "Scientific activity, forums and organisational communication",
            "Publish member/scientist profiles and contributed works with permission",
            "Statistics and site improvement (only with analytics consent)",
            "Comply with legal obligations and protect legal rights",
        ],
    )

    add_h(doc, "4. Legal bases (KVKK Art. 5)")
    add_p(doc, "Processing relies on one or more of the following, as applicable:")
    add_bullets(
        doc,
        [
            "Your explicit consent (e.g. analytics cookies; other cases where consent is required)",
            "Necessity for establishing or performing a contract (membership application process)",
            "Necessity for the controller’s legal obligation",
            "Legitimate interests of the controller (site security and continuity of operations), provided this does not harm your fundamental rights and freedoms",
        ],
    )
    add_table(
        doc,
        ["Purpose", "Typical legal basis"],
        [
            [
                "Publishing member/scientist profiles, photos and biographies",
                "Consent / permission given for publication",
            ],
            [
                "Publishing submitted books, articles or photo collections",
                "Consent / permission of the author or rights holder",
            ],
            [
                "Membership application processing",
                "Contractual necessity / legitimate interest in assessing applications",
            ],
            [
                "Responding to contact, donation or sponsorship enquiries",
                "Legitimate interest (responding to a request you initiated)",
            ],
            [
                "Operating and securing the Site",
                "Legitimate interest (security and functionality)",
            ],
            [
                "Google Analytics 4",
                "Explicit consent via the cookie banner",
            ],
        ],
    )

    add_h(doc, "5. Recipients and transfers")
    add_p(doc, "We do not sell personal data. Data may be shared, limited to purpose, with:")
    add_bullets(
        doc,
        [
            "WAAS membership / board commission and authorised volunteer team members",
            "Hosting, email and IT service providers acting as processors",
            "Google LLC — Google Analytics 4 (only if you consent; may involve transfers abroad)",
            "Competent public authorities where required by law",
            "Public visitors to the Site, for information you have agreed to publish (e.g. profile, photo, submitted works)",
        ],
    )
    add_p(
        doc,
        "Where data is transferred abroad, measures aligned with KVKK and related rules are intended.",
    )

    add_h(doc, "6. Retention")
    add_p(
        doc,
        "Data is kept only as long as necessary for the purpose and applicable legal "
        "retention periods. Membership applications are retained for a reasonable "
        "assessment and follow-up period. Profile and published works remain while "
        "publication is authorised or until removal is requested. Analytics data is "
        "deleted or anonymised according to consent and tool settings.",
    )

    add_h(doc, "7. Your rights (KVKK Art. 11)")
    add_p(doc, "As a data subject you may, among other rights:")
    add_bullets(
        doc,
        [
            "Learn whether your data is processed and request information",
            "Learn the purpose and whether use is consistent with that purpose",
            "Request correction of incomplete or inaccurate data",
            "Request erasure/destruction under the conditions in KVKK",
            "Request that corrections/erasures be notified to recipients",
            "Object to results of exclusively automated analysis",
            "Claim compensation for damage arising from unlawful processing",
            "Withdraw consent for future publication or analytics (without affecting prior lawful processing)",
        ],
    )
    add_p(
        doc,
        f"Contact: {EMAIL} or the postal address above. Please include enough "
        "information to verify your identity. Response times follow KVKK and "
        "applicable procedures.",
    )

    add_h(doc, "8. Membership form")
    add_p(
        doc,
        "Data submitted via the online membership form "
        f"({SITE}/en/application.html) is used only for the membership process. "
        "Before sending, applicants must confirm that they have read and agree to "
        "this notice. CV and photo are sent separately by email.",
    )

    add_h(doc, "9. Children")
    add_p(
        doc,
        "The Site and membership services are generally intended for an adult "
        "academic audience. We do not knowingly collect personal data from persons "
        "under 18.",
    )

    add_h(doc, "10. Security")
    add_p(
        doc,
        "We take reasonable technical and organisational measures to protect "
        "personal data against unauthorised access, loss or misuse. No method of "
        "transmission or storage is 100% secure.",
    )

    add_h(doc, "11. Changes")
    add_p(
        doc,
        "This notice may be updated. Material changes will be published on the Site; "
        "the “Last updated” date will change. The authoritative public version is "
        f"{SITE}/en/privacy.html.",
    )

    add_h(doc, "12. Contact")
    contact_block(doc)

    path = LEGAL / "01-Privacy-Policy.docx"
    doc.save(path)
    return path


def write_terms() -> Path:
    doc = Document()
    add_title(doc, "Terms of Use")
    add_subtitle(doc, f"{ORG_EN} ({BRAND})")
    add_subtitle(doc, f"Last updated: {UPDATED}")
    add_p(
        doc,
        "This English document is aligned with the live Terms of use on "
        f"{SITE}/en/terms.html.",
    )

    add_h(doc, "1. Parties and acceptance")
    add_p(
        doc,
        f"By accessing or using {SITE} (the “Site”), you agree to these Terms of Use. "
        "If you do not agree, please do not use the Site.",
    )
    add_p(
        doc,
        f"The Site is operated by {ORG_TR} ({ORG_EN}, {BRAND}), a Turkish association "
        f"(dernek) based in Kadıköy, Istanbul, Türkiye.",
    )

    add_h(doc, "2. About the Site")
    add_p(
        doc,
        "The Site connects, showcases and supports Azerbaijani scientists and scholars. "
        "It includes, among other things, association information, scientist directory "
        "and profiles, forums and activities, membership information, and sponsorship "
        "or donation pages, in Azerbaijani and English.",
    )

    add_h(doc, "3. Intellectual property")
    add_p(
        doc,
        f"Site design, text and original content created by {BRAND} (layout, graphics, "
        f"original articles, the DAAB/WAAS logo) are owned by {BRAND} and protected by "
        "copyright. You may not copy, reproduce or redistribute this content without "
        "prior written permission.",
    )
    add_p(
        doc,
        "Member-submitted content (biographies, photographs, books, articles and other "
        "works) remains the intellectual property of the respective author or rights "
        "holder. WAAS publishes such content only with the contributor’s permission, "
        "as described in the Copyright & Attribution Notice.",
    )
    add_p(
        doc,
        "Third-party trademarks or logos (e.g. partner institutions, universities) "
        "remain the property of their respective owners.",
    )

    add_h(doc, "4. Acceptable use")
    add_p(doc, "When using the Site, you agree not to:")
    add_bullets(
        doc,
        [
            "Scrape, systematically copy or republish the member directory or profile data for unauthorised commercial use",
            "Use any content from the Site to harass, defame or misrepresent any individual featured on it",
            "Attempt to gain unauthorised access to any part of the Site or its underlying systems",
            "Upload or submit content that is unlawful, infringing or misleading",
            "Misuse membership, sponsorship or contact forms",
        ],
    )

    add_h(doc, "5. Member and contributor submissions")
    add_p(doc, "If you submit a biography, photograph, article, book or other work for publication:")
    add_bullets(
        doc,
        [
            "You confirm that you own the rights, or have obtained the necessary permission, for WAAS to publish the material",
            "You grant WAAS a non-exclusive, royalty-free licence to display, reproduce and translate (into English/Azerbaijani as applicable) the material on the Site and in related WAAS materials until you request removal",
            f"You may request removal or correction at any time by contacting {EMAIL}",
        ],
    )

    add_h(doc, "6. Disclaimers")
    add_p(
        doc,
        "The Site is provided “as is.” While we make reasonable efforts to keep "
        "information accurate and current, WAAS does not guarantee completeness or "
        "accuracy of all content, much of which is supplied by members. Content is for "
        "scientific, organisational and public-information purposes and is not legal "
        "advice or an official government statement.",
    )
    add_p(
        doc,
        "Links to external sites are provided for convenience. WAAS is not responsible "
        "for the content or practices of external sites.",
    )

    add_h(doc, "7. Limitation of liability")
    add_p(
        doc,
        "To the fullest extent permitted by applicable law, WAAS shall not be liable "
        "for any indirect, incidental or consequential damages arising from your use "
        "of the Site, including reliance on any information published on it.",
    )

    add_h(doc, "8. Governing law")
    add_p(
        doc,
        "These Terms are governed by the laws of the Republic of Türkiye. Any disputes "
        "shall be subject to the jurisdiction of the courts of Istanbul, Türkiye, "
        "unless otherwise required by applicable mandatory law in your jurisdiction.",
    )

    add_h(doc, "9. Changes")
    add_p(
        doc,
        "We may revise these Terms from time to time. Continued use of the Site after "
        "changes are posted constitutes acceptance of the revised Terms. The "
        f"authoritative public version is {SITE}/en/terms.html.",
    )

    add_h(doc, "10. Contact")
    contact_block(doc)

    path = LEGAL / "02-Terms-of-Use.docx"
    doc.save(path)
    return path


def write_cookies() -> Path:
    doc = Document()
    add_title(doc, "Cookie Policy")
    add_subtitle(doc, f"{ORG_EN} ({BRAND})")
    add_subtitle(doc, f"Last updated: {UPDATED}")
    add_p(
        doc,
        "This English document is aligned with the live Cookie policy on "
        f"{SITE}/en/cookies.html and the Site’s cookie consent implementation.",
    )

    add_h(doc, "1. What are cookies?")
    add_p(
        doc,
        "Cookies are small text files stored on your device by your browser. They may "
        "support site functions, language preference and (with consent) visit statistics. "
        "The Site also uses browser local storage for certain preferences.",
    )

    add_h(doc, "2. Cookies and storage we use")
    add_table(
        doc,
        ["Category", "Purpose", "Examples / keys", "Consent required?"],
        [
            [
                "Essential",
                "Core site functions, security, remembering language and cookie choice",
                "daab-cookie-consent (localStorage); daab-lang (language preference)",
                "No — always active",
            ],
            [
                "Analytics (optional)",
                "Google Analytics 4 — page views and aggregate statistics",
                "GA4 cookies / tags (only after consent); measurement configured in site analytics settings",
                "Yes — via cookie banner",
            ],
        ],
    )
    add_p(
        doc,
        "The Site does not currently rely on separate “functional” or social-share "
        "cookie categories beyond the essential and optional analytics categories above. "
        "If additional third-party embeds that set cookies are introduced later, this "
        "policy and the banner categories should be updated.",
    )

    add_h(doc, "3. Google Analytics")
    add_p(
        doc,
        "If you consent, Google Analytics 4 may be enabled. This is a service provided "
        "by Google LLC and may involve transfers of data abroad. If you do not consent "
        "(or choose “Essential only”), the analytics script is not loaded.",
    )

    add_h(doc, "4. Local storage")
    add_p(
        doc,
        "Your cookie consent choice is stored in the browser’s localStorage under the "
        "key daab-cookie-consent so the banner is not shown on every visit. Language "
        "preference may be stored under daab-lang.",
    )

    add_h(doc, "5. How we obtain and manage consent")
    add_p(
        doc,
        "On your first visit, a cookie banner offers: Accept all, Essential only, "
        "Preferences, and Save choices. Essential cookies do not require consent. "
        "Analytics cookies run only with your consent. You can reopen cookie settings "
        "from the Cookie policy page (“Open cookie settings”).",
    )
    add_p(
        doc,
        "Suggested banner text (already in use on the Site): “We use essential cookies "
        "to run this site. Optional analytics cookies (Google Analytics) help us "
        "understand visits and are used only with your consent, in line with Turkish "
        "KVKK requirements.”",
    )

    add_h(doc, "6. Browser controls")
    add_p(
        doc,
        "In addition to the Site’s consent tool, you can control or delete cookies "
        "through your browser settings. Blocking all cookies or clearing local storage "
        "may affect language preference and may cause the consent banner to reappear.",
    )

    add_h(doc, "7. Related documents")
    add_bullets(
        doc,
        [
            f"Privacy notice: {SITE}/en/privacy.html",
            f"Cookie policy (web): {SITE}/en/cookies.html",
            f"Legal notice: {SITE}/en/legal-notice.html",
            f"Terms of use: {SITE}/en/terms.html",
        ],
    )

    add_h(doc, "8. Changes")
    add_p(
        doc,
        "We may update this Cookie Policy as the Site’s use of cookies changes. The "
        "“Last updated” date reflects the most recent revision.",
    )

    add_h(doc, "9. Contact")
    add_p(doc, f"Questions: {EMAIL}")
    contact_block(doc)

    path = LEGAL / "03-Cookie-Policy.docx"
    doc.save(path)
    return path


def write_copyright() -> Path:
    doc = Document()
    add_title(doc, "Copyright & Attribution Notice")
    add_subtitle(doc, f"{ORG_EN} ({BRAND})")
    add_subtitle(doc, f"Last updated: {UPDATED}")

    add_h(doc, "1. Site content")
    add_p(
        doc,
        f"Unless otherwise noted, the design, layout, graphics and original written "
        f"content of this Site are © {YEAR} {BRAND}. All Rights Reserved. This includes "
        "the DAAB/WAAS name and logo.",
    )

    add_h(doc, "2. Content published with permission")
    add_p(
        doc,
        "This Site features biographies, photographs, books, articles and other works "
        "contributed by or about Azerbaijani scientists, scholars and public figures "
        "(“Featured Individuals” and “Contributors”). Such content is published with "
        "the knowledge and permission of the Featured Individual or Contributor (or, "
        "for published books, with the author’s or rights holder’s consent to feature "
        "cover images, excerpts or full works).",
    )
    add_p(
        doc,
        "Copyright in each contributed work remains with its original author or rights "
        "holder. Publication on this Site does not transfer ownership of that copyright "
        f"to {BRAND}.",
    )
    add_p(
        doc,
        "Examples of contributed materials on the Site include scientist profile pages, "
        "board member photographs, forum materials, and book/media collections (e.g. "
        "works associated with individual scholars such as published books and photo "
        "collections). Written or emailed consent should be kept on file for each "
        "collection published in full or in substantial part.",
    )

    add_h(doc, "3. Third-party and licensed material")
    add_p(
        doc,
        "Where the Site uses icons, map tiles, fonts or other licensed assets not "
        "created by WAAS or its members, such material is used in accordance with its "
        "applicable licence, and attribution is provided where the licence requires it.",
    )

    add_h(doc, "4. Permitted use by visitors")
    add_p(
        doc,
        "You may view and share links to pages on this Site for personal, non-commercial "
        "and academic purposes. You may not:",
    )
    add_bullets(
        doc,
        [
            "Reproduce, republish or redistribute site content (including member photos, biographies or full-text books) without permission",
            "Use any Featured Individual’s name, photo or biography for commercial purposes",
            "Claim authorship of any work published on this Site",
        ],
    )

    add_h(doc, "5. Reporting a concern")
    add_p(
        doc,
        f"If you are a rights holder and believe your work has been published without "
        f"proper authorisation, or you would like content removed or corrected, contact "
        f"{EMAIL} with:",
    )
    add_bullets(
        doc,
        [
            "A description of the material and its location (URL) on the Site",
            "A description of your rights in that material",
            "Your contact information",
        ],
    )
    add_p(
        doc,
        "We will review and respond to legitimate requests promptly, and will remove "
        "or correct content where appropriate.",
    )

    add_h(doc, "6. Footer notice")
    add_p(
        doc,
        "The short-form notice used in the Site footer (English pages) is:",
    )
    add_p(doc, f"© {YEAR} WAAS — All Rights Reserved")
    add_p(
        doc,
        "Related public pages: Legal notice, Privacy, Cookies and Terms of use "
        f"(see {SITE}/en/).",
    )

    add_h(doc, "7. Contact")
    contact_block(doc)

    path = LEGAL / "04-Copyright-Attribution-Notice.docx"
    doc.save(path)
    return path


def main() -> None:
    LEGAL.mkdir(parents=True, exist_ok=True)
    written = [
        write_privacy(),
        write_terms(),
        write_cookies(),
        write_copyright(),
    ]
    for p in written:
        print(f"Wrote {p.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
