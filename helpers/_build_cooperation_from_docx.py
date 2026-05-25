#!/usr/bin/env python3
"""Build az/en forum/2024/cooperation.html from forum_2024/Əməkdaşlıq.docx."""
from __future__ import annotations

import html
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from _embed_static_nav import forum_nav_strip  # noqa: E402
from forum_en_common import FORUM_FOOTER_EN  # noqa: E402
from forum_en_cooperation import COOPERATION_EN  # noqa: E402

DOCX = ROOT / "forum_2024" / "Əməkdaşlıq.docx"
OUT_AZ = ROOT / "az" / "forum" / "2024" / "cooperation.html"
OUT_EN = ROOT / "en" / "forum" / "2024" / "cooperation.html"
ASSET = "../../../"
PAGE_ID = "forum-cooperation"
SIDEBAR_SCRIPT = f'<script src="{ASSET}js/daab-sidebar-timeline.js?v=1" defer></script>'

SECTIONS = (
    {"id": "contributions", "az_title": "Töhfələr və tərəfdaşlar"},
    {"id": "university-rectors", "az_title": "Universitet rektorları"},
)


def esc(s: str) -> str:
    return html.escape(s, quote=True)


def norm(text: str) -> str:
    return text.replace("\r", "").replace("\xa0", " ").strip()


def parse_docx(doc: Document) -> dict:
    lines = [norm(p.text) for p in doc.paragraphs if norm(p.text)]
    title = lines[0] if lines else ""
    items = lines[1:] if len(lines) > 1 else []
    table_rows: list[tuple[str, str]] = []
    if doc.tables:
        for row in doc.tables[0].rows[1:]:
            cells = [norm(c.text) for c in row.cells]
            if len(cells) >= 2 and cells[0]:
                table_rows.append((cells[0], cells[1]))
    return {"title": title, "items": items, "rectors": table_rows}


def table_html(rows: list[tuple[str, str]], *, lang: str) -> str:
    if lang == "en":
        meta = COOPERATION_EN
        data_rows = meta["rectors"]
        th_name, th_role = meta["table_name"], meta["table_role"]
    else:
        data_rows = rows
        th_name, th_role = "Ad, Soyad", "Vəzifə"
    body = "".join(
        f"<tr><td>{esc(name)}</td><td>{esc(role)}</td></tr>" for name, role in data_rows
    )
    return (
        f'<div class="program-table-wrap" role="region" aria-label="{esc(th_role)}">'
        f'<table class="program-table program-table--wide">'
        f"<thead><tr><th scope=\"col\">{esc(th_name)}</th>"
        f"<th scope=\"col\">{esc(th_role)}</th></tr></thead>"
        f"<tbody>{body}</tbody></table></div>"
    )


def contributions_body(data: dict, *, lang: str) -> str:
    if lang == "en":
        paragraphs = COOPERATION_EN["paragraphs"]
    else:
        paragraphs = data["items"]
    parts: list[str] = []
    for i, text in enumerate(paragraphs):
        if i == 7 and text.rstrip().endswith(":"):
            parts.append(f'<p class="card-lead">{esc(text)}</p>')
            continue
        parts.append(f'<p class="card-text">{esc(text)}</p>')
    return "\n".join(parts)


def section_card(sec: dict, body: str, *, lang: str) -> str:
    if lang == "en":
        titles = {
            "contributions": COOPERATION_EN["section_contributions"],
            "university-rectors": COOPERATION_EN["section_rectors"],
        }
        title = titles[sec["id"]]
    else:
        title = sec["az_title"]
    return f"""
<article class="news-card forum-cooperation-card" id="{esc(sec["id"])}">
<div class="card-header">
<h2 class="card-title">{esc(title)}</h2>
</div>
<div class="card-body">
{body}
</div>
</article>"""


def toc_items(*, lang: str) -> str:
    parts: list[str] = []
    for sec in SECTIONS:
        if lang == "en":
            label = COOPERATION_EN[
                "section_contributions"
                if sec["id"] == "contributions"
                else "section_rectors"
            ]
        else:
            label = sec["az_title"]
        parts.append(f'<li><a href="#{esc(sec["id"])}">{esc(label)}</a></li>')
    return "\n".join(parts)


def page_html(data: dict, *, lang: str) -> str:
    if lang == "en":
        meta = COOPERATION_EN
        hero_h1 = meta["hero_h1"]
        panel_title = meta["panel_title"]
        panel_copy = meta["panel_copy"]
        breadcrumb = meta["breadcrumb"]
        sidebar_label = meta["sidebar_label"]
        sidebar_aria = meta["sidebar_aria"]
        page_title = meta["page_title"]
        meta_desc = meta["meta_description"]
        skip = "Skip to content"
        bc_home = "Home"
        bc_activities = "Activities"
        bc_forum = "Forum 2024"
        footer_brand = "World Association of Azerbaijani Scientists"
        bc_aria = "Breadcrumb"
        panel_aria = "Contributions and cooperation summary"
        doc_lead = meta["doc_title"]
        footer_html = FORUM_FOOTER_EN
    else:
        hero_h1 = "Töhfələr <span>və əməkdaşlıq</span>"
        panel_title = "Tərəfdaşlar və töhfə verənlər"
        panel_copy = (
            "Dövlət qurumları, media, DAAB rəhbərliyi, universitetlər və alimlər — "
            "forumun reallaşmasına dəstək olanlar."
        )
        breadcrumb = "Töhfələr və əməkdaşlıq"
        sidebar_label = "🤝 Töhfələr və əməkdaşlıq"
        sidebar_aria = "Töhfələr menyusunu aç"
        page_title = "Töhfələr və əməkdaşlıq — DAAB"
        meta_desc = (
            "Xaricdə yaşayan azərbaycanlı alimlərin I Forumuna töhfə verən qurumlar və tərəfdaşlar."
        )
        skip = "Məzmuna keç"
        bc_home = "Ana səhifə"
        bc_activities = "Fəaliyyətimiz"
        bc_forum = "Forum 2024"
        footer_brand = "Dünya Azərbaycanlı Alimlər Birliyi"
        bc_aria = "Səhifə yolu"
        panel_aria = "Töhfələr və əməkdaşlıq haqqında qısa məlumat"
        doc_lead = data["title"]
        footer_html = f"""<footer class="footer-pro">
<div class="footer-inner">
<div class="footer-brand"><h3>{esc(footer_brand)}</h3></div>
<div class="footer-grid">
<div class="footer-col"><h4 class="footer-title">Əlaqə</h4><div class="footer-item">✉ <a href="mailto:bilik.birlik@gmail.com">bilik.birlik@gmail.com</a></div><div class="footer-item">☎ <span>+90 555 147 46 74</span></div><div class="footer-item">🌐 <a href="https://daab-waas.com" rel="noopener noreferrer" target="_blank">daab-waas.com</a></div></div>
<div class="footer-col"><h4 class="footer-title">Ünvan</h4><p class="footer-address">Feneryolu Mahallesi<br/>Gazi Muhtar Paşa Sokak No:44<br/>Kadıköy, İstanbul, Türkiyə</p></div>
<div class="footer-col"><h4 class="footer-title">Rəhbərlik</h4><p class="footer-leader"><strong>Prof. Dr. Məsud Əfəndiyev</strong><br/>DAAB İdarə Heyətinin Sədri<br/>Germany — James D. Murray Distinguished Professor</p></div>
</div>
</div>
<div class="footer-bottom">© 2026 DAAB / WAAS — All Rights Reserved</div>
</footer>"""

    contrib_body = contributions_body(data, lang=lang)
    rectors_body = table_html(data["rectors"], lang=lang)
    cards = section_card(SECTIONS[0], contrib_body, lang=lang) + section_card(
        SECTIONS[1], rectors_body, lang=lang
    )
    toc = toc_items(lang=lang)
    nav = forum_nav_strip(lang, active_nav_id="forum-2024")

    return f"""<!DOCTYPE html>
<html lang="{lang}" data-daab-lang="{lang}" data-daab-asset-root="{ASSET}" data-daab-page-id="{PAGE_ID}" data-daab-nav-mount="1">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0, viewport-fit=cover"/>
<title>{esc(page_title)}</title>
<meta name="description" content="{esc(meta_desc)}"/>
<link rel="preconnect" href="https://fonts.googleapis.com"/>
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin/>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&family=Playfair+Display:wght@700;800&display=swap" rel="stylesheet"/>
<link href="{ASSET}css/daab-common.css?v=24" rel="stylesheet"/>
<link href="{ASSET}css/daab-mobile.css?v=5" rel="stylesheet"/>
<link href="{ASSET}css/daab-search.css?v=1" rel="stylesheet"/>
<link href="{ASSET}css/daab-back-to-top.css?v=1" rel="stylesheet"/>
<link href="{ASSET}css/daab-lang.css?v=10" rel="stylesheet"/>
<link href="{ASSET}css/daab-nav-mega.css?v=13" rel="stylesheet"/>
<link href="{ASSET}css/daab-hero-summary.css?v=1" rel="stylesheet"/>
<link href="{ASSET}css/daab-sidebar-widget.css?v=3" rel="stylesheet"/>
<link href="{ASSET}css/daab-activities-layout.css?v=12" rel="stylesheet"/>
<link href="{ASSET}css/daab-forum-content.css?v=19" rel="stylesheet"/>
<script src="{ASSET}js/daab-mobile.js?v=1" defer></script>
<script src="{ASSET}js/daab-back-to-top.js?v=2" defer></script>
<script src="{ASSET}js/daab-i18n.js?v=12" defer></script>
<script src="{ASSET}js/daab-lang-position.js?v=7" defer></script>
<script src="{ASSET}js/daab-nav.js?v=10" defer></script>
<script src="{ASSET}js/daab-primary-nav.js?v=10" defer></script>
<script src="{ASSET}js/daab-section-nav.js?v=7" defer></script>
<script src="{ASSET}js/daab-shell.js?v=11" defer></script>
<script src="{ASSET}js/daab-search.js?v=4" defer></script>
</head>
<body>
<a class="skip" href="#content">{esc(skip)}</a>
{nav}
<div class="breadcrumbs forum-breadcrumbs" role="navigation" aria-label="{esc(bc_aria)}">
<a href="../../index.html">{esc(bc_home)}</a><span aria-hidden="true">›</span><a href="../../activities.html">{esc(bc_activities)}</a><span aria-hidden="true">›</span><a href="index.html">{esc(bc_forum)}</a><span aria-hidden="true">›</span><span class="forum-breadcrumbs-current" aria-current="page">{esc(breadcrumb)}</span>
</div>
<header class="page-hero">
<div class="hero-wrap shell">
<section class="hero-copy">
<h1>{hero_h1}</h1>
<p class="hero-doc-lead">{esc(doc_lead)}</p>
</section>
<aside aria-label="{esc(panel_aria)}" class="hero-panel">
<div class="panel-card">
<h2 class="panel-title">{esc(panel_title)}</h2>
<div class="panel-copy">{esc(panel_copy)}</div>
</div>
</aside>
</div>
</header>
<div class="content-wrap">
<aside class="sidebar">
<div class="sidebar-widget">
<div class="widget-head"><span>{sidebar_label}</span><button aria-controls="cooperationTOC" aria-expanded="false" aria-label="{esc(sidebar_aria)}" class="events-menu-toggle" type="button"><span></span><span></span><span></span></button></div>
<div class="widget-body">
<ul class="timeline-list" id="cooperationTOC">
{toc}
</ul>
</div>
</div>
</aside>
<main class="news-feed main" id="content">
{cards}
</main>
</div>
{footer_html}
{SIDEBAR_SCRIPT}
</body>
</html>
"""


def main() -> None:
    data = parse_docx(Document(DOCX))
    OUT_AZ.parent.mkdir(parents=True, exist_ok=True)
    OUT_AZ.write_text(page_html(data, lang="az"), encoding="utf-8", newline="\n")
    OUT_EN.write_text(page_html(data, lang="en"), encoding="utf-8", newline="\n")
    print(f"wrote {OUT_AZ.relative_to(ROOT)}")
    print(f"wrote {OUT_EN.relative_to(ROOT)} ({len(data['items'])} items, {len(data['rectors'])} rectors)")


if __name__ == "__main__":
    main()
