#!/usr/bin/env python3
"""Write documents/Legal/05-Legal-Compliance-Gap-Review.docx."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

try:
    from _paths import ROOT
except ImportError:
    from helpers._paths import ROOT  # type: ignore

OUT = ROOT / "documents" / "Legal" / "05-Legal-Compliance-Gap-Review.docx"


def add_toc_field(paragraph) -> None:
    """Insert a Word TOC field (update in Word: right-click → Update Field)."""
    run = paragraph.add_run()
    r = run._r

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")

    # Placeholder text until Word updates the field
    text = OxmlElement("w:t")
    text.text = "Right-click here → Update Field → Update entire table"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    r.append(begin)
    r.append(instr)
    r.append(sep)
    r.append(text)
    r.append(end)


def main() -> int:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str, *, bold: bool = False) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold

    def bullets(items: list[str]) -> None:
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    def numbered(items: list[str]) -> None:
        for item in items:
            doc.add_paragraph(item, style="List Number")

    doc.add_heading("DAAB / WAAS Legal Pages — Compliance Gap Review", 0)
    subtitle = doc.add_paragraph()
    subtitle.add_run(
        "Internally recognized legal standards checklist and recommendations for the public website legal suite "
        "(Privacy Notice, Terms of Use, Cookie Policy, Legal Notice / Imprint)."
    ).italic = True

    meta = doc.add_paragraph()
    meta.add_run(f"Date: {date.today().isoformat()}\n").bold = True
    meta.add_run(
        "Scope: Public legal pages on daab-waas.com (AZ/EN) and closely related site behaviour "
        "(cookie banner, membership form, analytics).\n"
    )
    meta.add_run(
        "Primary home-law frame: Türkiye — Personal Data Protection Law No. 6698 (KVKK / PDPL).\n"
    )
    meta.add_run(
        "International reference frames: GDPR / UK GDPR expectations, ePrivacy-style cookie practice, "
        "common imprint / terms standards for association websites."
    )

    note = doc.add_paragraph()
    note.add_run(
        "IMPORTANT DISCLAIMER: This document is a practical editorial and compliance-gap review for internal "
        "planning. It is not legal advice. Turkish counsel (and, where relevant, EU/UK counsel) should review "
        "and approve any wording before it is treated as legally binding or “compliant.”"
    ).bold = True

    contents_title = doc.add_paragraph()
    contents_title.add_run("Contents").bold = True
    contents_title.runs[0].font.size = Pt(16)
    toc_note = doc.add_paragraph()
    toc_note.add_run(
        "In Microsoft Word: right-click the table of contents field below → Update Field → Update entire table "
        "(to refresh page numbers). A static outline follows for immediate reading."
    ).italic = True
    toc_para = doc.add_paragraph()
    add_toc_field(toc_para)

    # Static outline (visible immediately; complements the Word TOC field)
    p("Document outline:", bold=True)
    bullets(
        [
            "1. Executive summary",
            "2. What you already have (solid base)",
            "3. Highest-priority gaps",
            "    3.1 Complaint authority",
            "    3.2 International transfers",
            "    3.3 Retention periods",
            "    3.4 Processors",
            "    3.5 GDPR / EU audience posture",
            "    3.6 Membership checkbox language",
            "4. Important content still missing on the pages",
            "5. Operational / technical compliance (not just text)",
            "6. Nice-to-have for internationally recognized polish",
            "7. Suggested priority order",
            "8. Bottom line",
            "9. Recommended next drafting step",
            "10. Document control",
            "11. Concrete draft text for Privacy Notice and Cookie Policy (AZ / EN)",
            "    11.1 How to use these drafts",
            "    11.2 Privacy Notice — English draft additions / replacements",
            "    11.3 Privacy Notice — Azerbaijani draft additions / replacements",
            "    11.4 Cookie Policy — English draft additions / replacements",
            "    11.5 Cookie Policy — Azerbaijani draft additions / replacements",
            "    11.6 Suggested membership form checkbox wording (AZ / EN)",
            "    11.7 Implementation checklist after counsel approval",
            "12. Recommended next steps (action plan)",
            "    12.1 Counsel review",
            "    12.2 Confirm facts internally",
            "    12.3 Apply approved text to the site",
            "    12.4 Technical alignment",
            "    12.5 Medium priority (after the above)",
            "    12.6 Internal compliance pack",
            "    12.7 Optional follow-up with the website team",
            "13. Implications of deploying the site to production",
            "    13.1 What production changes",
            "    13.2 What is already in decent shape for go-live",
            "    13.3 Main implications / residual risks if deployed as-is",
            "    13.4 What is not blocked by the legal gaps",
            "    13.5 Practical recommendation",
            "    13.6 Board checklist — safe to deploy / fix within 30 days",
        ]
    )
    doc.add_page_break()

    h("1. Executive summary")
    p(
        "Publicly, the site already has a credible baseline for a Turkish association website: imprint/operator "
        "identity, privacy notice, cookie policy with analytics opt-in, terms of use, and discoverability via "
        "footer and Legal menu. What is still missing for genuinely internationally recognized compliance is "
        "mostly: a supervisory complaint path, concrete transfer/retention/processor detail, a clearer EU/GDPR "
        "posture, cleaner consent vs contract wording on membership, and the internal contracts/processes that "
        "make the policies true in practice."
    )

    h("2. What you already have (solid base)")
    p("The following public-facing elements are already in place and form a solid foundation:")
    bullets(
        [
            "Legal notice / imprint with operator identity, address, and contact details",
            "Privacy notice covering controller, data categories, purposes, legal bases, rights, children, and security",
            "Cookie policy plus consent banner with analytics opt-in (essential always on)",
            "Terms of use covering intellectual property, acceptable use, liability, and governing law",
            "Footer and Legal menu discoverability for the four legal documents",
            "Membership form privacy acknowledgment / confirmation",
            "Plain-language explanations and key takeaways on legal pages (usability strength)",
        ]
    )

    h("3. Highest-priority gaps")

    h("3.1 Complaint authority (almost always expected)", 2)
    p(
        "The privacy notice lists data-subject rights, but does not clearly state where to complain if a request "
        "is refused or mishandled."
    )
    bullets(
        [
            "Missing for Türkiye: right to lodge a complaint with the Kişisel Verileri Koruma Kurulu (Personal Data Protection Board).",
            "If EU visitors are in scope: also provide a path to an EU supervisory authority (GDPR Art. 77).",
        ]
    )

    h("3.2 International transfers are too vague", 2)
    p(
        "Current wording that Google Analytics may transfer data abroad and that “measures aligned with PDPL are "
        "intended” is not enough for internationally recognized practice."
    )
    p("Missing:", bold=True)
    bullets(
        [
            "Named transfer mechanism (e.g. adequacy decision / Standard Contractual Clauses / Google’s applicable terms)",
            "Clear statement of Google’s role (processor vs independent controller — state your position clearly)",
            "Link to Google’s privacy/terms documentation for GA4",
            "Concrete cookie names, providers, and retention periods (e.g. _ga, _ga_*)",
        ]
    )

    h("3.3 Retention periods are not operational", 2)
    p("“As long as necessary / reasonable period” is a common weak spot.")
    p("Missing concrete or criteria-based periods for:", bold=True)
    bullets(
        [
            "Membership applications",
            "Enquiry emails",
            "Server / access logs",
            "Analytics retention setting in GA4",
            "Published profiles after withdrawal of permission",
        ]
    )

    h("3.4 Processors are unnamed", 2)
    p(
        "Hosting, email, and IT providers are mentioned only generically. International practice usually expects "
        "named key processors, or at least a clear “list available on request” plus categories."
    )
    p("Especially important for:", bold=True)
    bullets(
        [
            "Hosting provider",
            "Mail handling (mail.php / inbox processing)",
            "Any CDN",
            "Google Analytics 4",
        ]
    )

    h("3.5 GDPR / EU audience posture is missing", 2)
    p(
        "The audience is international scientists. If membership is offered to, or profiles are published about, "
        "people in the EU/EEA/UK, GDPR or UK GDPR may apply even though the association is Turkish."
    )
    p("Missing from the privacy notice:", bold=True)
    bullets(
        [
            "Whether GDPR / UK GDPR may also apply",
            "Mapping of rights (access, erasure, restriction, portability, objection)",
            "Whether an EU/UK representative is required (Art. 27) — often relevant if you regularly deal with EU data subjects and have no EU establishment",
            "Clear statement that Turkish law remains primary for the association, with mandatory foreign rights preserved where applicable",
        ]
    )

    h("3.6 Membership checkbox language is risky", 2)
    p(
        "The form effectively asks people to “agree” to the privacy notice for processing. Under modern privacy "
        "regimes, best practice is:"
    )
    bullets(
        [
            "Reading the notice is not blanket consent for everything",
            "Membership processing should rely mainly on contract / application necessity",
            "Consent should be reserved for optional uses (analytics, public profile publication, marketing)",
        ]
    )
    p("Current wording can look like bundled consent and should be reviewed with counsel.")

    h("4. Important content still missing on the pages")
    p("The table below summarises page-level content gaps.")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Area"
    hdr[1].text = "What is missing"
    for area, missing in [
        (
            "Imprint",
            "Dernek registration / file number, city of registration, tax number if used publicly; optionally MERSİS-style identifiers if applicable",
        ),
        (
            "Privacy",
            "Dedicated privacy contact / DPO-style contact (even if same inbox with subject “Privacy / KVKK”)",
        ),
        (
            "Privacy",
            "Sources of data (directly from user vs third parties / public sources)",
        ),
        (
            "Privacy",
            "Sensitive data stance (CVs/photos may include special-category or biometric-adjacent data)",
        ),
        (
            "Privacy",
            "Breach notification stance (when you will notify the Board / affected people)",
        ),
        (
            "Cookies",
            "Exact GA cookies, duration, first/third party, withdrawal, how often consent is re-asked",
        ),
        (
            "Cookies / banner",
            "Keep “Reject analytics / Essential only” as clear and prominent as accept",
        ),
        (
            "Terms",
            "Formal IP complaint / takedown procedure with response timeline (currently informal email only)",
        ),
        (
            "Terms / Privacy",
            "Donation & sponsorship data flows (bank details, receipts, or payment processors, if any)",
        ),
        (
            "Accessibility",
            "Public accessibility statement (WCAG 2.x commitment) — increasingly expected internationally",
        ),
        (
            "Related documents",
            "Privacy and Terms pages lack a Related documents section (cookies/legal-notice already have one)",
        ),
    ]:
        row = table.add_row().cells
        row[0].text = area
        row[1].text = missing

    doc.add_paragraph()

    h("5. Operational / technical compliance (not just text)")
    p("These matter as much as the published HTML pages:")
    numbered(
        [
            "GA4 configuration — consent mode, disable ads/personalization, document IP anonymization / data-retention settings, no tags before consent.",
            "Processor contracts — Google data processing / controller-processor terms; hosting and email DPAs.",
            "Internal ROPA — record of processing activities (required under serious privacy regimes; not necessarily a public page).",
            "Email for CVs/photos — disclose that email is not a secure channel; prefer encrypted upload or clearer instructions.",
            "Access control — who in the association can open membership mail / form submissions.",
            "Incident process — who decides, Board notification checklist under KVKK when applicable.",
            "Consent records — store timestamp/version of privacy notice accepted with applications.",
        ]
    )

    h("6. Nice-to-have for internationally recognized polish")
    bullets(
        [
            "Version history or changelog of legal pages (not only “Last updated”)",
            "Keep plain-language summaries synced when substantive legal text changes",
            "Multilingual consistency: EN “PDPL” / AZ “KVKK” is fine, but banners and tooltips should use the same naming as the pages",
            "Explicit “no advertising / no sale / no profiling for ads” statement in one prominent place",
            "If YouTube, maps, or social embeds are added later: update the cookie policy and banner categories before go-live",
        ]
    )

    h("7. Suggested priority order")
    numbered(
        [
            "Add Board complaint path + clearer international-transfer / GA disclosures",
            "Put concrete retention rules and named processors into the privacy notice",
            "Fix membership form legal basis / consent wording with counsel",
            "Decide GDPR/EU posture and add a short section if you target EU scientists",
            "Strengthen imprint identifiers + formal takedown procedure",
            "Add accessibility statement + donation data flows (if applicable)",
            "Back the public pages with internal docs (ROPA, DPAs, breach playbook)",
        ]
    )

    h("8. Bottom line")
    p(
        "Publicly, DAAB/WAAS already looks like a serious association site with the right four legal documents. "
        "What remains for genuine internationally recognized compliance is:"
    )
    bullets(
        [
            "Supervisory complaint path",
            "Concrete transfer, retention, and processor detail",
            "Clearer EU / GDPR posture where relevant",
            "Cleaner consent vs contract wording on membership",
            "Internal contracts and processes that make the published policies accurate in practice",
        ]
    )

    h("9. Recommended next drafting step")
    p(
        "Section 11 below contains concrete AZ/EN draft text for the highest-priority privacy and cookie updates. "
        "After counsel review, implement approved wording via helpers/_build_legal_pages.py and align the membership "
        "application checkbox with the chosen legal bases."
    )

    h("10. Document control")
    bullets(
        [
            "Document ID: 05-Legal-Compliance-Gap-Review",
            "Location: documents/Legal/",
            "Related public pages: /az|en/privacy.html, /cookies.html, /terms.html, /legal-notice.html",
            "Related source builder: helpers/_build_legal_pages.py",
            "Regenerate this file: python helpers/_write_legal_compliance_review_docx.py",
            "Status: Internal planning draft + proposed page wording + next-steps + production implications / board checklist — pending legal counsel review",
        ]
    )

    # ------------------------------------------------------------------
    # Part B — concrete draft text
    # ------------------------------------------------------------------
    doc.add_page_break()
    h("11. Concrete draft text for Privacy Notice and Cookie Policy (AZ / EN)")
    note2 = doc.add_paragraph()
    note2.add_run(
        "STATUS: Proposed website wording only. Items in [SQUARE BRACKETS] must be confirmed or replaced by the "
        "association / counsel before publication. This text has not yet been applied to the live HTML pages."
    ).bold = True

    h("11.1 How to use these drafts", 2)
    bullets(
        [
            "Insert or replace the matching sections on /az/privacy.html, /en/privacy.html, /az/cookies.html, /en/cookies.html.",
            "Keep existing takeaways/plain-language boxes unless counsel asks to revise them.",
            "After approval, update helpers/_build_legal_pages.py (source of truth) and rebuild.",
            "Also update the membership form checkbox wording in Section 11.6 (application pages).",
        ]
    )

    h("11.2 Privacy Notice — English draft additions / replacements", 2)

    h("A. Privacy contact (add after §1 Data controller)", 3)
    p(
        "Privacy contact. For personal-data requests under this notice (access, correction, erasure, withdrawal of "
        "consent, or questions), please email info@daab-waas.com with the subject line “Privacy / PDPL request”. "
        "We may ask for information reasonably needed to verify your identity. [If a dedicated Data Protection "
        "Officer or privacy lead is appointed, add name/role here.]"
    )

    h("B. Sources of data (new section; suggested as §2A after data categories)", 3)
    p("We obtain personal data from the following sources, as applicable:")
    bullets(
        [
            "Directly from you — for example when you browse the Site (technical data), submit a membership application, send an email, or ask us to publish a profile or contributed work;",
            "From you or a person authorised by you — for example a CV, photograph or biography submitted for membership or publication;",
            "From publicly available academic or organisational sources — only where we have a lawful basis and, for profile publication, appropriate permission.",
        ]
    )

    h("C. Recipients, processors and international transfers (replace / expand current §5)", 3)
    p(
        "We do not sell personal data. We may share personal data, limited to what is needed for the purpose, with:"
    )
    bullets(
        [
            "WAAS membership / board commission members and authorised volunteer team members who handle applications or content;",
            "IT and communications providers acting on our instructions as processors, including: [HOSTING PROVIDER — legal name and country]; [EMAIL / MAILBOX PROVIDER — legal name and country]; and, where used, [CDN OR OTHER IT PROVIDER];",
            "Google LLC (and its affiliates) — Google Analytics 4, only if you consent to analytics cookies; Google acts as a processor for this analytics service under Google’s applicable terms;",
            "Competent public authorities where required by law;",
            "Public visitors to the Site, for information you have agreed to publish (for example a scientist profile, photograph or contributed work).",
        ]
    )
    p(
        "International transfers. Some providers (in particular Google LLC for Google Analytics 4) may process data "
        "in countries outside Türkiye, including the United States. Where personal data is transferred abroad, we "
        "rely on safeguards recognised under PDPL and related rules, including [adequacy decision / Standard "
        "Contractual Clauses / Google’s data processing terms — COUNSEL TO CONFIRM THE CORRECT MECHANISM]. "
        "Further information about Google’s processing is available in Google’s privacy documentation: "
        "https://policies.google.com/privacy and Google Analytics terms: "
        "https://marketingplatform.google.com/about/analytics/terms/us/."
    )
    p(
        "A current list of key processors is available on request at info@daab-waas.com (subject “Privacy — processors”)."
    )

    h("D. Retention (replace current §6)", 3)
    p(
        "We keep personal data only as long as needed for the purpose and any applicable legal retention duty. "
        "Unless a longer period is required by law or an ongoing dispute, we apply the following guidance periods "
        "(counsel/association to confirm):"
    )
    bullets(
        [
            "Membership applications and related correspondence: up to [24 months] after the application is decided or after our last substantive contact about it, whichever is later;",
            "General enquiry, donation or sponsorship emails: up to [24 months] after the last substantive contact;",
            "Server / security logs: typically up to [90 days], unless needed longer to investigate a security incident;",
            "Google Analytics data: according to the retention setting configured for our GA4 property (currently [CONFIRM GA4 RETENTION, e.g. 14 months]) and your consent status;",
            "Published profiles and contributed works: for as long as publication remains authorised; after a valid removal request, we aim to remove or anonymise the material within [30 days], subject to technical and backup cycles;",
            "Cookie consent records stored in your browser: until you clear site data or change your choice; we do not keep a separate server-side copy of the banner choice unless [CONFIRM IF APPLICATION/SERVER LOGS STORE CONSENT].",
        ]
    )

    h("E. Your rights — add complaint path (expand current §7)", 3)
    p(
        "In addition to the rights listed above, if you believe your personal data has been processed unlawfully, "
        "you may lodge a complaint with the Personal Data Protection Authority of Türkiye "
        "(Kişisel Verileri Koruma Kurulu — “the Board”). Contact details and procedures are published by the Board: "
        "https://www.kvkk.gov.tr. We encourage you to contact us first at info@daab-waas.com so we can try to resolve "
        "your concern."
    )
    p(
        "If the GDPR or UK GDPR applies to you (see the section “Visitors and members in the EU/EEA/UK” below), "
        "you may also lodge a complaint with your local supervisory authority."
    )

    h("F. Visitors and members in the EU/EEA/UK (new section)", 3)
    p(
        "WAAS is established in Türkiye and this notice is primarily based on Türkiye’s Personal Data Protection "
        "Law No. 6698 (PDPL / KVKK). Because our scientific community is international, some visitors and members "
        "may be located in the European Union, European Economic Area or United Kingdom."
    )
    p(
        "Where the GDPR or UK GDPR applies to our processing of your personal data, you may also have rights under "
        "those laws, including rights of access, rectification, erasure, restriction of processing, data portability "
        "and objection, and the right to withdraw consent where processing is based on consent. Those rights apply "
        "in addition to, and do not reduce, any mandatory rights you have under PDPL."
    )
    p(
        "[COUNSEL TO CONFIRM:] WAAS [has appointed / has not appointed] an EU/UK representative under GDPR/UK GDPR "
        "Article 27. [If appointed, insert representative name, address and email.] Until a representative is "
        "appointed, please use the privacy contact details in this notice."
    )

    h("G. Sensitive data (new short section)", 3)
    p(
        "We do not seek to collect special categories of personal data (such as health, religion, political opinions "
        "or biometric data used for identification) through the Site. Please do not include such information in "
        "applications, CVs or emails unless it is strictly necessary and you choose to provide it. Photographs and "
        "biographies published with permission are processed for the purpose of presenting scientists and association "
        "activities; we rely on the permission / consent arrangements described in this notice."
    )

    h("H. Personal data breaches (new short section)", 3)
    p(
        "If we become aware of a personal data breach that is likely to affect your rights, we will assess the "
        "incident and, where required under PDPL and related rules, notify the Board and/or affected individuals "
        "without undue delay. You can report a suspected incident to info@daab-waas.com with the subject "
        "“Privacy — security incident”."
    )

    h("I. Membership applications — legal bases clarification (expand current membership section)", 3)
    p(
        "Data you submit in the online membership application is used to receive, assess and respond to your "
        "application and related membership correspondence. The primary legal bases are necessity for taking steps "
        "at your request prior to a membership relationship / contractual process, and our legitimate interest in "
        "assessing applications for a scientific association, provided this does not override your rights."
    )
    p(
        "The checkbox on the application form confirms that you have read this Privacy Notice. It is not a blanket "
        "consent for all processing described on the Site. Separate consent is used where required — for example "
        "optional analytics cookies, and publication of a public profile, photograph or contributed works."
    )
    p(
        "CVs and photographs sent by ordinary email are transmitted over email systems that may not be end-to-end "
        "encrypted. Please avoid including unnecessary sensitive information. [OPTIONAL: Prefer submission via "
        "secure upload — COUNSEL/IT TO CONFIRM IF AVAILABLE.]"
    )

    h("11.3 Privacy Notice — Azerbaijani draft additions / replacements", 2)

    h("A. Məxfilik əlaqəsi (§1-dən sonra)", 3)
    p(
        "Məxfilik əlaqəsi. Bu bildiriş üzrə şəxsi məlumat sorğuları (məlumat alma, düzəliş, silinmə, razılığın "
        "geri götürülməsi və ya suallar) üçün info@daab-waas.com ünvanına “Məxfilik / KVKK sorğusu” mövzusu ilə "
        "yazın. Şəxsiyyətinizi təsdiqləmək üçün ağlabatan məlumat istəyə bilərik. [Əgər ayrıca məlumatların "
        "qorunması məsul şəxsi təyin olunubsa, ad/vəzifəni bura əlavə edin.]"
    )

    h("B. Məlumat mənbələri (yeni bölmə)", 3)
    p("Şəxsi məlumatları, vəziyyətdən asılı olaraq, aşağıdakı mənbələrdən əldə edirik:")
    bullets(
        [
            "Birbaşa sizdən — məsələn saytı ziyarət etdikdə (texniki məlumat), üzvlük müraciəti göndərdikdə, e-poçt yazdıqda və ya profil/töhfə əsərinin dərcini xahiş etdikdə;",
            "Sizdən və ya sizin səlahiyyət verdiyiniz şəxsdən — məsələn üzvlük və ya dərc üçün göndərilən CV, foto və ya bioqrafiya;",
            "İctimai əlçatan akademik və ya təşkilati mənbələrdən — yalnız qanuni əsas olduqda və profil dərcı üçün müvafiq icazə olduqda.",
        ]
    )

    h("C. Alıcılar, məlumat emaledənlər və beynəlxalq ötürmə (mövcud §5-in genişləndirilməsi)", 3)
    p(
        "Şəxsi məlumatları satmırıq. Məlumatlarınız yalnız məqsədə uyğun həcmdə aşağıdakılara ötürülə bilər:"
    )
    bullets(
        [
            "Üzvlük / idarə heyəti komissiyası üzvlərinə və müraciət və ya məzmunla məşğul olan səlahiyyətli könüllü komanda üzvlərinə;",
            "Təlimatlarımız əsasında məlumat emaledən (processor) qismində çıxış edən İT və rabitə xidməti göstərənlərə, o cümlədən: [HOSTİNQ PROVAYDERİ — hüquqi ad və ölkə]; [E-POÇT / POÇT QUTUSU PROVAYDERİ — hüquqi ad və ölkə]; və istifadə olunarsa [CDN VƏ YA DİGƏR İT PROVAYDER];",
            "Google LLC (və filialları) — Google Analytics 4, yalnız analitika kukilərinə razılıq verdikdə; Google bu analitika xidməti üçün tətbiq olunan şərtlər əsasında məlumat emaledən qismində çıxış edir;",
            "Qanun tələb etdikdə səlahiyyətli dövlət orqanlarına;",
            "Dərcinə razı olduğunuz məlumatlar üçün saytın ictimai ziyarətçilərinə (məsələn, alim profili, foto, töhfə əsərləri).",
        ]
    )
    p(
        "Beynəlxalq ötürmə. Bəzi xidmət göstərənlər (xüsusən Google Analytics 4 üçün Google LLC) məlumatları "
        "Türkiyədən kənar ölkələrdə, o cümlədən ABŞ-da emal edə bilər. Şəxsi məlumatlar xaricə ötürülərsə, KVKK "
        "və əlaqəli qaydalara uyğun mühafizə tədbirlərinə əsaslanırıq, o cümlədən [yetərlilik qərarı / Standart "
        "Müqavilə Müddəaları / Google-un məlumat emalı şərtləri — DÜZGÜN MEXANİZMİ VƏKİL TƏSDİQLƏMƏLİDİR]. "
        "Google-un emalı haqqında əlavə məlumat: https://policies.google.com/privacy və Google Analytics şərtləri: "
        "https://marketingplatform.google.com/about/analytics/terms/us/."
    )
    p(
        "Əsas məlumat emaledənlərin cari siyahısı info@daab-waas.com ünvanına “Məxfilik — emaledənlər” mövzusu ilə "
        "sorğu göndərməklə əldə edilə bilər."
    )

    h("D. Saxlama müddətləri (mövcud §6-nın əvəzi)", 3)
    p(
        "Şəxsi məlumatları yalnız məqsəd üçün lazım olan müddət və tətbiq olunan qanuni saxlama öhdəliyi qədər "
        "saxlayırıq. Qanun və ya davam edən mübahisə daha uzun müddət tələb etmədikcə, aşağıdakı istiqamətləndirici "
        "müddətləri tətbiq edirik (birlik/vəkil tərəfindən təsdiqlənməlidir):"
    )
    bullets(
        [
            "Üzvlük müraciətləri və əlaqəli yazışmalar: müraciət barədə qərar verildikdən və ya son mahiyyətli əlaqədən sonra (hansı sonradırsa) ən çoxu [24 ay];",
            "Ümumi sorğu, ianə və ya sponsorluq e-poçtları: son mahiyyətli əlaqədən sonra ən çoxu [24 ay];",
            "Server / təhlükəsizlik jurnalları: adətən ən çoxu [90 gün], təhlükəsizlik hadisəsinin araşdırılması üçün daha uzun tələb olunmazsa;",
            "Google Analytics məlumatları: GA4 mülkiyyətimizdə təyin olunmuş saxlama parametrinə uyğun (hazırda [GA4 SAXLAMA MÜDDƏTİNİ TƏSDİQLƏYİN, məs. 14 ay]) və razılıq statusunuza uyğun;",
            "Dərc olunmuş profillər və töhfə əsərləri: dərc icazəsi qüvvədə qaldığı müddət; etibarlı silinmə tələbindən sonra materialı [30 gün] ərzində silməyə və ya anonimləşdirməyə çalışırıq (texniki və ehtiyat nüsxə dövrləri nəzərə alınmaqla);",
            "Brauzerdə saxlanan kuki razılığı: sayt məlumatlarını təmizləyənə və ya seçiminizi dəyişənə qədər; banner seçiminin ayrıca server nüsxəsini saxlamırıq, əgər [ƏRİZƏ/SERVER JURNALLARI RAZILIĞI SAXLAYIRSA — TƏSDİQLƏYİN] istisna deyilsə.",
        ]
    )

    h("E. Hüquqlar — şikayət yolu (mövcud §7-yə əlavə)", 3)
    p(
        "Yuxarıda göstərilən hüquqlarla yanaşı, şəxsi məlumatlarınızın qanunsuz işləndiyini düşünürsinizsə, "
        "Türkiyənin Kişisel Verileri Koruma Kurulu’na (“Kurul”) şikayət edə bilərsiniz. Əlaqə və prosedurlar "
        "Kurulun saytında dərc olunur: https://www.kvkk.gov.tr. Əvvəlcə info@daab-waas.com ünvanına yazmağınızı "
        "tövsiyə edirik ki, məsələni həll etməyə çalışaq."
    )
    p(
        "Əgər sizin üçün GDPR və ya UK GDPR tətbiq olunursa (aşağıdakı “Aİ/EEA/BK-dəki ziyarətçi və üzvlər” "
        "bölməsinə baxın), yerli nəzarət orqanına da şikayət edə bilərsiniz."
    )

    h("F. Aİ/EEA/BK-dəki ziyarətçi və üzvlər (yeni bölmə)", 3)
    p(
        "WAAS Türkiyədə yaradılıb və bu bildiriş əsasən 6698 saylı Şəxsi Məlumatların Qorunması Qanununa "
        "(KVKK / PDPL) əsaslanır. Elmi icmamız beynəlxalq olduğu üçün bəzi ziyarətçi və üzvlər Avropa İttifaqı, "
        "Avropa İqtisadi Bölgəsi və ya Birləşmiş Krallıqda ola bilər."
    )
    p(
        "Şəxsi məlumatlarınızın işlənməsinə GDPR və ya UK GDPR tətbiq olunduğu hallarda, həmin qanunlar üzrə də "
        "hüquqlarınız ola bilər — o cümlədən məlumat alma, düzəliş, silinmə, işlənmənin məhdudlaşdırılması, "
        "məlumatların daşınabilirliyi, etiraz və razılığa əsaslanan hallarda razılığın geri götürülməsi. Bu hüquqlar "
        "KVKK üzrə məcburi hüquqlarınızı azaltmır."
    )
    p(
        "[VƏKİL TƏSDİQLƏMƏLİDİR:] WAAS GDPR/UK GDPR-nin 27-ci maddəsi üzrə Aİ/BK nümayəndəsi [təyin edib / təyin "
        "etməyib]. [Təyin olunubsa, ad, ünvan və e-poçtu yazın.] Nümayəndə təyin olunana qədər bu bildirişdəki "
        "məxfilik əlaqə vasitələrindən istifadə edin."
    )

    h("G. Həssas məlumatlar (qısa yeni bölmə)", 3)
    p(
        "Sayt vasitəsilə xüsusi kateqoriyalı şəxsi məlumatları (məsələn, sağlamlıq, din, siyasi baxışlar və ya "
        "identifikasiya üçün biometrik məlumatlar) qəsdən toplamırıq. Ərizə, CV və ya e-poçtlara belə məlumatları "
        "yalnız ciddi zəruri olduqda və öz seçiminizlə əlavə etməyin. İcazə ilə dərc olunan foto və bioqrafiyalar "
        "alimlərin və birlik fəaliyyətinin təqdimatı məqsədi ilə işlənir; bu bildirişdə təsvir olunan icazə / "
        "razılıq qaydalarına əsaslanırıq."
    )

    h("H. Məlumat pozuntuları (qısa yeni bölmə)", 3)
    p(
        "Hüquqlarınıza təsir edə biləcək şəxsi məlumat pozuntusundan xəbərdar olsaq, hadisəni qiymətləndirəcək "
        "və KVKK ilə əlaqəli qaydalar tələb etdikdə Kurula və/və ya təsirə məruz qalan şəxslərə lüzumsuz "
        "gecikmədən məlumat verəcəyik. Şübhəli hadisəni info@daab-waas.com ünvanına “Məxfilik — təhlükəsizlik "
        "hadisəsi” mövzusu ilə bildirin."
    )

    h("I. Üzvlük müraciətləri — hüquqi əsasların dəqiqləşdirilməsi", 3)
    p(
        "Onlayn üzvlük müraciətində göndərdiyiniz məlumatlar müraciətinizi qəbul etmək, qiymətləndirmək və "
        "cavablandırmaq, habelə əlaqəli üzvlük yazışmaları üçün istifadə olunur. Əsas hüquqi əsaslar: sizin "
        "tələbinizlə üzvlük münasibətindən əvvəl addımlar atmaq / müqavilə prosesi üçün zərurət və elmi birliyə "
        "müraciətləri qiymətləndirməyə dair qanuni maraqlarımızdır — bu, sizin hüquqlarınızı üstələməmək şərtilə."
    )
    p(
        "Müraciət formasındakı qeyd xanası bu Məxfilik bildirişini oxuduğunuzu təsdiqləyir. Bu, saytda təsvir "
        "olunan bütün işlənmələr üçün ümumi razılıq deyil. Ayrı razılıq tələb olunan hallarda istifadə olunur — "
        "məsələn, istəyə bağlı analitika kukiləri və ictimai profil, foto və ya töhfə əsərlərinin dərcı."
    )
    p(
        "Adi e-poçtla göndərilən CV və fotolar ucdan-uca şifrələnməyən e-poçt sistemləri ilə ötürülə bilər. "
        "Lüzumsuz həssas məlumat əlavə etməyin. [İSTƏYƏ BAĞLI: Təhlükəsiz yükləmə yolu — İT/VƏKİL TƏSDİQLƏSİN.]"
    )

    h("11.4 Cookie Policy — English draft additions / replacements", 2)

    h("A. Expand the cookies table (replace current §2 table body as needed)", 3)
    p("Proposed table rows:")
    bullets(
        [
            "Essential — Purpose: core site functions, security, language preference, remembering cookie choice. Examples: daab-cookie-consent (localStorage); daab-lang (language). Party: first-party (WAAS). Duration: until cleared / changed. Consent: not required.",
            "Analytics (optional) — Purpose: aggregate visit statistics via Google Analytics 4 (measurement ID G-K778896NSH). Examples: _ga; _ga_<container-id> and related GA4 tags/cookies set by Google. Party: third-party (Google LLC). Typical duration: up to [2 years] for _ga / according to Google and our GA4 settings. Consent: required via cookie banner.",
        ]
    )
    p(
        "We do not use advertising, social-media tracking or sale-of-data cookie categories on the Site. If we later "
        "add third-party embeds that set cookies (for example video or map embeds), we will update this policy and "
        "the banner categories before those features go live."
    )

    h("B. Google Analytics detail (expand current §3)", 3)
    p(
        "If you consent, we load Google Analytics 4 (GA4). The analytics script is not loaded until consent is given "
        "(or is removed if you later choose “Essential only”). GA4 may set first- and/or third-party cookies and "
        "collect technical and usage data such as approximate location derived from IP, device/browser type, pages "
        "viewed and timestamps. We configure GA4 for aggregated site statistics; we do not use it to show you ads."
    )
    p(
        "[COUNSEL/IT TO CONFIRM AND THEN PUBLISH:] We [enable / will enable] Google consent mode and [disable / keep "
        "disabled] advertising personalization signals for this property. IP anonymization / data-retention settings "
        "are configured in the GA4 admin interface as follows: [DESCRIBE ACTUAL SETTINGS]."
    )
    p(
        "Google’s role and transfers are described in the Privacy Notice section on recipients and international "
        "transfers. Google privacy information: https://policies.google.com/privacy."
    )

    h("C. How we obtain, withdraw and refresh consent (expand current §5)", 3)
    p(
        "On your first visit, the cookie banner offers clear choices, including accepting analytics, essential-only "
        "(rejecting analytics), and saving custom preferences. Essential cookies do not require consent. Analytics "
        "cookies run only with your consent."
    )
    p(
        "You can change or withdraw analytics consent at any time using the “Open cookie settings” control on this "
        "page (and equivalent controls where shown on the Site). Withdrawal does not make earlier processing unlawful."
    )
    p(
        "We may ask you to confirm cookie choices again if you clear site data, use another browser/device, or if we "
        "materially change cookie categories. [Optional policy: We re-prompt at least every [12 months] — CONFIRM.]"
    )

    h("11.5 Cookie Policy — Azerbaijani draft additions / replacements", 2)

    h("A. Kuki cədvəlinin genişləndirilməsi", 3)
    p("Təklif olunan sətirlər:")
    bullets(
        [
            "Zəruri — Məqsəd: əsas funksiyalar, təhlükəsizlik, dil seçimi, kuki seçiminin yadda saxlanması. Nümunələr: daab-cookie-consent (localStorage); daab-lang. Tərəf: birinci tərəf (WAAS). Müddət: silinənə / dəyişdirilənə qədər. Razılıq: tələb olunmur.",
            "Analitika (istəyə bağlı) — Məqsəd: Google Analytics 4 ilə ümumi statistika (ölçmə ID: G-K778896NSH). Nümunələr: _ga; _ga_<konteyner-id> və Google tərəfindən qoyulan əlaqəli GA4 teq/kukiləri. Tərəf: üçüncü tərəf (Google LLC). Tipik müddət: _ga üçün ən çoxu [2 il] / Google və GA4 parametrlərimizə uyğun. Razılıq: kuki pəncərəsi ilə tələb olunur.",
        ]
    )
    p(
        "Saytda reklam, sosial media izləmə və ya məlumat satışı kuki kateqoriyalarından istifadə etmirik. Sonradan "
        "kuki qoyan üçüncü tərəf məzmunu (məsələn, video və ya xəritə) əlavə olunarsa, bu siyasət və pəncərə "
        "kateqoriyaları həmin funksiyalar işə düşməzdən əvvəl yenilənəcək."
    )

    h("B. Google Analytics ətraflı (§3-ün genişləndirilməsi)", 3)
    p(
        "Razılıq verdikdə Google Analytics 4 (GA4) yüklənir. Analitika skripti razılıq verilməyənə qədər "
        "yüklənmir (sonradan “Yalnız zəruri” seçsəniz, söndürülür). GA4 birinci və/və ya üçüncü tərəf kukiləri "
        "qoya və IP-dən əldə olunan təqribi yer, cihaz/brauzer növü, baxılan səhifələr və vaxt damğası kimi "
        "texniki və istifadə məlumatlarını toplaya bilər. GA4-ü ümumi statistika üçün konfiqurasiya edirik; "
        "reklam göstərmək üçün istifadə etmirik."
    )
    p(
        "[VƏKİL/İT TƏSDİQLƏYİB DƏRC ETMƏLİDİR:] Bu mülkiyyət üçün Google razılıq rejimini [aktivləşdiririk / "
        "aktivləşdirəcəyik] və reklam fərdiləşdirmə siqnallarını [söndürürük / sönülü saxlayırıq]. IP-nin "
        "anonimləşdirilməsi / məlumat saxlama parametrləri GA4 idarə panelində belədir: [FAKTİKİ PARAMETRLƏRİ YAZIN]."
    )
    p(
        "Google-un rolu və ötürmələr Məxfilik bildirişinin alıcılar və beynəlxalq ötürmə bölməsində izah olunur. "
        "Google məxfilik məlumatı: https://policies.google.com/privacy."
    )

    h("C. Razılığın alınması, geri götürülməsi və yenilənməsi (§5-in genişləndirilməsi)", 3)
    p(
        "İlk ziyarətdə kuki pəncərəsi aydın seçimlər təqdim edir — analitikanı qəbul etmək, yalnız zəruri "
        "(analitikanı rədd etmək) və fərdi seçimləri saxlamaq. Zəruri kukilər üçün razılıq tələb olunmur. "
        "Analitika kukiləri yalnız razılığınızla işləyir."
    )
    p(
        "Analitika razılığını istənilən vaxt bu səhifədəki “Kuki seçimlərini aç” düyməsi (və saytda göstərilən "
        "ekvivalent idarəetmələr) ilə dəyişə və ya geri götürə bilərsiniz. Geri götürmə əvvəlki qanuni işlənməni "
        "qanunsuz etmir."
    )
    p(
        "Sayt məlumatlarını təmizləsəniz, başqa brauzer/cihaz istifadə etsəniz və ya kuki kateqoriyalarını "
        "əhəmiyyətli dərəcədə dəyişsek, seçimlərinizi yenidən təsdiqləməyinizi xahiş edə bilərik. "
        "[İstəyə bağlı siyasət: Ən azı hər [12 ayda] bir yenidən soruşuruq — TƏSDİQLƏYİN.]"
    )

    h("11.6 Suggested membership form checkbox wording (AZ / EN)", 2)
    p(
        "Replace the current “I have read and agree to processing…” style wording with a notice-acknowledgment "
        "plus, if needed, a separate publication consent later in the process."
    )
    p("English (proposed):", bold=True)
    p(
        "I have read the Privacy Notice and understand that the data I submit in this application will be used to "
        "assess and respond to my membership application as described there."
    )
    p("Azerbaijani (proposed):", bold=True)
    p(
        "Məxfilik bildirişini oxudum və bu müraciətdə göndərdiyim məlumatların orada göstərildiyi kimi üzvlük "
        "müraciətimin qiymətləndirilməsi və cavablandırılması üçün istifadə olunacağını başa düşürəm."
    )
    p(
        "Optional separate consent (only if a public profile will be published at application stage — usually "
        "better collected later):"
    )
    p("EN: I consent to publication of my name, photograph and biography on the WAAS website if my application is accepted. I may withdraw this consent later by emailing info@daab-waas.com.")
    p("AZ: Müraciətim qəbul olunarsa, adımın, fotomun və bioqrafiyamın WAAS saytında dərc olunmasına razılıq verirəm. Bu razılığı sonradan info@daab-waas.com ünvanına yazaraq geri götürə bilərəm.")

    h("11.7 Implementation checklist after counsel approval", 2)
    numbered(
        [
            "Fill all [SQUARE BRACKET] placeholders with confirmed facts (hosting, email provider, GA retention, Art. 27 representative decision, transfer mechanism).",
            "Paste approved EN/AZ privacy sections into helpers/_build_legal_pages.py privacy bodies.",
            "Paste approved EN/AZ cookie sections into the cookies bodies (table + GA + consent refresh).",
            "Update az/application.html and en/application.html checkbox labels to the approved notice-acknowledgment text.",
            "Rebuild legal pages; sync Deployment; bump asset/i18n cache versions if needed.",
            "Verify cookie banner still blocks GA until consent; document GA admin settings in the association’s internal compliance folder.",
        ]
    )

    h("12. Recommended next steps (action plan)")
    p(
        "Follow these steps in order. Do not publish Section 11 draft wording as live legal text until counsel "
        "has signed it off and all [SQUARE BRACKET] facts are filled."
    )

    h("12.1 Counsel review", 2)
    p(
        "Send this document (especially Section 11) to Turkish counsel. Ask them to approve or rewrite the draft "
        "wording and fill every [SQUARE BRACKET] item (hosting, email provider, transfer mechanism, GA retention, "
        "Art. 27 representative decision, and similar)."
    )

    h("12.2 Confirm facts internally", 2)
    p("Gather and confirm:")
    bullets(
        [
            "Hosting provider — legal name and country",
            "Email / mailbox provider — legal name and country",
            "GA4 retention setting and consent-mode / ad-personalization configuration",
            "Whether WAAS will appoint an EU/UK representative (GDPR/UK GDPR Art. 27)",
            "Concrete retention periods the association is willing to commit to publicly",
        ]
    )

    h("12.3 Apply approved text to the site", 2)
    p("After counsel approval:")
    numbered(
        [
            "Put the approved privacy and cookie wording into helpers/_build_legal_pages.py",
            "Rebuild the legal pages",
            "Update the membership checkbox on az/application.html and en/application.html",
            "Sync Deployment/",
        ]
    )

    h("12.4 Technical alignment", 2)
    bullets(
        [
            "Verify Google Analytics loads only after consent",
            "Set and document GA4 retention; keep advertising personalization disabled unless counsel approves otherwise",
            "Keep processor / DPA paperwork with Google, the hosting provider, and the email provider",
        ]
    )

    h("12.5 Medium priority (after the above)", 2)
    bullets(
        [
            "Imprint identifiers (dernek registration number and related official details)",
            "Formal IP takedown procedure with response timeline",
            "Accessibility statement",
            "Donation data flows, if bank details or payment processors are used",
        ]
    )

    h("12.6 Internal compliance pack", 2)
    p(
        "Create and maintain internal records so the published policies match real practice:"
    )
    bullets(
        [
            "ROPA (record of processing activities)",
            "Breach / incident playbook",
            "Access rules for who may open membership emails and application data",
        ]
    )

    h("12.7 Optional follow-up with the website team", 2)
    p(
        "After the association fills the bracketed facts and counsel approves Section 11, the website team can "
        "apply the approved wording to the site builders and republish the legal pages."
    )

    h("13. Implications of deploying the site to production")
    p(
        "Deploying now is technically fine for a normal production launch. The main implications are legal and "
        "process-related, not that the site will fail to work. This section is for internal planning and board "
        "discussion. It is not legal clearance to go live — that remains for counsel."
    )

    h("13.1 What production changes", 2)
    p("Once the site is live on daab-waas.com:")
    bullets(
        [
            "Your Privacy Notice, Cookie Policy, Terms, and Legal notice become the public rules you are holding out to visitors and applicants",
            "The cookie banner and GA4 start affecting real users (consent choices and analytics data)",
            "The membership form starts collecting real personal data under the notice people see",
            "Footer and Legal menu links make those documents easy to find — good for transparency, and also means they may be scrutinized",
        ]
    )
    p(
        "Production does not create new technical blockers; it makes your current legal text and practices operative."
    )

    h("13.2 What is already in decent shape for go-live", 2)
    p("You already have the core public suite most association sites need:")
    bullets(
        [
            "Operator identity / imprint",
            "Privacy notice (KVKK / PDPL-framed)",
            "Cookie policy with analytics opt-in",
            "Terms of use",
            "Discoverability in the footer and Legal menu",
        ]
    )
    p(
        "For many non-commercial association sites, that is enough to launch, with the understanding that "
        "counsel-strengthening continues afterward."
    )

    h("13.3 Main implications / residual risks if deployed as-is", 2)
    p(
        "These are the practical consequences of going live before Section 11 drafts are approved and applied:"
    )
    numbered(
        [
            "Supervisory complaint path missing — If someone asks for access/erasure and is unhappy, the privacy page does not clearly point them to the KVKK Board. That is a gap, not usually a must-block-launch item, but it is one of the easiest fixes after counsel okays wording.",
            "Transfers / processors / retention are still vague — With GA4 live in production, analytics data is processed under relatively general wording. Risk rises with traffic volume and with EU users. Mitigate by keeping analytics consent-only (already implemented) and tightening the text soon.",
            "Membership checkbox wording — Going live means real applicants acknowledge processing under the current language. If counsel later prefers notice-acknowledgment (not bundled consent), change the form wording and keep a note of when the old wording was used.",
            "GDPR / EU posture not spelled out — If many scientists in the EU apply or are profiled, GDPR expectations may apply even though WAAS is a Turkish association. Launch is still possible; the implication is you may need clearer rights text and an Art. 27 decision sooner rather than later.",
            "Operational reality must match the pages — Once live, statements such as “we don’t sell data,” “analytics only with consent,” and retention promises should match what you actually do. Weak spots today are mostly documentation and detail, not missing pages.",
            "No counsel sign-off yet on the new drafts — The Word-document drafts in Section 11 are not live. Deploying the current site does not publish those drafts. Do not treat this Word file as the live policy until approved wording is implemented on the website.",
        ]
    )

    h("13.4 What is not blocked by the legal gaps", 2)
    p("You can still deploy:")
    bullets(
        [
            "Static pages, scientists catalogue, forums, flyers, and related content",
            "Cookie banner behaviour already built",
            "AZ/EN language switch and footer legal links",
        ]
    )
    p(
        "Deployment packaging that omits helpers/ and documents/ is correct — this compliance Word file stays "
        "internal and should not be uploaded to the public web root."
    )

    h("13.5 Practical recommendation", 2)
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = "Table Grid"
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Approach"
    hdr2[1].text = "Implication"
    for approach, implication in [
        (
            "Deploy now",
            "Reasonable if you accept current legal pages as an interim public baseline and continue counsel review.",
        ),
        (
            "Deploy now + quick follow-up",
            "Best balance: launch, then apply counsel-approved Section 11 privacy/cookie updates soon.",
        ),
        (
            "Wait for full counsel + all brackets filled",
            "Lowest legal residual risk; delays launch for items that are improvements, not missing core pages.",
        ),
    ]:
        row = table2.add_row().cells
        row[0].text = approach
        row[1].text = implication
    doc.add_paragraph()
    p(
        "Bottom line: Production deploy is not reckless given what is already published, but it locks in the "
        "current privacy/cookie/terms as your public position. The implication is urgency — finish counsel review "
        "and apply the highest-priority privacy/cookie upgrades after launch (or before, if you prefer lower "
        "residual risk)."
    )

    h("13.6 Board checklist — safe to deploy / fix within 30 days", 2)
    p(
        "Use this as an internal go-live discussion aid. “Safe to deploy” here means: core public legal pages exist, "
        "analytics is consent-gated, and residual gaps are tracked for prompt follow-up. It is not a lawyer’s "
        "certificate of compliance."
    )

    p("Safe to deploy now (already in place or acceptable interim baseline):", bold=True)
    bullets(
        [
            "☐ Privacy Notice, Cookie Policy, Terms of Use, and Legal notice are published in AZ and EN",
            "☐ Legal links are visible in the site footer and under About → Legal",
            "☐ Cookie banner offers Essential only / analytics opt-in; GA does not load without consent",
            "☐ Membership form includes a privacy acknowledgment linking to the Privacy Notice",
            "☐ Association contact email and postal address appear on legal pages",
            "☐ Deployment package excludes helpers/ and documents/ (internal files not published)",
            "☐ Someone in the association is assigned to monitor info@daab-waas.com for privacy / rights requests",
        ]
    )

    p("Fix within 30 days after go-live (high priority follow-up):", bold=True)
    bullets(
        [
            "☐ Counsel reviews Section 11 of this document and returns approved wording",
            "☐ Fill all [SQUARE BRACKET] facts (hosting, email provider, transfer mechanism, GA retention, Art. 27 decision)",
            "☐ Publish approved Privacy Notice updates: complaint path to KVKK Board; processors/transfers; retention periods; EU/UK posture",
            "☐ Publish approved Cookie Policy updates: GA cookie names/durations; withdraw/refresh consent language",
            "☐ Align membership checkbox wording with counsel-approved notice-acknowledgment text",
            "☐ Document GA4 admin settings (retention, no ads personalization, consent behaviour) in an internal note",
            "☐ Confirm who may access membership applications / CVs and how long they are kept in practice",
        ]
    )

    p("Schedule within 90 days (medium priority):", bold=True)
    bullets(
        [
            "☐ Add dernek registration / imprint identifiers if counsel recommends publishing them",
            "☐ Add formal IP / copyright takedown procedure with response timeline",
            "☐ Decide accessibility statement scope and publish if adopted",
            "☐ Review donation/sponsorship data flows if payment or bank details are collected",
            "☐ Create internal ROPA and a simple breach/incident playbook",
        ]
    )

    p("Do not do at go-live:", bold=True)
    bullets(
        [
            "☐ Do not upload this Word file or other documents/Legal materials to the public website",
            "☐ Do not enable analytics without consent, advertising cookies, or new third-party embeds without updating the Cookie Policy first",
            "☐ Do not treat unapproved Section 11 drafts as the live legal text",
        ]
    )

    end = doc.add_paragraph()
    end.add_run(
        "End of document. After counsel approval, implement through helpers/_build_legal_pages.py so AZ/EN pages "
        "stay in sync. Regenerate this Word file anytime with: python helpers/_write_legal_compliance_review_docx.py"
    ).italic = True

    doc.save(OUT)
    print(f"Wrote {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
