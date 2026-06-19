#!/usr/bin/env python3
"""Generate documents/docx/DAAB-Comprehensive-Codebase-Review-2026-06-18.docx."""
from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from _paths import ROOT

OUT_DOCX = ROOT / "documents" / "docx" / "DAAB-Comprehensive-Codebase-Review-2026-06-18.docx"
REVIEW_DATE = "18 June 2026"

NAVY = RGBColor(0x09, 0x4D, 0x78)
BLUE = RGBColor(0x00, 0x69, 0xB4)
INK = RGBColor(0x1A, 0x2E, 0x3D)
MUTED = RGBColor(0x34, 0x5D, 0x76)
HEADER_FILL = "094D78"
ALT_FILL = "F4F7FA"

PRIORITY_COLORS = {
    "High": RGBColor(0xC9, 0x5A, 0x00),
    "Medium": RGBColor(0x8A, 0x6D, 0x00),
    "Low": RGBColor(0x2E, 0x6B, 0x3A),
    "Info": RGBColor(0x34, 0x5D, 0x76),
}


@dataclass(frozen=True)
class Finding:
    issue_id: str
    category: str
    where: str
    problem: str
    why_matters: str
    correction: str
    priority: str


FINDINGS: list[Finding] = [
    # --- SEO & metadata ---
    Finding(
        "SEO-01",
        "SEO & metadata",
        "All 60 indexable az/ and en/ pages; <!-- daab-seo --> blocks",
        "No og:image or twitter:image meta tags anywhere on the site.",
        "Social and messaging apps show text-only link previews without a visual card.",
        "Add a shared 1200×630 image (or daab-logo.png) via helpers/_inject_seo_head.py.",
        "High",
    ),
    Finding(
        "SEO-02",
        "SEO & metadata",
        "index.html (language gateway at site root)",
        "Gateway page JS-redirects to az/ or en/ but lacks full hreflang and Open Graph blocks.",
        "Crawlers and share bots hitting / receive thin metadata before redirect.",
        "Add hreflang + OG pointing to language homes, or server 302 with noindex on gateway.",
        "Medium",
    ),
    Finding(
        "SEO-03",
        "SEO & metadata",
        "az/membership.html; en/membership.html",
        "Legacy noindex redirect stubs to membership_value.html; no h1; partial meta.",
        "Old bookmarks may land on thin redirect pages (low SEO risk by design).",
        "Accept as-is, or use server 301 to membership_value.html and remove stubs.",
        "Low",
    ),
    # --- Links & assets ---
    Finding(
        "LINK-01",
        "Links & assets",
        "az/en membership_flyer.html; az/en sponsors_flyer.html",
        "QR code images loaded from api.qrserver.com at runtime.",
        "Third-party dependency; broken QR if API is down; privacy/offline print concerns.",
        "Generate static QR PNGs under images/qr/ and reference locally.",
        "Medium",
    ),
    Finding(
        "LINK-02",
        "Links & assets",
        "Sitewide footers (e.g. az/index.html contact block)",
        "Phone +90 555 147 46 74 is plain text; no tel: hyperlink.",
        "Mobile users cannot tap-to-call from footer contact.",
        'Wrap in <a href="tel:+905551474674">.</a>',
        "Medium",
    ),
    Finding(
        "LINK-03",
        "Links & assets",
        "Site web root; all pages use images/daab-logo.png as favicon",
        "No favicon.ico at root; browsers request /favicon.ico and may get 404.",
        "Minor UX noise and extra failed requests in server logs.",
        "Add favicon.ico at web root or configure host rewrite to images/daab-logo.png.",
        "Low",
    ),
    # --- Performance ---
    Finding(
        "PERF-01",
        "Performance",
        "az/en forum/2024/presentations.html (~145 KB); sessions_organization.html (~103 KB); "
        "rector_speeches.html (~76 KB); az/en activities.html (~75 KB)",
        "Large archival content fully inlined in HTML.",
        "Slower download/parse on mobile; harder to maintain bilingual copies.",
        "Optional: extract bodies to JSON + client render (scientist profiles pattern).",
        "Medium",
    ),
    Finding(
        "PERF-02",
        "Performance",
        "i18n/scientists-profiles.json (~316 KB); az/en scientists/profiles.html; "
        "scientists-list-preview.js on list page",
        "All 83 scientist profiles fetched and parsed in one JSON request.",
        "Noticeable stall on mid-tier phones when opening profiles or list previews.",
        "Chunk JSON by letter/country or virtualize card grid; share one cached fetch.",
        "Medium",
    ),
    Finding(
        "PERF-03",
        "Performance",
        "images/ (~1,812 files, ~91 MB); gallery and board photos 190–250 KB each",
        "Several images exceed typical web weight targets.",
        "Slower LCP on gallery, activities, and board photo pages.",
        "Batch resize/compress; prefer WebP where supported; keep daab-perf.js lazy loading.",
        "Medium",
    ),
    Finding(
        "PERF-04",
        "Performance",
        "i18n/search-index.json (~380 KB); js/daab-search.js lazy load",
        "Large single search index blob.",
        "Acceptable today (loads on search open); may grow with more content.",
        "Split by locale or prefix when index grows further.",
        "Low",
    ),
    Finding(
        "PERF-05",
        "Performance",
        "az/en scientists/profiles.html — 20 script tags, 16 stylesheets",
        "Heavy shared chrome stack per page visit.",
        "Parse and connection overhead on mobile despite defer.",
        "Audit required scripts; long-term optional chrome bundle.",
        "Low",
    ),
    # --- CSS architecture ---
    Finding(
        "CSS-01",
        "CSS architecture",
        "css/daab-forum-content.css ~670–677; az/en forum/2024/anas_leadership_speeches.html",
        "Comma selector lists bare html[data-daab-page-id=forum-anas-leadership-speeches] "
        "without .speech-body-lead descendant.",
        "Typography/layout rules intended for .speech-body-lead apply to entire <html> element.",
        "Fix selector to html[...] .speech-body-lead only; audit minified comma-list generator.",
        "High",
    ),
    Finding(
        "CSS-02",
        "CSS architecture",
        "daab-activities-page.css (642 !important); daab-common.css (428); daab-nav-mega.css (197); "
        "daab-charter-page.css (192)",
        "High !important count; sticky nav split between daab-common and daab-sticky-chrome.",
        "Unpredictable cascade; small CSS edits can break unrelated pages.",
        "Pick one chrome positioning model; reduce !important incrementally.",
        "High",
    ),
    Finding(
        "CSS-03",
        "CSS architecture",
        "daab-activities-layout.css; daab-charter-page.css; daab-nav-mega.css; daab-forum-content.css",
        "Near-identical sticky-sidebar / TOC / widget stacks duplicated across four files.",
        "Forum, activities, and charter pages drift when only one file is patched.",
        "Extract shared daab-sticky-sidebar.css or extend daab-sidebar-widget.css.",
        "Medium",
    ),
    Finding(
        "CSS-04",
        "CSS architecture",
        "daab-tokens.css --z-breadcrumbs:100 vs daab-sticky-chrome.css --z-breadcrumbs:9998",
        "Design token overridden at runtime in another global sheet.",
        "Breadcrumb and stacking context bugs are hard to diagnose.",
        "Single z-index scale in daab-tokens.css; remove override.",
        "Medium",
    ),
    Finding(
        "CSS-05",
        "CSS architecture",
        "22+ distinct @media widths; worst in daab-activities-page.css, daab-hub-cards.css",
        "Documented breakpoints 1060px/1180px not consistently used.",
        "Inconsistent tablet behaviour between similar page types.",
        "Normalize breakpoints to design tokens; shared gallery-collapse rule.",
        "Medium",
    ),
    Finding(
        "CSS-06",
        "CSS architecture",
        "az/en membership_flyer.html; az/en sponsors_flyer.html (~line 35)",
        "Redundant <link> to daab-tokens.css already inlined in daab-common.css.",
        "Extra HTTP request and duplicate token load on flyer pages.",
        "Remove redundant daab-tokens.css link.",
        "Low",
    ),
    Finding(
        "CSS-07",
        "CSS architecture",
        "css/daab-membership-application.css",
        "~70 hardcoded hex colours vs ~47 CSS custom properties.",
        "Form styling drifts from daab-tokens.css design system.",
        "Map colours and spacing to shared tokens.",
        "Low",
    ),
    # --- JavaScript ---
    Finding(
        "JS-01",
        "JavaScript",
        "az/en activities.html — daab-lang-position.js + daab-sidebar-timeline.js",
        "Two scroll-spy modules manage hash/active section with different offsets (32% vs 35%).",
        "Wrong TOC highlight and competing URL hash updates on long activities page.",
        "Use one owner per page — prefer daab-sidebar-timeline.js only.",
        "Medium",
    ),
    Finding(
        "JS-02",
        "JavaScript",
        "az/en charter.html ~510–513 inline IntersectionObserver + daab-lang-position.js",
        "Duplicate TOC spy logic on charter page.",
        "Same class of active-link bugs as activities.",
        "Replace inline observer with shared sidebar timeline module.",
        "Medium",
    ),
    Finding(
        "JS-03",
        "JavaScript",
        "js/scientists-list-catalog.js ~108–113",
        "window.DAAB_COLLATION.compare/sort used without localeCompare fallback.",
        "Runtime error if daab-collation.js load order changes.",
        "Add fallback like other catalog scripts.",
        "Medium",
    ),
    Finding(
        "JS-04",
        "JavaScript",
        "daab-sidebar-timeline.js; daab-forum-2026-toc.js; charter inline TOC",
        "~90% identical TOC/spy engines maintained separately.",
        "Fixes must be applied in multiple places; regression risk.",
        "Extract shared DAAB_SIDEBAR_SPY module.",
        "Medium",
    ),
    Finding(
        "JS-05",
        "JavaScript",
        "js/daab-search.js mountOverlay() ~303–304",
        "input.placeholder set without null guard after getElementById.",
        "Edge-case throw if overlay DOM injection fails.",
        "Add if (!input) return; after query.",
        "Low",
    ),
    Finding(
        "JS-06",
        "JavaScript",
        "daab-back-to-top.js; multiple resize handlers across modules",
        "Duplicate scroll/resize listeners (window + document on back-to-top).",
        "Redundant layout work and jank on long forum/activities pages.",
        "rAF-throttled shared DAAB_LAYOUT_SYNC bus; drop duplicate document listener.",
        "Low",
    ),
    # --- Accessibility ---
    Finding(
        "A11Y-01",
        "Accessibility",
        "Sitewide footers (e.g. az/index.html ✉ ☎ 🌐 ~275–287)",
        "Decorative emoji in footer contact rows not wrapped with aria-hidden.",
        "Screen readers announce emoji before link text (noisy UX).",
        "Wrap emoji in <span aria-hidden=\"true\"> via footer generator.",
        "Medium",
    ),
    Finding(
        "A11Y-02",
        "Accessibility",
        "az/en scientists/list.html — #filterCountry, #filterIxtilas, #filterDegree, "
        "#filterCins, #groupBy, #sortBy",
        "Filter <select> elements lack aria-label (search input is labelled).",
        "Screen reader users hear unlabeled filter controls.",
        "Add aria-label per filter in list page shell generator.",
        "Medium",
    ),
    Finding(
        "A11Y-03",
        "Accessibility",
        "js/daab-membership-application.js; az/en application.html science checkboxes",
        "Science-field errors not associated via fieldset/legend and aria-describedby.",
        "Form errors may not be announced clearly (WCAG 3.3.1 gap).",
        "Add fieldset/legend wrapper and role=alert on error region.",
        "Medium",
    ),
    Finding(
        "A11Y-04",
        "Accessibility",
        "Search overlay; photo lightbox (daab-photos-gallery.js); foundation gallery lightbox",
        "Background content not marked inert/aria-hidden while modals open.",
        "Screen reader browse mode can reach page behind open modal.",
        "Set inert on <main> while modal open (polyfill if needed).",
        "Medium",
    ),
    Finding(
        "A11Y-05",
        "Accessibility",
        "css/daab-search.css .search-empty; placeholder colour #8aa0b7",
        "Contrast ~4.0:1 on white (borderline WCAG AA for small text).",
        "Low-vision users may struggle with empty-state and placeholder text.",
        "Darken to ≥4.5:1 using verified --muted token.",
        "Low",
    ),
    Finding(
        "A11Y-06",
        "Accessibility",
        "js/daab-sponsors-page.js email validation",
        "field-error not linked with aria-describedby; no live region for errors.",
        "Validation errors not announced to assistive technology.",
        "Wire error id + aria-live=polite on error container.",
        "Low",
    ),
    # --- HTML & content ---
    Finding(
        "HTML-01",
        "HTML & content",
        "az/en forum/2024 presentations, sessions, speeches pages",
        "Large bilingual HTML duplicates for archival forum content.",
        "AZ/EN edits require two large files; content drift risk.",
        "Structured content source + generator, or accept as static archive.",
        "Medium",
    ),
    Finding(
        "HTML-02",
        "HTML & content",
        "az/en foundation.html image gallery",
        "Images use onclick=openLightbox(this); keyboard path partial.",
        "Not ideal semantic/interaction pattern for accessible gallery.",
        "Use <button> wrappers or delegated listeners on accessible controls.",
        "Low",
    ),
    # --- Forms & CTAs ---
    Finding(
        "FORM-01",
        "Forms & CTAs",
        "js/daab-membership-application.js step scroll ~624–642",
        "Step scroll uses --daab-nav-height + 120, ignores --daab-sticky-top-stack.",
        "Step headers can hide under sticky chrome + breadcrumbs on scroll.",
        "Use DAAB_LANG_POSITION.scrollToAnchor or shared stack helper.",
        "Medium",
    ),
    Finding(
        "FORM-02",
        "Forms & CTAs",
        "az/en application.html post-submit instructions",
        "CV and photo still require separate email to bilik.birlik@gmail.com.",
        "UX friction; applicants may think upload is complete after form submit.",
        "Phase 2: file upload backend (documented in membership storage strategy).",
        "Low",
    ),
    # --- Navigation & i18n ---
    Finding(
        "NAV-01",
        "Navigation & i18n",
        "i18n/nav.json membership group (3 items)",
        "No standalone charter/terms link under membership (charter is under About).",
        "Product/IA decision only — may be intentional.",
        "Confirm IA with stakeholders; add link only if required.",
        "Low",
    ),
    # --- Responsiveness ---
    Finding(
        "RESP-01",
        "Responsiveness",
        "daab-activities-page.css duplicate @media 768px; activities-layout + sidebar-widget",
        "Triplicated mobile widget collapse rules across activity sheets.",
        "Tablet layout inconsistencies on activities timeline pages.",
        "Consolidate mobile collapse into one stylesheet.",
        "Medium",
    ),
    # --- Configuration & docs ---
    Finding(
        "DOC-01",
        "Configuration & docs",
        "documents/DAAB-Codebase-Review-2026-06-followup.md; older review docx",
        "Internal docs reference removed prominent-figures module (~400 pages).",
        "Confusing for future maintainers; docs do not match 64-page live site.",
        "Archive or update internal documentation to current architecture.",
        "Low",
    ),
]

RESOLVED: list[tuple[str, str]] = [
    (
        "Broken internal links",
        "_validate_site.py — 4,042 references, 0 broken local paths.",
    ),
    (
        "Scientist profile HTML bloat",
        "Client-render via scientists-profiles-render.js + i18n/scientists-profiles.json (~11 KB shell).",
    ),
    (
        "Prominent-figures orphan module",
        "Removed from repository and Deployment package (June 2026).",
    ),
    (
        "Scientist photos/QR on production",
        "Deployment rebuild with --include-images; assetRoot() fix in render script.",
    ),
    (
        "Forum inline TOC pushState drift",
        "Forum 2024 pages migrated to daab-sidebar-timeline.js.",
    ),
    (
        "Photo lightbox focus trap",
        "daab-photos-gallery.js Tab trap at lines ~573–585.",
    ),
    (
        "Canonical / hreflang coverage",
        "60/60 indexable pages have canonical and hreflang (az, en, x-default).",
    ),
    (
        "Bilingual route parity",
        "_validate_bilingual.py — 30 AZ/EN pairs aligned with sitemap.xml.",
    ),
    (
        "All images have alt text",
        "Automated scan: 0 <img> without alt on deploy HTML.",
    ),
    (
        "AZ foundation page title",
        "Localised title DAAB — Birliyin təsisi (no longer English-only).",
    ),
]

BASELINE: list[str] = [
    "helpers/_validate_site.py — OK (64 HTML pages, 4,042 refs)",
    "helpers/_deploy_preflight.py — OK",
    "helpers/_validate_bilingual.py — OK (30 page pairs)",
    "helpers/_audit_repo_health.py — OK (no ?v= drift)",
    "helpers/_validate_cv_cards.py — OK (83 profiles, client-render mode)",
    "helpers/_check_name_order.py — OK (83/83 catalog parity)",
    "Artifact consistency audit — 0 errors (nav/i18n)",
]

CATEGORY_ORDER = [
    "SEO & metadata",
    "Links & assets",
    "Performance",
    "CSS architecture",
    "JavaScript",
    "Accessibility",
    "HTML & content",
    "Forms & CTAs",
    "Navigation & i18n",
    "Responsiveness",
    "Configuration & docs",
]

PHASES = [
    (
        "Phase 1 — High impact",
        "SEO-01 og:image; CSS-01 ANAS selector fix; CSS-02 chrome/!important strategy",
        "1–2 days",
    ),
    (
        "Phase 2 — UX & accessibility",
        "A11Y-01 footer emoji; A11Y-02 filter labels; LINK-02 tel: links; A11Y-03 application fieldset",
        "1 day",
    ),
    (
        "Phase 3 — Performance",
        "PERF-03 image compression; LINK-01 static flyer QR; PERF-02 scientists JSON chunking",
        "1–2 days",
    ),
    (
        "Phase 4 — JS consolidation",
        "JS-01/02 single scroll-spy; JS-03 collation fallback; JS-04 sidebar spy module",
        "2–3 days",
    ),
    (
        "Phase 5 — CSS debt",
        "CSS-03 sticky-sidebar extraction; CSS-04 z-index tokens; CSS-05 breakpoints",
        "Ongoing",
    ),
    (
        "Phase 6 — Polish",
        "LINK-03 favicon.ico; SEO-02 gateway; DOC-01 docs refresh; CSS-06 flyer token links",
        "Half day",
    ),
]


def shade_cell(cell, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: int = 8,
    color: RGBColor | None = None,
) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run2._r.append(instr)
    run3 = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)
    run4 = paragraph.add_run("1")
    run5 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)


def priority_summary() -> dict[str, int]:
    counts: dict[str, int] = {}
    for f in FINDINGS:
        counts[f.priority] = counts.get(f.priority, 0) + 1
    return counts


def add_finding_table(doc: Document, findings: list[Finding]) -> None:
    headers = ["ID", "Where", "Problem", "Why it matters", "Correction", "Priority"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[i],
            h,
            bold=True,
            size=7,
            color=RGBColor(0xFF, 0xFF, 0xFF),
        )
        shade_cell(table.rows[0].cells[i], HEADER_FILL)

    for idx, item in enumerate(findings):
        row = table.add_row().cells
        if idx % 2 == 1:
            for cell in row:
                shade_cell(cell, ALT_FILL)
        set_cell_text(row[0], item.issue_id, bold=True, size=7)
        set_cell_text(row[1], item.where, size=7)
        set_cell_text(row[2], item.problem, size=7)
        set_cell_text(row[3], item.why_matters, size=7)
        set_cell_text(row[4], item.correction, size=7)
        set_cell_text(
            row[5],
            item.priority,
            bold=True,
            size=7,
            color=PRIORITY_COLORS.get(item.priority, INK),
        )

    widths = [Cm(1.2), Cm(2.8), Cm(3.0), Cm(2.8), Cm(3.2), Cm(1.2)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w


def build_document() -> None:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ""
    hr = hp.add_run("DAAB / WAAS — Comprehensive Codebase Review")
    hr.font.size = Pt(9)
    hr.font.color.rgb = MUTED

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    fr = fp.add_run(f"Review date: {REVIEW_DATE}   |   Page ")
    fr.font.size = Pt(9)
    fr.font.color.rgb = MUTED
    add_page_number(fp)

    title = doc.add_heading("DAAB / WAAS Website", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_heading("Comprehensive Codebase Review — Cleanup Report", level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Review date: {REVIEW_DATE}\n"
        f"Scope: 64 deploy HTML pages, css/, js/, i18n/, images/ (~91 MB)\n"
        f"Repository: DAAB-WAAS static bilingual site (az/, en/)\n"
        f"Document generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        f"Open findings: {len(FINDINGS)}"
    )
    mr.font.color.rgb = MUTED
    mr.font.size = Pt(10)

    doc.add_page_break()

    # Executive summary
    doc.add_heading("1. Executive summary", level=1)
    counts = priority_summary()
    doc.add_paragraph(
        "This report documents findings from a comprehensive review of the DAAB/WAAS "
        "bilingual static website. All core automated validators pass: no broken local "
        "links, bilingual routes aligned with sitemap, scientist catalog consistent, and "
        "nav/i18n artifacts in sync. Remaining work centres on SEO polish (social preview "
        "images), CSS architecture debt, accessibility refinements, performance tuning on "
        "heavy forum pages, and JavaScript consolidation (duplicate scroll-spy logic)."
    )
    doc.add_paragraph(
        "Priority breakdown: "
        + "; ".join(f"{k}: {counts[k]}" for k in ("High", "Medium", "Low") if k in counts)
    )

    doc.add_heading("1.1 Automated validation baseline (passing)", level=2)
    for line in BASELINE:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("1.2 Resolved since earlier reviews", level=2)
    for title_text, detail in RESOLVED:
        p = doc.add_paragraph(style="List Bullet")
        r1 = p.add_run(f"{title_text}: ")
        r1.bold = True
        p.add_run(detail)

    doc.add_page_break()

    # Recommended phases
    doc.add_heading("2. Recommended remediation phases", level=1)
    pt = doc.add_table(rows=1, cols=3)
    pt.style = "Table Grid"
    for i, h in enumerate(["Phase", "Scope", "Effort"]):
        set_cell_text(pt.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(pt.rows[0].cells[i], HEADER_FILL)
    for phase, scope, effort in PHASES:
        cells = pt.add_row().cells
        set_cell_text(cells[0], phase, bold=True, size=9)
        set_cell_text(cells[1], scope, size=9)
        set_cell_text(cells[2], effort, size=9)

    doc.add_page_break()

    # Findings by category
    doc.add_heading("3. Detailed findings by category", level=1)
    doc.add_paragraph(
        "Each finding lists: where it appears, what the problem is, why it matters, "
        "how to correct it, and priority (High / Medium / Low)."
    )
    for category in CATEGORY_ORDER:
        cat = [f for f in FINDINGS if f.category == category]
        if not cat:
            continue
        doc.add_heading(category, level=2)
        doc.add_paragraph(f"Findings in this category: {len(cat)}")
        add_finding_table(doc, cat)
        doc.add_paragraph()

    doc.add_page_break()

    # Site metrics appendix
    doc.add_heading("4. Appendix A — Site metrics snapshot", level=1)
    metrics = [
        ("Deploy HTML pages", "64 (az/ + en/ + gateway + 404)"),
        ("Total HTML size (az + en)", "~1.9 MB"),
        ("Largest HTML page", "az/forum/2024/presentations.html (~145 KB)"),
        ("Scientists profiles shell", "~11 KB HTML + ~316 KB JSON"),
        ("Search index", "~380 KB (lazy loaded)"),
        ("Images", "~1,812 files, ~91 MB"),
        ("CSS !important (top file)", "daab-activities-page.css — 642 occurrences"),
        ("og:image coverage", "0 / 60 indexable pages"),
        ("hreflang coverage", "60 / 60 indexable pages"),
        ("tel: links in footers", "0"),
        ("External QR dependency", "4 flyer pages (api.qrserver.com)"),
    ]
    mt = doc.add_table(rows=1, cols=2)
    mt.style = "Table Grid"
    set_cell_text(mt.rows[0].cells[0], "Metric", bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
    set_cell_text(mt.rows[0].cells[1], "Value", bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
    shade_cell(mt.rows[0].cells[0], HEADER_FILL)
    shade_cell(mt.rows[0].cells[1], HEADER_FILL)
    for name, value in metrics:
        cells = mt.add_row().cells
        set_cell_text(cells[0], name, bold=True, size=9)
        set_cell_text(cells[1], value, size=9)

    doc.add_heading("5. Appendix B — Pre-deploy checklist", level=1)
    for step in [
        "python helpers/_validate_site.py",
        "python helpers/_deploy_preflight.py",
        "python helpers/_validate_bilingual.py",
        "python helpers/_validate_cv_cards.py",
        "python helpers/_build_deployment_folder.py --include-images",
        "Local smoke test: gateway → AZ home → scientists profiles → forum page → application",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("6. Appendix C — Issue index", level=1)
    idx = doc.add_table(rows=1, cols=4)
    idx.style = "Table Grid"
    for i, h in enumerate(["ID", "Category", "Priority", "Problem (brief)"]):
        set_cell_text(idx.rows[0].cells[i], h, bold=True, size=8, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(idx.rows[0].cells[i], HEADER_FILL)
    for item in FINDINGS:
        cells = idx.add_row().cells
        set_cell_text(cells[0], item.issue_id, size=8)
        set_cell_text(cells[1], item.category, size=8)
        set_cell_text(
            cells[2],
            item.priority,
            size=8,
            color=PRIORITY_COLORS.get(item.priority, INK),
        )
        brief = item.problem[:100] + ("…" if len(item.problem) > 100 else "")
        set_cell_text(cells[3], brief, size=8)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))


def main() -> int:
    build_document()
    print(f"Wrote {OUT_DOCX.relative_to(ROOT)}")
    print(f"  Findings: {len(FINDINGS)}")
    print(f"  Resolved items documented: {len(RESOLVED)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
