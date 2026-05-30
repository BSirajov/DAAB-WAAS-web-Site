#!/usr/bin/env python3
"""Build forum speech pages from forum_2024 rectors docx.

- az|en/forum/2024/rector_speeches.html — university rectors (excludes ANAS leadership)
- az|en/forum/2024/anas_leadership_speeches.html — İsa Həbibbəyli & Rasim Əliquliyev
"""
from __future__ import annotations

import html
import json
import re
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

PANEL_SUMMARIES_PATH = ROOT / "i18n" / "page-panel-summaries.json"

DOCX_AZ = ROOT / "forum_2024" / "AZƏRBAYCAN UNİVERSİTETLƏRİNİN REKTORLARININ NİTQLƏRİ.docx"
DOCX_EN = ROOT / "forum_2024" / "Azerbaijani_University_Rectors_Speeches_EN.docx"
ASSET = "../../../"
SIDEBAR_SCRIPT = f'<script src="{ASSET}js/daab-sidebar-timeline.js?v=2" defer></script>'

ANAS_SECTION_IDS = frozenset({"isa-hebibbeyli", "rasim-eliquliyev"})
SLUG_OVERRIDES = {
    "isa-habibbayli": "isa-hebibbeyli",
    "rasim-aliguliyev": "rasim-eliquliyev",
}

from _embed_static_nav import forum_nav_strip  # noqa: E402
from _site_wide_cleanup import SCRIPT_VERSIONS, STYLE_VERSIONS  # noqa: E402
from forum_en_common import FORUM_FOOTER_EN  # noqa: E402

AZ_SLUG = str.maketrans(
    {
        "Ə": "e",
        "ə": "e",
        "İ": "i",
        "I": "i",
        "Ö": "o",
        "ö": "o",
        "Ü": "u",
        "ü": "u",
        "Ş": "s",
        "ş": "s",
        "Ç": "c",
        "ç": "c",
        "Ğ": "g",
        "ğ": "g",
    }
)

SIGNOFF_RE = re.compile(
    r"^(Təşəkkür edirəm!|Hörmətlə,?$|Gələn forumda görüşənədək!|"
    r"Thank you\.?|Until we meet at the next forum!)$",
    re.I,
)

PS_SKIP_RE = re.compile(r"^P\.S\.", re.I)

ROADMAP_STOP_RE = re.compile(
    r"(YOL XƏRİTƏSİ|ROADMAP FOR THE DEVELOPMENT)",
    re.I,
)

# Section headings that must not become speech body text.
BODY_SKIP_RE = re.compile(
    r"^(DİASPOR HƏRƏKATINDA MÜHÜM İSTİQAMƏT|"
    r"AZƏRBAYCANLI ALİMLƏRİN QLOBAL BİRLİYİ|"
    r"AN IMPORTANT DIRECTION IN THE DIASPORA MOVEMENT|"
    r"THE GLOBAL UNITY OF AZERBAIJANI SCIENTISTS)",
    re.I,
)

EN_SKIP_H1_RE = re.compile(
    r"^(SPEECHES BY RECTORS|TABLE OF CONTENTS|ENGLISH TRANSLATION|"
    r"TRANSLATED FROM|AN IMPORTANT DIRECTION|THE GLOBAL UNITY)",
    re.I,
)

EN_ANAS_NAME_RE = re.compile(r"^(Isa HABIBBAYLI|Rasim ALIGULIYEV)$", re.I)

PAGE_SPECS = {
    "rector": {
        "out_az": ROOT / "az" / "forum" / "2024" / "rector_speeches.html",
        "out_en": ROOT / "en" / "forum" / "2024" / "rector_speeches.html",
        "page_id": "forum-rector-speeches",
        "filter": lambda sid: sid not in ANAS_SECTION_IDS,
        "toc_id": "rectorTOC",
        "copy": {
            "az": {
                "title": "Rektorlar — DAAB",
                "description": (
                    "Forum 2024 — Azərbaycan universitet rektorlarının nitqləri."
                ),
                "breadcrumb": "Rektorlar",
                "hero_h1": "Rektorlar",
                "hero_subtitle": "Azərbaycan universitet rektorlarının Forum 2024-dəki nitqləri",
                "sidebar_label": "Rektorlar",
                "sidebar_toggle": "Rektorlar siyahısını aç",
            },
            "en": {
                "title": "Rectors — WAAS",
                "description": (
                    "Forum 2024 — speeches by rectors of Azerbaijani universities."
                ),
                "breadcrumb": "Rectors",
                "hero_h1": "Rectors",
                "hero_subtitle": "Speeches by rectors of Azerbaijani universities at Forum 2024",
                "sidebar_label": "Rectors",
                "sidebar_toggle": "Open rectors list",
            },
        },
    },
    "anas": {
        "out_az": ROOT / "az" / "forum" / "2024" / "anas_leadership_speeches.html",
        "out_en": ROOT / "en" / "forum" / "2024" / "anas_leadership_speeches.html",
        "page_id": "forum-anas-leadership-speeches",
        "filter": lambda sid: sid in ANAS_SECTION_IDS,
        "toc_id": "anasTOC",
        "copy": {
            "az": {
                "title": "AMEA rəhbərliyi — DAAB",
                "description": (
                    "Forum 2024 — Azərbaycan Milli Elmlər Akademiyasının rəhbərliyinin nitqləri."
                ),
                "breadcrumb": "AMEA rəhbərliyi",
                "hero_h1": "AMEA <span>rəhbərliyi</span>",
                "hero_subtitle": (
                    "Azərbaycan Milli Elmlər Akademiyasının rəhbərliyinin Forum 2024-dəki nitqləri"
                ),
                "sidebar_label": "AMEA rəhbərliyi",
                "sidebar_toggle": "AMEA rəhbərliyi siyahısını aç",
            },
            "en": {
                "title": "ANAS Leadership — WAAS",
                "description": (
                    "Forum 2024 — speeches by leaders of the Azerbaijan National Academy of Sciences."
                ),
                "breadcrumb": "ANAS Leadership",
                "hero_h1": "ANAS <span>Leadership</span>",
                "hero_subtitle": (
                    "Speeches by leaders of the Azerbaijan National Academy of Sciences at Forum 2024"
                ),
                "sidebar_label": "ANAS Leadership",
                "sidebar_toggle": "Open ANAS leadership list",
            },
        },
    },
}

FOOTER_AZ = {
    "footer_brand": "Dünya Azərbaycanlı Alimlər Birliyi",
    "footer_contact": "Əlaqə",
    "footer_address": "Ünvan",
    "footer_leadership": "Rəhbərlik",
    "footer_chair": "DAAB İdarə Heyətinin Sədri",
    "footer_chair_name": "Prof. Dr. Məsud Əfəndiyev",
    "footer_bottom": "© 2026 DAAB / WAAS — All Rights Reserved",
}

FOOTER_EN = {
    "footer_brand": "World Association of Azerbaijani Scientists",
    "footer_contact": "Contact",
    "footer_address": "Address",
    "footer_leadership": "Leadership",
    "footer_chair": "Chair of the WAAS Board of Directors",
    "footer_chair_name": "Prof. Dr. Mesud Afandiyev",
    "footer_bottom": "© 2026 WAAS — All Rights Reserved",
}


def esc(s: str) -> str:
    return html.escape(s, quote=True)


def _panel_summaries() -> dict:
    return json.loads(PANEL_SUMMARIES_PATH.read_text(encoding="utf-8")).get("pages") or {}


def speech_hero_panel_html(page_id: str, lang: str) -> str:
    entry = _panel_summaries().get(page_id) or {}
    text = entry.get(lang) or entry.get("en") or ""
    if isinstance(text, list):
        text = " ".join(str(part).strip() for part in text if str(part).strip())
    titles = entry.get("panelTitle") or {}
    arias = entry.get("panelAria") or {}
    title = titles.get(lang) or titles.get("en") or ""
    aria = arias.get(lang) or arias.get("en") or title
    return (
        f'<aside aria-label="{esc(aria)}" class="hero-panel">\n'
        f'<div class="panel-card">\n'
        f'<h2 class="panel-title">{esc(title)}</h2>\n'
        f'<div class="panel-copy">\n'
        f'<p class="panel-copy-lead">{esc(text)}</p>\n'
        f"</div>\n"
        f"</div>\n"
        f"</aside>"
    )


def norm(text: str) -> str:
    return text.replace("\r", "").replace("\xa0", " ").strip()


def name_slug(name: str) -> str:
    s = name.translate(AZ_SLUG).lower()
    slug = re.sub(r"[^a-z0-9]+", "-", s).strip("-")
    return SLUG_OVERRIDES.get(slug, slug)


def new_section(name: str) -> dict:
    return {
        "id": name_slug(name),
        "name": name,
        "role": "",
        "blocks": [],
    }


def is_list_paragraph(text: str, style: str) -> bool:
    if style.startswith("List"):
        return True
    return text.startswith("- ") or text.startswith("– ")


def is_quote_line(text: str) -> bool:
    return (
        text.startswith("Belə bir aforizm var:")
        or text.startswith("There is an aphorism")
        or text.startswith('"')
        or text.startswith("\u201c")
        or text.startswith("\u2018")
    )


def append_body(cur: dict, text: str, style: str) -> None:
    if PS_SKIP_RE.match(text):
        return
    if BODY_SKIP_RE.match(text):
        return
    if style.startswith("Heading 2") and cur.get("role"):
        cur["blocks"].append({"list": False, "text": text, "subhead": True})
        return
    if style.startswith("Heading 3") and cur.get("role"):
        cur["blocks"].append({"list": False, "text": text, "subhead": True})
        return
    list_text = text.lstrip("-– ").strip() if is_list_paragraph(text, style) else text
    cur["blocks"].append(
        {"list": is_list_paragraph(text, style), "text": list_text}
    )


def parse_speeches_az() -> list[dict]:
    doc = Document(str(DOCX_AZ))
    sections: list[dict] = []
    cur: dict | None = None

    for para in doc.paragraphs:
        text = norm(para.text)
        if not text:
            continue
        style = para.style.name if para.style else ""

        if style == "Heading 1" and ROADMAP_STOP_RE.search(text):
            break

        if style.startswith("Heading 2"):
            if cur:
                sections.append(cur)
            cur = new_section(text)
            continue

        if not cur:
            continue

        if style.startswith("Heading 3") and not cur["role"]:
            cur["role"] = text
            continue

        append_body(cur, text, style)

    if cur:
        sections.append(cur)
    return sections


def parse_speeches_en() -> list[dict]:
    doc = Document(str(DOCX_EN))
    sections: list[dict] = []
    cur: dict | None = None
    in_anas_zone = False

    for para in doc.paragraphs:
        text = norm(para.text)
        if not text:
            continue
        style = para.style.name if para.style else ""

        if style == "Heading 1" and ROADMAP_STOP_RE.search(text):
            break

        if style == "Heading 1":
            if EN_SKIP_H1_RE.match(text):
                in_anas_zone = "AN IMPORTANT DIRECTION" in text.upper()
                continue
            if cur:
                sections.append(cur)
            cur = new_section(text)
            in_anas_zone = False
            continue

        if style.startswith("Heading 2"):
            if EN_ANAS_NAME_RE.match(text) or (
                in_anas_zone and not cur
            ):
                if cur:
                    sections.append(cur)
                cur = new_section(text)
                continue
            if cur and not cur["role"]:
                cur["role"] = text
                continue
            if cur:
                append_body(cur, text, style)
            continue

        if not cur:
            continue

        append_body(cur, text, style)

    if cur:
        sections.append(cur)
    return sections


def toc_item_html(section: dict, page_key: str) -> str:
    sid = esc(section["id"])
    name = esc(section["name"])
    titles = (section.get("role") or "").strip()
    titles_html = (
        f'<span class="rector-toc-titles">{esc(titles)}</span>' if titles else ""
    )
    return (
        f'<li><a class="rector-toc-link" href="#{sid}">'
        f'<span class="rector-toc-name">{name}</span>{titles_html}</a></li>'
    )


def strip_duplicate_role_blocks(section: dict) -> None:
    """Drop a leading body paragraph that repeats the Heading 3 academic title."""
    role = (section.get("role") or "").strip()
    if not role or not section.get("blocks"):
        return
    first = section["blocks"][0]
    if first.get("list"):
        return
    if norm(first["text"]) == role:
        section["blocks"].pop(0)


def blocks_to_html(blocks: list[dict]) -> str:
    parts: list[str] = []
    list_buf: list[str] = []

    def flush_list() -> None:
        nonlocal list_buf
        if not list_buf:
            return
        items = "".join(f"<li>{esc(t)}</li>" for t in list_buf)
        parts.append(f'<ul class="content-list">{items}</ul>')
        list_buf = []

    for block in blocks:
        text = block["text"]
        if block["list"]:
            list_buf.append(text)
            continue
        flush_list()
        if block.get("subhead"):
            parts.append(f'<p class="card-text"><strong>{esc(text)}</strong></p>')
        elif SIGNOFF_RE.match(text):
            parts.append(f'<p class="card-signoff">{esc(text)}</p>')
        elif is_quote_line(text):
            parts.append(f'<p class="card-quote">{esc(text)}</p>')
        else:
            parts.append(f'<p class="card-text">{esc(text)}</p>')

    flush_list()
    return "\n".join(parts)


def speech_card(section: dict) -> str:
    role_html = ""
    if section["role"]:
        role_html = (
            f'<p class="card-subtitle" role="doc-subtitle">{esc(section["role"])}</p>'
        )
    return f"""
<article class="news-card" id="{esc(section["id"])}">
<div class="card-header">
<h2 class="card-title">{esc(section["name"])}</h2>
{role_html}
</div>
<div class="card-body">
{blocks_to_html(section["blocks"])}
</div>
</article>"""


def footer_html(lang: str) -> str:
    if lang == "en":
        return FORUM_FOOTER_EN
    c = FOOTER_AZ
    return f"""<footer class="footer-pro">
<div class="footer-inner">
<div class="footer-brand"><h3>{c["footer_brand"]}</h3></div>
<div class="footer-grid">
<div class="footer-col"><h4 class="footer-title">{c["footer_contact"]}</h4><div class="footer-item">✉ <a href="mailto:info@daab-waas.com">info@daab-waas.com</a></div><div class="footer-item">☎ <span>+90 555 147 46 74</span></div><div class="footer-item">🌐 <a href="https://daab-waas.com" rel="noopener noreferrer" target="_blank">daab-waas.com</a></div></div>
<div class="footer-col"><h4 class="footer-title">{c["footer_address"]}</h4><p class="footer-address">Feneryolu Mahallesi<br/>Gazi Muhtar Paşa Sokak No:44<br/>Kadıköy, İstanbul, Türkiyə</p></div>
<div class="footer-col"><h4 class="footer-title">{c["footer_leadership"]}</h4><p class="footer-leader"><strong>{c["footer_chair_name"]}</strong><br/>{c["footer_chair"]}<br/>Germany — James D. Murray Distinguished Professor</p></div>
</div>
</div>
<div class="footer-bottom">{c["footer_bottom"]}</div>
</footer>"""


def build_html(lang: str, sections: list[dict], spec: dict, page_key: str) -> str:
    c = {**spec["copy"][lang], **(FOOTER_EN if lang == "en" else FOOTER_AZ)}
    page_id = spec["page_id"]
    toc_id = spec["toc_id"]
    speech_photos_css = ""
    if page_id in ("forum-anas-leadership-speeches", "forum-rector-speeches"):
        speech_photos_css = (
            f'<link href="{ASSET}css/daab-speech-photos.css?v=2" rel="stylesheet"/>\n'
        )
    nav = forum_nav_strip(lang, active_nav_id="forum-2024")
    toc_items = "".join(toc_item_html(s, page_key) for s in sections)
    cards = "\n".join(speech_card(s) for s in sections)

    if lang == "az":
        crumbs = (
            '<a href="../../index.html">Ana səhifə</a><span aria-hidden="true">›</span>'
            '<a href="../../activities.html">Fəaliyyətimiz</a><span aria-hidden="true">›</span>'
            '<a href="index.html">Forum 2024</a><span aria-hidden="true">›</span>'
            f'<span class="forum-breadcrumbs-current" aria-current="page">{c["breadcrumb"]}</span>'
        )
        skip = "Məzmuna keç"
        crumb_aria = "Səhifə yolu"
    else:
        crumbs = (
            '<a href="../../index.html">Home</a><span aria-hidden="true">›</span>'
            '<a href="../../activities.html">Activities</a><span aria-hidden="true">›</span>'
            '<a href="index.html">Forum 2024</a><span aria-hidden="true">›</span>'
            f'<span class="forum-breadcrumbs-current" aria-current="page">{c["breadcrumb"]}</span>'
        )
        skip = "Skip to content"
        crumb_aria = "Breadcrumb"

    hub_panel = speech_hero_panel_html(page_id, lang)
    sv = SCRIPT_VERSIONS
    st = STYLE_VERSIONS
    return f"""<!DOCTYPE html>
<html lang="{lang}" data-daab-lang="{lang}" data-daab-asset-root="{ASSET}" data-daab-page-id="{page_id}" data-daab-nav-mount="1">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0, viewport-fit=cover"/>
<title>{esc(c["title"])}</title>
<meta name="description" content="{esc(c["description"])}"/>
<link rel="preconnect" href="https://fonts.googleapis.com"/>
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin/>
<link href="https://fonts.googleapis.com/css2?family=Inter:opsz,wght@14..32,400..900&family=Playfair+Display:wght@700;800&display=swap" rel="stylesheet"/>
<link href="{ASSET}css/daab-common.css?v={st["daab-common.css"]}" rel="stylesheet"/>
<link href="{ASSET}css/daab-mobile.css?v={st["daab-mobile.css"]}" rel="stylesheet"/>
<link href="{ASSET}css/daab-sticky-chrome.css?v={st["daab-sticky-chrome.css"]}" rel="stylesheet"/>
<link href="{ASSET}css/daab-search.css?v={st["daab-search.css"]}" rel="stylesheet"/>
<link href="{ASSET}css/daab-back-to-top.css?v={st["daab-back-to-top.css"]}" rel="stylesheet"/>
<link href="{ASSET}css/daab-lang.css?v={st["daab-lang.css"]}" rel="stylesheet"/>
<link href="{ASSET}css/daab-nav-mega.css?v={st["daab-nav-mega.css"]}" rel="stylesheet"/>
<link href="{ASSET}css/daab-forum-section-nav.css?v={st["daab-forum-section-nav.css"]}" rel="stylesheet"/>
<link href="{ASSET}css/daab-hero-summary.css?v={st["daab-hero-summary.css"]}" rel="stylesheet"/>
<link href="{ASSET}css/daab-sidebar-widget.css?v={st["daab-sidebar-widget.css"]}" rel="stylesheet"/>
<link href="{ASSET}css/daab-activities-layout.css?v={st["daab-activities-layout.css"]}" rel="stylesheet"/>
<link href="{ASSET}css/daab-forum-content.css?v={st["daab-forum-content.css"]}" rel="stylesheet"/>
{speech_photos_css}<script src="{ASSET}js/daab-mobile.js?v={sv["daab-mobile.js"]}" defer></script>
<script src="{ASSET}js/daab-sticky-chrome.js?v={sv["daab-sticky-chrome.js"]}" defer></script>
<script src="{ASSET}js/daab-back-to-top.js?v={sv["daab-back-to-top.js"]}" defer></script>
<script src="{ASSET}js/daab-i18n.js?v={sv["daab-i18n.js"]}" defer></script>
<script src="{ASSET}js/daab-lang-position.js?v={sv["daab-lang-position.js"]}" defer></script>
<script src="{ASSET}js/daab-design-tokens.js?v={sv["daab-design-tokens.js"]}" defer></script>
<script src="{ASSET}js/daab-nav.js?v={sv["daab-nav.js"]}" defer></script>
<script src="{ASSET}js/daab-primary-nav.js?v={sv["daab-primary-nav.js"]}" defer></script>
<script src="{ASSET}js/daab-section-nav.js?v={sv["daab-section-nav.js"]}" defer></script>
<script src="{ASSET}js/daab-shell.js?v={sv["daab-shell.js"]}" defer></script>
<script src="{ASSET}js/daab-page-subtitle.js?v={sv["daab-page-subtitle.js"]}" defer></script>
<script src="{ASSET}js/daab-search.js?v={sv["daab-search.js"]}" defer></script>
</head>
<body>
<a class="skip" href="#content">{skip}</a>
{nav}
<div class="breadcrumbs forum-breadcrumbs" role="navigation" aria-label="{crumb_aria}">
{crumbs}
</div>
<header class="page-hero">
<div class="hero-wrap shell">
<section class="hero-copy">
<h1 aria-describedby="page-hero-subtitle">{c["hero_h1"]}</h1>
<p class="page-hero-subtitle" id="page-hero-subtitle" role="doc-subtitle">{esc(c["hero_subtitle"])}</p>
</section>
{hub_panel}
</div>
</header>
<div class="content-wrap">
<aside class="sidebar">
<div class="sidebar-widget">
<div class="widget-head"><span>{c["sidebar_label"]}</span><button aria-controls="{toc_id}" aria-expanded="false" aria-label="{esc(c["sidebar_toggle"])}" class="events-menu-toggle" type="button"><span></span><span></span><span></span></button></div>
<div class="widget-body">
<ul class="timeline-list" id="{toc_id}">
{toc_items}
</ul>
</div>
</div>
</aside>
<main class="news-feed main" id="content">
{cards}
</main>
</div>
{footer_html(lang)}
{SIDEBAR_SCRIPT}
</body>
</html>
"""


def build() -> None:
    if not DOCX_AZ.is_file():
        raise SystemExit(f"Missing source: {DOCX_AZ}")
    if not DOCX_EN.is_file():
        raise SystemExit(f"Missing source: {DOCX_EN}")

    by_lang = {
        "az": parse_speeches_az(),
        "en": parse_speeches_en(),
    }
    for lang, sections in by_lang.items():
        if not sections:
            raise SystemExit(f"No speech sections parsed from {lang} docx")

    for key, spec in PAGE_SPECS.items():
        for lang in ("az", "en"):
            sections = [s for s in by_lang[lang] if spec["filter"](s["id"])]
            if not sections:
                raise SystemExit(f"No sections for page {key} ({lang})")
            for section in sections:
                strip_duplicate_role_blocks(section)

            out = spec["out_az"] if lang == "az" else spec["out_en"]
            out.parent.mkdir(parents=True, exist_ok=True)
            out.write_text(
                build_html(lang, sections, spec, key), encoding="utf-8", newline="\n"
            )
            print(f"wrote {out.relative_to(ROOT)} ({len(sections)} speeches)")
            if key in ("anas", "rector"):
                from _speech_photos_lib import refresh_page

                refresh_page(out, spec["toc_id"])


if __name__ == "__main__":
    build()
