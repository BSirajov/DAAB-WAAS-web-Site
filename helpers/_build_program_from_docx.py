"""Build az/forum/2024/program.html from forum_2024/Proqram.docx (Yeniliklər UI)."""
from __future__ import annotations

import html
import re
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

DOCX = ROOT / "forum_2024" / "Proqram.docx"
OUT = ROOT / "az" / "forum" / "2024" / "program.html"
ASSET = "../../../"

from _embed_static_nav import forum_nav_strip  # noqa: E402

DAY_RE = re.compile(r"^(\d{1,2})\s+SENTYABR\s+2024\s*$", re.I)
TIME_ROW_RE = re.compile(
    r"^(\d{1,2}:\d{2}(?:\s*[–—-]\s*\d{1,2}:\d{2})?)\s+(.+)$"
)
TIME_ONLY_RE = re.compile(r"^\d{1,2}:\d{2}(?:\s*[–—-]\s*\d{1,2}:\d{2})?\s*$")

SIDEBAR_SCRIPT = """
<script>
(function () {
  const links = Array.from(document.querySelectorAll('.timeline-list a[href^="#"]'));
  const ids = links.map(a => a.getAttribute('href').slice(1));
  const cards = ids.map(id => document.getElementById(id)).filter(Boolean);
  const eventsWidget = document.querySelector('.sidebar-widget');
  const eventsToggle = document.querySelector('.events-menu-toggle');
  const mobileQuery = window.matchMedia('(max-width: 1060px)');

  function activate(link) {
    links.forEach(a => a.classList.remove('tl-active'));
    if (link) link.classList.add('tl-active');
  }

  function closeEventsMenu() {
    if (!eventsWidget || !eventsToggle) return;
    eventsWidget.classList.remove('events-open');
    eventsToggle.setAttribute('aria-expanded', 'false');
  }

  function toggleEventsMenu() {
    if (!eventsWidget || !eventsToggle) return;
    const open = eventsWidget.classList.toggle('events-open');
    eventsToggle.setAttribute('aria-expanded', open ? 'true' : 'false');
  }

  function jumpToTarget(event) {
    const link = event.currentTarget;
    const id = link.getAttribute('href').slice(1);
    const target = document.getElementById(id);
    if (!target) return;
    event.preventDefault();
    activate(link);
    const Pos = window.DAAB_LANG_POSITION;
    if (Pos && Pos.scrollToAnchor) {
      Pos.scrollToAnchor(id, false);
    } else {
      target.scrollIntoView({ block: 'start', behavior: 'auto' });
    }
    history.pushState(null, '', link.getAttribute('href'));
    if (mobileQuery.matches) closeEventsMenu();
  }

  function onScroll() {
    const mid = window.scrollY + window.innerHeight * 0.35;
    let active = null;
    for (let i = cards.length - 1; i >= 0; i--) {
      if (cards[i] && cards[i].offsetTop <= mid) {
        active = i;
        break;
      }
    }
    activate(active !== null ? links[ids.indexOf(cards[active].id)] : null);
  }

  links.forEach(link => link.addEventListener('click', jumpToTarget));
  if (eventsToggle) {
    eventsToggle.addEventListener('click', event => {
      event.stopPropagation();
      toggleEventsMenu();
    });
  }
  document.addEventListener('click', event => {
    if (!mobileQuery.matches || !eventsWidget || !eventsWidget.classList.contains('events-open')) return;
    if (eventsWidget.contains(event.target)) return;
    closeEventsMenu();
  });
  document.addEventListener('keydown', event => {
    if (event.key === 'Escape') closeEventsMenu();
  });
  window.addEventListener('scroll', onScroll, { passive: true });
  onScroll();
})();
</script>"""


def esc(s: str) -> str:
    return html.escape(s, quote=True)


def norm(text: str) -> str:
    return text.replace("\r", "").replace("\xa0", " ").strip()


def table_to_html(table, *, four_col: bool = False) -> str:
    rows = table.rows
    if not rows:
        return ""
    if four_col:
        head = "<thead><tr><th>Vaxt</th><th>Məruzəçi</th><th>Ölkə</th><th>Mövzu</th></tr></thead>"
        body_rows = []
        for row in rows[1:]:
            cells = [norm(c.text) for c in row.cells]
            while len(cells) < 4:
                cells.append("")
            body_rows.append(
                "<tr>"
                f"<td>{esc(cells[0])}</td>"
                f"<td>{esc(cells[1])}</td>"
                f"<td>{esc(cells[2])}</td>"
                f"<td>{esc(cells[3])}</td>"
                "</tr>"
            )
    else:
        head = "<thead><tr><th>Vaxt</th><th>Tədbir</th></tr></thead>"
        body_rows = []
        for row in rows:
            cells = [norm(c.text) for c in row.cells]
            if len(cells) >= 2:
                t, d = cells[0], cells[1]
            else:
                t, d = cells[0], ""
            body_rows.append(f"<tr><td>{esc(t)}</td><td>{esc(d)}</td></tr>")
    tbody = "".join(body_rows)
    cls = "program-table program-table--wide" if four_col else "program-table"
    return (
        f'<div class="program-table-wrap"><table class="{cls}">'
        f"{head}<tbody>{tbody}</tbody></table></div>"
    )


def parse_days(doc: Document) -> list[dict]:
    days: list[dict] = []
    current: dict | None = None
    items: list[dict] = []

    def flush_day() -> None:
        nonlocal current, items
        if current:
            days.append({**current, "items": items})
        items = []

    for para in doc.paragraphs:
        text = norm(para.text)
        if not text:
            continue
        m_day = DAY_RE.match(text)
        if m_day:
            flush_day()
            n = m_day.group(1)
            titles = {
                "9": "9 sentyabr 2024 — Bakı",
                "10": "10 sentyabr 2024 — Strateji sessiyalar",
                "11": "11 sentyabr 2024 — Xankəndi və Şuşa",
            }
            current = {"id": f"gun-{n}", "title": titles.get(n, f"{n} sentyabr 2024")}
            continue
        if not current:
            continue

        style = str(para.style.name) if para.style else ""
        if "List" in style:
            items.append({"type": "li", "text": text})
            continue

        m_time = TIME_ROW_RE.match(text)
        if m_time:
            items.append(
                {
                    "type": "slot",
                    "time": m_time.group(1).strip(),
                    "text": m_time.group(2).strip(),
                }
            )
            continue
        if TIME_ONLY_RE.match(text):
            items.append({"type": "slot", "time": text.strip(), "text": ""})
            continue

        if len(text) < 100 and text == text.upper() and re.search(r"[A-ZƏÜÖĞŞÇİ]", text):
            items.append({"type": "subhead", "text": text})
        else:
            items.append({"type": "p", "text": text})

    flush_day()
    return days


def items_to_html(items: list[dict], extra_table: str = "") -> str:
    parts: list[str] = []
    slot_buf: list[tuple[str, str]] = []

    def flush_slots() -> None:
        if not slot_buf:
            return
        rows = "".join(
            f"<tr><td>{esc(t)}</td><td>{esc(d)}</td></tr>" for t, d in slot_buf
        )
        parts.append(
            '<div class="program-table-wrap"><table class="program-table">'
            "<thead><tr><th>Vaxt</th><th>Tədbir</th></tr></thead>"
            f"<tbody>{rows}</tbody></table></div>"
        )
        slot_buf.clear()

    insert_table = bool(extra_table)
    table_placed = False

    i = 0
    while i < len(items):
        item = items[i]
        if item["type"] == "slot":
            slot_buf.append((item["time"], item["text"]))
            if (
                insert_table
                and not table_placed
                and "08:40" in item["time"]
                and "10:00" in item["time"]
            ):
                flush_slots()
                parts.append(extra_table)
                table_placed = True
            i += 1
            continue

        flush_slots()
        if item["type"] == "li":
            lis = []
            while i < len(items) and items[i]["type"] == "li":
                lis.append(f"<li>{esc(items[i]['text'])}</li>")
                i += 1
            parts.append(f'<ul class="content-list">{"".join(lis)}</ul>')
            continue
        if item["type"] == "subhead":
            parts.append(f'<h3 class="program-subhead">{esc(item["text"])}</h3>')
        else:
            parts.append(f'<p class="card-text">{esc(item["text"])}</p>')
        i += 1

    flush_slots()
    if insert_table and not table_placed:
        parts.append(extra_table)
    return "\n".join(parts)


def day_card(day: dict, body_html: str) -> str:
    return f"""
<article class="news-card program-day-card" id="{esc(day["id"])}">
<div class="card-header">
<h2 class="card-title">{esc(day["title"])}</h2>
</div>
<div class="card-body program-day-body">
{body_html}
</div>
</article>"""


def build() -> None:
    doc = Document(str(DOCX))
    days = parse_days(doc)
    presentations_html = table_to_html(doc.tables[0], four_col=True) if doc.tables else ""
    karabakh_html = table_to_html(doc.tables[1]) if len(doc.tables) > 1 else ""

    cards: list[str] = []
    for day in days:
        extra = ""
        if day["id"] == "gun-10" and presentations_html:
            extra = presentations_html
        body = items_to_html(day["items"], extra_table=extra)
        if day["id"] == "gun-11" and karabakh_html:
            body = karabakh_html + body
        cards.append(day_card(day, body))

    toc_items = "".join(
        f'<li><a href="#{esc(d["id"])}">{esc(d["title"])}</a></li>' for d in days
    )
    nav = forum_nav_strip("az", active_nav_id="forum-2024")

    html_out = f"""<!DOCTYPE html>
<html lang="az" data-daab-lang="az" data-daab-asset-root="{ASSET}" data-daab-page-id="forum-program" data-daab-nav-mount="1">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0, viewport-fit=cover"/>
<title>Proqram — DAAB</title>
<meta name="description" content="Xaricdə Yaşayan Azərbaycanlı Alimlərin 2024-cü il Forumu — 9–11 sentyabr proqramı (Bakı, Xankəndi, Şuşa)."/>
<link rel="preconnect" href="https://fonts.googleapis.com"/>
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin/>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&family=Playfair+Display:wght@700;800&display=swap" rel="stylesheet"/>
<link href="{ASSET}css/daab-common.css?v=24" rel="stylesheet"/>
<link href="{ASSET}css/daab-mobile.css?v=5" rel="stylesheet"/>
<link href="{ASSET}css/daab-search.css?v=3" rel="stylesheet"/>
<link href="{ASSET}css/daab-back-to-top.css?v=1" rel="stylesheet"/>
<link href="{ASSET}css/daab-lang.css?v=10" rel="stylesheet"/>
<link href="{ASSET}css/daab-nav-mega.css?v=23" rel="stylesheet"/>
<link href="{ASSET}css/daab-forum-section-nav.css?v=1" rel="stylesheet"/>
<link href="{ASSET}css/daab-hero-summary.css?v=1" rel="stylesheet"/>
<link href="{ASSET}css/daab-sidebar-widget.css?v=3" rel="stylesheet"/>
<link href="{ASSET}css/daab-activities-layout.css?v=7" rel="stylesheet"/>
<link href="{ASSET}css/daab-forum-content.css?v=12" rel="stylesheet"/>
<script src="{ASSET}js/daab-mobile.js?v=1" defer></script>
<script src="{ASSET}js/daab-back-to-top.js?v=2" defer></script>
<script src="{ASSET}js/daab-i18n.js?v=12" defer></script>
<script src="{ASSET}js/daab-lang-position.js?v=7" defer></script>
<script src="{ASSET}js/daab-nav.js?v=9" defer></script>
<script src="{ASSET}js/daab-primary-nav.js?v=9" defer></script>
<script src="{ASSET}js/daab-section-nav.js?v=12" defer></script>
<script src="{ASSET}js/daab-shell.js?v=11" defer></script>
<script src="{ASSET}js/daab-search.js?v=4" defer></script>
</head>
<body>
<a class="skip" href="#content">Məzmuna keç</a>
{nav}
<div class="breadcrumbs forum-breadcrumbs" role="navigation" aria-label="Səhifə yolu">
<a href="../../index.html">Ana səhifə</a><span aria-hidden="true">›</span><a href="index.html">Forum 2024</a><span aria-hidden="true">›</span><span class="forum-breadcrumbs-current" aria-current="page">Proqram</span>
</div>
<header class="page-hero">
<div class="hero-wrap shell">
<section class="hero-copy">
<h1>Proqram</h1>
</section>
<aside aria-label="Proqram haqqında qısa məlumat" class="hero-panel">
<div class="panel-card">
<h2 class="panel-title">Proqram</h2>
<div class="panel-copy">9–11 sentyabr 2024-cü il ərzində Bakı, Xankəndi və Şuşa üzrə açılış iclası, plenar sessiyalar, məruzələr və yekun tədbirlərin tam cədvəli.</div>
</div>
</aside>
</div>
</header>
<div class="content-wrap">
<aside class="sidebar">
<div class="sidebar-widget">
<div class="widget-head"><span>📅 Proqram</span><button aria-controls="programTOC" aria-expanded="false" aria-label="Proqram menyusunu aç" class="events-menu-toggle" type="button"><span></span><span></span><span></span></button></div>
<div class="widget-body">
<ul class="timeline-list" id="programTOC">
{toc_items}
</ul>
</div>
</div>
</aside>
<main class="news-feed main" id="content">
{"".join(cards)}
</main>
</div>
<footer class="footer-pro">
<div class="footer-inner">
<div class="footer-brand"><h3>Dünya Azərbaycanlı Alimlər Birliyi</h3></div>
<div class="footer-grid">
<div class="footer-col"><h4 class="footer-title">Əlaqə</h4><div class="footer-item">✉ <a href="mailto:info@daab-waas.com">info@daab-waas.com</a></div><div class="footer-item">☎ <span>+90 555 147 46 74</span></div><div class="footer-item">🌐 <a href="https://daab-waas.com" rel="noopener noreferrer" target="_blank">daab-waas.com</a></div></div>
<div class="footer-col"><h4 class="footer-title">Ünvan</h4><p class="footer-address">Feneryolu Mahallesi<br/>Gazi Muhtar Paşa Sokak No:44<br/>Kadıköy, İstanbul, Türkiyə</p></div>
<div class="footer-col"><h4 class="footer-title">Rəhbərlik</h4><p class="footer-leader"><strong>Prof. Dr. Məsud Əfəndiyev</strong><br/>DAAB İdarə Heyətinin Sədri<br/>Germany — James D. Murray Distinguished Professor</p></div>
</div>
</div>
<div class="footer-bottom">© 2026 DAAB / WAAS — All Rights Reserved</div>
</footer>
{SIDEBAR_SCRIPT}
</body>
</html>
"""
    OUT.parent.mkdir(parents=True, exist_ok=True)
    OUT.write_text(html_out, encoding="utf-8", newline="\n")
    print(f"wrote {OUT.relative_to(ROOT)} ({len(days)} days)")


if __name__ == "__main__":
    build()
