#!/usr/bin/env python3
"""
Generate documents/docx/DAAB-Website-Codebase-Review-Report.docx
from the June 2026 full-site technical review findings.
"""
from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from _paths import ROOT

OUT_DOCX = ROOT / "documents" / "docx" / "DAAB-Website-Codebase-Review-Report.docx"

# Review identification timestamp (automated validation run window)
REVIEW_IDENTIFIED_AT = "2026-06-07 23:21:53"
REVIEW_PERIOD = "5–7 June 2026"

NAVY = RGBColor(0x09, 0x4D, 0x78)
BLUE = RGBColor(0x00, 0x69, 0xB4)
INK = RGBColor(0x1A, 0x2E, 0x3D)
MUTED = RGBColor(0x34, 0x5D, 0x76)
HEADER_FILL = "094D78"
ALT_FILL = "F4F7FA"

SEVERITY_COLORS = {
    "Critical": RGBColor(0xB4, 0x22, 0x22),
    "High": RGBColor(0xC9, 0x5A, 0x00),
    "Medium": RGBColor(0x8A, 0x6D, 0x00),
    "Low": RGBColor(0x2E, 0x6B, 0x3A),
    "Informational": RGBColor(0x34, 0x5D, 0x76),
}


@dataclass(frozen=True)
class Issue:
    issue_id: str
    category: str
    description: str
    affected: str
    severity: str
    recommended_fix: str
    expected_result: str


ISSUES: list[Issue] = [
    # --- Internal links & routing ---
    Issue(
        "LINK-02",
        "Internal links",
        "sitemap.xml does not list Treasury placeholder pages that exist in routes.json and are "
        "linked from navigation.",
        "sitemap.xml; az/industrial_revolutions.html; en/industrial_revolutions.html; "
        "az/major_scientific_inventions.html; en/major_scientific_inventions.html",
        "High",
        "Regenerate sitemap.xml via helpers/_build_bilingual_tree.py (or the project's "
        "standard sitemap build step) after confirming routes.json entries.",
        "python helpers/_validate_bilingual.py passes; search engines discover all public routes.",
    ),
    Issue(
        "LINK-03",
        "Internal links",
        "Image asset filename contains a space, which can cause encoding inconsistencies "
        "across servers and tools.",
        "images/scientists-photos/rufat azizov.png; az/en/forum/2024/rector_speeches.html; "
        "helpers/_speech_photos_lib.py",
        "Medium",
        "Rename file to rufat-azizov.png; update photo map and all HTML src references in one "
        "scripted pass; run helpers/_validate_site.py.",
        "Stable image URLs; no broken photos on rector speeches pages.",
    ),
    Issue(
        "LINK-04",
        "Internal links",
        "Automated path validation reports no broken local asset references across 23,003 "
        "checks on 469 HTML pages (positive baseline).",
        "Sitewide az/, en/, css/, js/, images/",
        "Informational",
        "Maintain current practice: run python helpers/_validate_site.py after every HTML/CSS/JS "
        "path change.",
        "Continued zero broken-path regressions before deploy.",
    ),
    # --- Navigation ---
    Issue(
        "NAV-01",
        "Navigation",
        "Primary navigation is defined in i18n/nav.json and rebuilt by daab-primary-nav.js, but "
        "~400+ pages embed a static HTML nav snapshot. Snapshots are often stale (e.g. flat "
        "sponsors link while nav.json defines a sponsorship dropdown including Donate).",
        "i18n/nav.json; js/daab-primary-nav.js; embedded <nav> in az/en/**/*.html "
        "(data-daab-nav-placeholder=\"1\")",
        "Medium",
        "After nav.json changes, regenerate page shells via the site-wide HTML sync helper; "
        "long-term: use empty nav placeholder only and rely on JS (or server include) for nav HTML.",
        "Fallback nav matches JSON when JS is slow; no flash of outdated menu structure.",
    ),
    Issue(
        "NAV-02",
        "Navigation",
        "Scientists catalogue (list + profiles) is not in the primary top navigation bar; "
        "discovery relies on home page cards and Forum section navigation.",
        "i18n/nav.json (primary array); az/scientists/list.html; az/scientists/profiles.html",
        "Medium",
        "Product decision: add a scientists group to nav.json primary, or strengthen home/forum "
        "funnels and site search; regenerate nav snapshots if primary nav changes.",
        "Scientists section is discoverable according to agreed information architecture.",
    ),
    Issue(
        "NAV-03",
        "Navigation",
        "Donate page exists in routes and nav.json sponsorship group but is absent from many "
        "embedded nav fallbacks that only show a flat link to sponsors.html.",
        "i18n/nav.json; az/donate.html; en/donate.html; embedded nav on profile/forum pages",
        "Low",
        "Regenerate embedded nav from nav.json; verify daab-primary-nav.js mounts on all "
        "data-daab-nav-mount pages including sponsors after refactor.",
        "Donate is reachable from sponsorship dropdown on every standard page.",
    ),
    # --- HTML structure ---
    Issue(
        "HTML-01",
        "HTML structure",
        "sponsors.html uses a bespoke HTML document with large inline <style> blocks instead of "
        "the shared site shell used by donate.html and other main pages.",
        "az/sponsors.html; en/sponsors.html",
        "High",
        "Refactor onto standard shell: daab-common.css, daab-mobile.css, data-daab-nav-mount, "
        "shared header/footer; move page-specific rules to css/daab-sponsors-page.css if needed.",
        "Sponsors pages share nav, language switcher, mobile menu, and search with the rest of the site.",
    ),
    Issue(
        "HTML-03",
        "HTML structure",
        "membership.html in az/ and en/ intentionally omit shared CSS/JS because they are "
        "meta-refresh redirect stubs to membership_value.html.",
        "az/membership.html; en/membership.html",
        "Low",
        "No change required for UX; optionally teach _validate_site.py to skip pages with "
        "data-daab-legacy-redirect to reduce noise.",
        "Validator output reflects intentional stubs without false alarms.",
    ),
    Issue(
        "HTML-04",
        "HTML structure",
        "Root index.html language gateway uses older cache-bust query versions than inner pages.",
        "index.html (daab-common.css?v=56, daab-i18n.js?v=18 vs v=64/v=21 on az/index.html)",
        "Low",
        "Bump ?v= on gateway assets via helpers/_site_wide_cleanup.py when shared files change.",
        "Gateway page loads current CSS/JS after deployments.",
    ),
    # --- CSS and styling ---
    Issue(
        "CSS-01",
        "CSS and styling",
        "sponsors.html defines its own :root design tokens (DM Sans, custom blues/golds) "
        "parallel to daab-common.css / design tokens, causing visual inconsistency.",
        "az/sponsors.html; en/sponsors.html (inline styles)",
        "High",
        "Replace inline token block with shared css/daab-common.css and css/daab-design-tokens.js "
        "alignment; retain only sponsors-specific layout in a dedicated stylesheet.",
        "Sponsors page typography and colours match DAAB design system.",
    ),
    Issue(
        "CSS-02",
        "CSS and styling",
        "No conflicting duplicate stylesheets at repository root (daab-common.css correctly lives "
        "under css/ only) — positive baseline.",
        "css/ directory",
        "Informational",
        "Continue enforcing file-organization rules; do not recreate root-level CSS duplicates.",
        "Maintainable single source of truth for shared styles.",
    ),
    # --- JavaScript functionality ---
    Issue(
        "JS-01",
        "JavaScript functionality",
        "Mobile filter toolbar still references filterGroup for encyclopedia page, but HTML and "
        "catalog JS were renamed to filterCategory. Category filter is not counted in the mobile "
        "filter badge.",
        "js/daab-scientists-toolbar-mobile.js (line ~38); az/en/encyclopedia.html",
        "High",
        "Change filterGroup to filterCategory in filterSelectIds() for pageId === encyclopedia; "
        "bump ?v= on affected pages.",
        "Mobile filter badge correctly reflects active category filter.",
    ),
    Issue(
        "JS-02",
        "JavaScript functionality",
        "deployment/ package ships outdated prominent-figures-catalog.js (v=8, filterGroup) and "
        "catalog data (group/groupLabel) while source repo uses v=11 and category/categoryLabel.",
        "deployment/az/encyclopedia.html; deployment/en/encyclopedia.html; "
        "deployment/js/prominent-figures-catalog*.js",
        "Critical",
        "Rebuild deployment after every release: python helpers/_build_deployment_folder.py "
        "(add --include-images when images change).",
        "Production package matches source; encyclopedia filters and table sort work in deployment.",
    ),
    Issue(
        "JS-03",
        "JavaScript functionality",
        "Prominent figures catalog, table view, period inference, and AZ/EN link parity validated "
        "successfully in source (201 profiles, 0 missing links).",
        "js/prominent-figures-catalog.js; js/prominent-figures-catalog-data.js; "
        "helpers/_check_catalog_links.py",
        "Informational",
        "Keep catalog rebuild in workflow: python helpers/_build_prominent_figures_catalog.py "
        "after profile HTML changes.",
        "Catalog stays synchronized with profile pages.",
    ),
    Issue(
        "JS-04",
        "JavaScript functionality",
        "Scientists name order and catalog data consistency validated (83 cards, 0 mismatches).",
        "az/scientists/profiles.html; js/scientists-catalog-data.js; helpers/_check_name_order.py",
        "Informational",
        "Re-run _check_name_order.py when scientists data changes.",
        "Displayed names remain consistent across list and card views.",
    ),
    # --- Multilingual consistency ---
    Issue(
        "I18N-01",
        "Multilingual consistency",
        "Encyclopedia page branding mixes Ensiklopediya with nav label Xəziné / Görkəmli "
        "şəxsiyyətlər (Treasury / Prominent Figures).",
        "az/encyclopedia.html; en/encyclopedia.html; <title>, hero <h1>, toolbar aria-label",
        "Medium",
        "Align titles and hero copy with Treasury terminology; use Ensiklopediya only as "
        "subtitle if desired; update meta descriptions consistently in AZ and EN.",
        "User-facing labels match navigation and i18n/ui.json treasury keys.",
    ),
    Issue(
        "I18N-02",
        "Multilingual consistency",
        "i18n/search-index.json encyclopedia summaries still refer to sorting by group after "
        "UI rename to category.",
        "i18n/search-index.json",
        "Low",
        "Rebuild search index; replace group with category in encyclopedia summary strings.",
        "Site search snippets match current filter terminology.",
    ),
    Issue(
        "I18N-03",
        "Multilingual consistency",
        "i18n/ui.json contains duplicate label keys (encyclopedia and prominentFigures with "
        "identical text), increasing maintenance risk.",
        "i18n/ui.json; js/daab-primary-nav.js PAGE_LABEL_KEYS",
        "Low",
        "Consolidate to one canonical key; update PAGE_LABEL_KEYS and any HTML data-nav-id "
        "references.",
        "Single source for prominent-figures nav labels.",
    ),
    Issue(
        "I18N-04",
        "Multilingual consistency",
        "Strict EN profile scan reports ~78 issues: untranslated AZ paragraphs, mixed AZ field "
        "labels (sənət, müqavimət), AZ date strings (XV əsr, təx.), and quotes in EN profile pages.",
        "en/prominent_figures/**/*.html (e.g. biruni.html, ebulfez_elchibey.html, musa_celil.html); "
        "helpers/_build_en_prominent_figures.py; helpers/_prominent_figure_en_strings.py",
        "High",
        "Triage strict-scan output; fix EN build templates and string maps; rebuild EN profiles; "
        "whitelist acceptable proper nouns (Türkiye) in scanner; add _scan_en_profile_strict.py "
        "to pre-deploy gate.",
        "EN profiles read as fully translated; strict scan below agreed threshold.",
    ),
    Issue(
        "I18N-05",
        "Multilingual consistency",
        "Heuristic untranslated-prose scan reports 0 issues, but strict scan finds substantive "
        "gaps — weak heuristic alone is insufficient.",
        "helpers/_scan_en_untranslated_prose.py vs helpers/_scan_en_profile_strict.py",
        "Medium",
        "Use strict scan in CI; optionally tighten heuristic rules for AZ character runs in EN body.",
        "Translation regressions caught before deploy.",
    ),
    Issue(
        "I18N-06",
        "Multilingual consistency",
        "AZ home page hero subtitle contains typo: alimlərlin should be alimlərin.",
        "az/index.html (page-hero-subtitle)",
        "Low",
        "Correct Azerbaijani copy; verify EN equivalent if translated from same source.",
        "Professional AZ copy on the primary landing page.",
    ),
    Issue(
        "I18N-07",
        "Multilingual consistency",
        "python helpers/_validate_bilingual.py fails until sitemap is regenerated; otherwise "
        "page pairs and required i18n scripts are present on audited routes.",
        "Sitewide az/en page pairs; sitemap.xml",
        "High",
        "Fix sitemap (LINK-02); re-run _validate_bilingual.py until pass.",
        "Bilingual validation passes as part of release checklist.",
    ),
    # --- Responsiveness ---
    Issue(
        "RESP-01",
        "Responsiveness",
        "sponsors.html lacks viewport-fit=cover used on standard pages for notched mobile devices.",
        "az/sponsors.html; en/sponsors.html",
        "Low",
        "Add viewport-fit=cover to meta viewport on sponsors pages during shell refactor.",
        "Consistent safe-area behaviour on iOS/Android notched screens.",
    ),
    Issue(
        "RESP-02",
        "Responsiveness",
        "sponsors.html does not load daab-mobile.css or daab-mobile.js, so shared mobile nav "
        "and touch optimisations do not apply.",
        "az/sponsors.html; en/sponsors.html",
        "High",
        "Include daab-mobile.css and daab-mobile.js as part of HTML-01 shell integration.",
        "Mobile navigation and touch targets work on sponsors pages.",
    ),
    Issue(
        "RESP-03",
        "Responsiveness",
        "Standard pages (az/en) include viewport meta and daab-mobile "
        "layer — positive baseline.",
        "az/**/*.html; en/**/*.html (excluding redirect stubs)",
        "Informational",
        "Preserve viewport and mobile includes on all new pages via page templates/build helpers.",
        "New pages remain mobile-ready by default.",
    ),
    # --- Accessibility ---
    Issue(
        "A11Y-01",
        "Accessibility",
        "Scientists profile QR code images use empty alt attributes (83 instances).",
        "az/scientists/profiles.html; en/scientists/profiles.html",
        "Medium",
        "Set alt to descriptive text e.g. QR code for [scientist name] via build step or catalog data.",
        "Screen readers announce purpose of QR images.",
    ),
    Issue(
        "A11Y-02",
        "Accessibility",
        "Forum rector speeches table-of-contents thumbnails use alt=\"\" (12 per locale); names "
        "appear in adjacent link text (decorative pattern acceptable but improvable).",
        "az/en/forum/2024/rector_speeches.html",
        "Low",
        "Optionally set alt to rector name for consistency, or keep empty with documented "
        "decorative pattern.",
        "Improved screen-reader clarity for portrait thumbnails.",
    ),
    Issue(
        "A11Y-03",
        "Accessibility",
        "Standard pages include skip-to-content links and nav ARIA labels — positive baseline.",
        "Sitewide main pages",
        "Informational",
        "Maintain skip links and aria-label on nav when adding new page templates.",
        "Keyboard users can bypass repetitive navigation.",
    ),
    # --- Performance ---
    Issue(
        "PERF-01",
        "Performance",
        "deployment/ images/ may be empty or stale because _build_deployment_folder.py does not "
        "copy images/ by default (preserves existing deployment/images/).",
        "helpers/_build_deployment_folder.py; deployment/images/",
        "Medium",
        "Run build with --include-images when image assets change; document in deploy checklist.",
        "Fresh deploys include required images without manual copy steps.",
    ),
    Issue(
        "PERF-02",
        "Performance",
        "Cache-bust ?v= parameters on shared CSS/JS are inconsistent between gateway, inner "
        "pages, and deployment package.",
        "index.html; az/en/**/*.html; deployment/**",
        "Medium",
        "Use helpers/_site_wide_cleanup.py after shared asset edits; rebuild deployment.",
        "Browsers fetch updated assets after deploy; fewer stale-cache support issues.",
    ),
    Issue(
        "PERF-03",
        "Performance",
        "Encyclopedia uses deferred scripts, batched card rendering, and lazy search index "
        "loading — positive baseline from recent optimisations.",
        "js/prominent-figures-catalog.js; js/daab-search.js; az/en/encyclopedia.html",
        "Informational",
        "Preserve defer attributes and batching when extending catalog features.",
        "Large catalog remains responsive on mid-range devices.",
    ),
    # --- Security ---
    Issue(
        "SEC-01",
        "Security",
        "Static site with no server-side form handlers in repo; membership application appears "
        "client-side / embed — no exposed secrets found in HTML/JS during review.",
        "az/application.html; repository root",
        "Informational",
        "Ensure production forms post only to approved endpoints; never commit API keys or "
        ".env files; keep helpers/ and documents/ out of deploy package.",
        "No credential leakage via public static hosting.",
    ),
    Issue(
        "SEC-02",
        "Security",
        "External links using target=\"_blank\" should include rel=\"noopener\" — spot-check "
        "found rel=\"noopener\" on footer external links (positive pattern).",
        "Footer blocks on main pages",
        "Informational",
        "Audit new external links for rel=\"noopener noreferrer\" when added.",
        "Tab-nabbing risk mitigated on external links.",
    ),
    # --- Content quality ---
    Issue(
        "CONTENT-01",
        "Content quality",
        "Treasury sub-pages industrial_revolutions and major_scientific_inventions are "
        "intentional placeholder content (coming soon) but are linked from live navigation.",
        "az/industrial_revolutions.html; en/industrial_revolutions.html; "
        "az/major_scientific_inventions.html; en/major_scientific_inventions.html",
        "Low",
        "Keep placeholders until content ready; ensure sitemap and meta descriptions set "
        "expectations; consider noindex until substantive content exists.",
        "Users understand pages are in preparation; SEO reflects placeholder status if desired.",
    ),
    Issue(
        "CONTENT-02",
        "Content quality",
        "hazirlanir.html is a profile preparation stub, correctly excluded from catalog data.",
        "az/prominent_figures/azturk/hazirlanir.html; en/prominent_figures/azturk/hazirlanir.html",
        "Low",
        "Do not link from catalog until profile is complete; remove stub when replaced by real profile.",
        "No dead-end profiles surfaced in encyclopedia search/filter.",
    ),
    Issue(
        "CONTENT-03",
        "Content quality",
        "Section anchor validation passes for activities, charter, foundation, and mission "
        "page pairs (30 pairs checked).",
        "az/en activities, charter, foundation, mission pages",
        "Informational",
        "Re-run helpers/_validate_section_anchors.py when section IDs change.",
        "In-page section links continue to resolve correctly.",
    ),
    # --- UI/UX consistency ---
    Issue(
        "UX-01",
        "UI/UX consistency",
        "Sponsorship area presents inconsistently: JSON nav uses dropdown (Sponsors + Donate) "
        "while many pages show only a single flat sponsors link in embedded fallback nav.",
        "Sitewide nav; az/sponsors.html; az/donate.html",
        "Medium",
        "Complete NAV-01 nav regeneration; integrate sponsors page into shared chrome (HTML-01).",
        "Unified sponsorship UX across locales and pages.",
    ),
    Issue(
        "UX-02",
        "UI/UX consistency",
        "Encyclopedia card/table view toggle and column resize are implemented in source but "
        "not in stale deployment package.",
        "az/en/encyclopedia.html; css/daab-encyclopedia-page.css; js/daab-table-resize.js",
        "Critical",
        "Rebuild deployment (JS-02); verify toggle persistence (sessionStorage daab-encyclopedia-view).",
        "Users on production see card/table toggle and resizable columns.",
    ),
    Issue(
        "UX-03",
        "UI/UX consistency",
        "Category filter label renamed from group to category in source UI but deployment and "
        "search index still use old terminology.",
        "See JS-02, I18N-02",
        "Medium",
        "Apply JS-02 and I18N-02 together in one release.",
        "Filters, table headers, and search use category consistently.",
    ),
]

CATEGORY_ORDER = [
    "Internal links",
    "Navigation",
    "HTML structure",
    "CSS and styling",
    "JavaScript functionality",
    "Multilingual consistency",
    "Responsiveness",
    "Accessibility",
    "Performance",
    "Security",
    "Content quality",
    "UI/UX consistency",
]


def shade_cell(cell, fill_hex: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: int = 9,
    color: RGBColor | None = None,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run2._r.append(instr)
    run3 = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)
    run4 = paragraph.add_run("1")
    run5 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)


def severity_summary() -> dict[str, int]:
    counts: dict[str, int] = {}
    for issue in ISSUES:
        counts[issue.severity] = counts.get(issue.severity, 0) + 1
    return counts


def add_issue_table(doc: Document, issues: list[Issue]) -> None:
    headers = [
        "ID",
        "Description",
        "Affected files / components",
        "Severity",
        "Recommended fix",
        "Expected result",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(table.rows[0].cells[i], HEADER_FILL)

    for idx, issue in enumerate(issues):
        row = table.add_row().cells
        if idx % 2 == 1:
            for cell in row:
                shade_cell(cell, ALT_FILL)
        set_cell_text(row[0], issue.issue_id, bold=True, size=8)
        set_cell_text(row[1], issue.description, size=8)
        set_cell_text(row[2], issue.affected, size=8)
        sev_color = SEVERITY_COLORS.get(issue.severity, INK)
        set_cell_text(row[3], issue.severity, bold=True, size=8, color=sev_color)
        set_cell_text(row[4], issue.recommended_fix, size=8)
        set_cell_text(row[5], issue.expected_result, size=8)

    # Column widths (approximate)
    widths = [Cm(1.6), Cm(4.2), Cm(3.5), Cm(1.8), Cm(4.0), Cm(3.5)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w


def build_document() -> None:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Header / footer
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ""
    hr = hp.add_run("DAAB / WAAS — Codebase Review Report (technical)")
    hr.font.size = Pt(9)
    hr.font.color.rgb = MUTED

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    fr = fp.add_run(f"Issues identified: {REVIEW_IDENTIFIED_AT}   |   Page ")
    fr.font.size = Pt(9)
    fr.font.color.rgb = MUTED
    add_page_number(fp)

    # Title block
    title = doc.add_heading("DAAB / WAAS Website", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading("Full Codebase Review — Issues & Remediation Strategies", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Issues identified: {REVIEW_IDENTIFIED_AT}\n"
        f"Review period: {REVIEW_PERIOD}\n"
        f"Repository: DAAB-WAAS static site (az/, en/, css/, js/, images/)\n"
        f"Document generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    mr.font.color.rgb = MUTED
    mr.font.size = Pt(10)

    doc.add_page_break()

    # Executive summary
    doc.add_heading("Executive summary", level=1)
    counts = severity_summary()
    doc.add_paragraph(
        f"This report documents {len(ISSUES)} findings from an automated and manual review of "
        f"the DAAB/WAAS bilingual static website. Validators checked 469 HTML pages and "
        f"23,003 local asset references. No broken local paths were detected in source. "
        f"Critical gaps centre on stale deployment artifacts, sitemap coverage, EN profile "
        f"translation leftovers, and inconsistent page chrome on sponsors.html."
    )
    doc.add_paragraph(
        "Severity breakdown: "
        + "; ".join(f"{k}: {counts[k]}" for k in sorted(counts.keys(), key=lambda s: (
            ["Critical", "High", "Medium", "Low", "Informational"].index(s)
            if s in ["Critical", "High", "Medium", "Low", "Informational"] else 99
        )))
    )

    doc.add_heading("Validation tools executed", level=2)
    for item in [
        "helpers/_validate_site.py — pass (8 intentional warnings on membership redirect stubs)",
        "helpers/_validate_bilingual.py — fail (sitemap missing Treasury placeholder pages)",
        "helpers/_check_catalog_links.py — pass (201/201 prominent figure hrefs)",
        "helpers/_check_name_order.py — pass (83 scientists, 0 mismatches)",
        "helpers/_validate_section_anchors.py — pass (30 page pairs)",
        "helpers/_validate_cv_cards.py — pass",
        "helpers/_scan_en_untranslated_prose.py — pass (heuristic; weak alone)",
        "helpers/_scan_en_profile_strict.py — ~78 issues in EN profiles",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Recommended remediation phases", level=2)
    phases = [
        ("Phase 0 — Deploy safety (immediate)", "Fix JS-01 filterCategory bug; rebuild deployment (JS-02); smoke-test encyclopedia."),
        ("Phase 1 — SEO & routing", "Regenerate sitemap (LINK-02); pass _validate_bilingual.py."),
        ("Phase 2 — Terminology & copy", "I18N-01, I18N-06; regenerate nav snapshots (NAV-01); rebuild search index (I18N-02)."),
        ("Phase 3 — EN profiles", "I18N-04 triage and rebuild via _build_en_prominent_figures.py; strict scan gate."),
        ("Phase 4 — Page chrome", "HTML-01/CSS-01/RESP-02 sponsors integration."),
        ("Phase 5 — Polish", "LINK-03 image rename; A11Y-01 QR alt text; PERF-01 images in deploy."),
        ("Phase 6 — IA decision", "NAV-02 scientists in primary nav (product decision)."),
    ]
    pt = doc.add_table(rows=1, cols=2)
    pt.style = "Table Grid"
    set_cell_text(pt.rows[0].cells[0], "Phase", bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
    set_cell_text(pt.rows[0].cells[1], "Scope", bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
    shade_cell(pt.rows[0].cells[0], HEADER_FILL)
    shade_cell(pt.rows[0].cells[1], HEADER_FILL)
    for phase, scope in phases:
        cells = pt.add_row().cells
        set_cell_text(cells[0], phase, bold=True, size=9)
        set_cell_text(cells[1], scope, size=9)

    doc.add_page_break()

    # Findings by category
    doc.add_heading("Detailed findings by category", level=1)
    for category in CATEGORY_ORDER:
        cat_issues = [i for i in ISSUES if i.category == category]
        if not cat_issues:
            continue
        doc.add_heading(category, level=2)
        doc.add_paragraph(f"Findings in this category: {len(cat_issues)}")
        add_issue_table(doc, cat_issues)
        doc.add_paragraph()

    doc.add_page_break()

    # Appendix
    doc.add_heading("Appendix A — Pre-deploy checklist", level=1)
    for step in [
        "python helpers/_validate_site.py",
        "python helpers/_validate_bilingual.py",
        "python helpers/_check_catalog_links.py",
        "python helpers/_scan_en_profile_strict.py (under agreed threshold)",
        "python helpers/_build_deployment_folder.py [--include-images]",
        "Local smoke test: gateway → AZ home → encyclopedia → EN profile → sponsors",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Appendix B — Issue index", level=1)
    idx_table = doc.add_table(rows=1, cols=4)
    idx_table.style = "Table Grid"
    for i, h in enumerate(["ID", "Category", "Severity", "Brief description"]):
        set_cell_text(idx_table.rows[0].cells[i], h, bold=True, size=8, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(idx_table.rows[0].cells[i], HEADER_FILL)
    for issue in ISSUES:
        cells = idx_table.add_row().cells
        set_cell_text(cells[0], issue.issue_id, size=8)
        set_cell_text(cells[1], issue.category, size=8)
        set_cell_text(cells[2], issue.severity, size=8, color=SEVERITY_COLORS.get(issue.severity))
        brief = issue.description[:120] + ("…" if len(issue.description) > 120 else "")
        set_cell_text(cells[3], brief, size=8)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"Wrote {OUT_DOCX.relative_to(ROOT)}")
    print(f"  Issues: {len(ISSUES)}")
    print(f"  Identified: {REVIEW_IDENTIFIED_AT}")


def main() -> int:
    build_document()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
