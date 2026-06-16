#!/usr/bin/env python3
"""Build az/en forum/2024/stories.html from forum_2024 story Word documents."""
from __future__ import annotations

import html
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from _embed_static_nav import forum_nav_strip  # noqa: E402
from forum_en_common import FORUM_FOOTER_EN  # noqa: E402
from forum_en_stories import STORIES_EN  # noqa: E402

DOCX_AZ = ROOT / "forum_2024" / "Forum_haqqinda_hekayələr_AZ.docx"
DOCX_EN = ROOT / "forum_2024" / "Stories_about_the_Forum_EN.docx"
DOCX_LEGACY = ROOT / "forum_2024" / "Hekayələr.docx"
OUT_AZ = ROOT / "az" / "forum" / "2024" / "stories.html"
OUT_EN = ROOT / "en" / "forum" / "2024" / "stories.html"
ASSET = "../../../"
PAGE_ID = "forum-bagli-hekayeler"
ELDAR_PHOTO_SRC = f"{ASSET}images/scientists-photos/eldar-ehedov.jpg"

SECTION_IMAGES = {
    "nur": "NUR.jpg",
    "veten-hissleri": "VƏTƏN HİSSLƏRİ.jpg",
    "cidir-duzu": "CIDIR DÜZÜ.jpg",
    "xedice": "XƏDİCƏ.jpg",
}

# Sidebar icons — matched to each story’s theme (light, homeland, plateau, piano).
SECTION_ICONS = {
    "nur": "✨",
    "veten-hissleri": "❤️",
    "cidir-duzu": "🏇",
    "xedice": "🎹",
}

_AZ_TRANS = str.maketrans(
    {
        "İ": "i",
        "I": "i",
        "ı": "i",
        "Ə": "e",
        "ə": "e",
        "Ü": "u",
        "ü": "u",
        "Ö": "o",
        "ö": "o",
        "Ş": "s",
        "ş": "s",
        "Ç": "c",
        "ç": "c",
        "Ğ": "g",
        "ğ": "g",
    }
)

AZ_HEADERS: dict[str, tuple[str, str]] = {
    "nur": ("nur", "NUR"),
    "veten hissleri": ("veten-hissleri", "VƏTƏN HİSSLƏRİ"),
    "veten duygusu": ("veten-hissleri", "VƏTƏN HİSSLƏRİ"),
    "cidir duzu": ("cidir-duzu", "CIDIR DÜZÜ"),
    "xedice": ("xedice", "XƏDİCƏ"),
}

# How many paragraphs stay beside the floated image (char budget + min/max per story).
SECTION_WRAP_TARGET_CHARS: dict[str, int] = {
    "nur": 2600,
    "veten-hissleri": 2100,
    "cidir-duzu": 3600,
    "xedice": 1700,
}
SECTION_WRAP_MIN_PARAS: dict[str, int] = {
    "nur": 3,
    "veten-hissleri": 3,
    "cidir-duzu": 4,
    "xedice": 3,
}
SECTION_WRAP_MAX_PARAS: dict[str, int] = {
    "nur": 5,
    "veten-hissleri": 5,
    "cidir-duzu": 7,
    "xedice": 5,
}

EN_HEADERS: dict[str, tuple[str, str]] = {
    "NUR": ("nur", "NUR (LIGHT)"),
    "A SENSE OF HOMELAND": ("veten-hissleri", "FEELINGS FOR THE HOMELAND"),
    "JIDIR DUZU": ("cidir-duzu", "JIDIR PLAIN"),
    "KHADIJA": ("xedice", "KHADIJA"),
}

SIDEBAR_SCRIPT = f'<script src="{ASSET}js/daab-sidebar-timeline.js?v=1" defer></script>'


def esc(s: str) -> str:
    return html.escape(s, quote=True)


def norm(text: str) -> str:
    return text.replace("\r", "").replace("\xa0", " ").strip()


def header_key_az(text: str) -> str:
    folded = unicodedata.normalize("NFKC", text.strip()).translate(_AZ_TRANS).casefold()
    return folded.replace("ğ", "g").replace("ı", "i")


def parse_az_docx(doc: Document) -> dict:
    lines = [norm(p.text) for p in doc.paragraphs if norm(p.text)]
    title = lines[0] if lines else ""
    sections: list[dict] = []
    current: dict | None = None

    for line in lines[1:]:
        hk = header_key_az(line)
        if hk in AZ_HEADERS:
            if current:
                sections.append(current)
            sid, display = AZ_HEADERS[hk]
            current = {"id": sid, "title": display, "paragraphs": []}
            continue
        if current is not None:
            current["paragraphs"].append(line)

    if current:
        sections.append(current)

    return {"title": title, "sections": sections}


def parse_en_docx(doc: Document) -> dict:
    lines = [norm(p.text) for p in doc.paragraphs if norm(p.text)]
    title = lines[0] if lines else ""
    sections: list[dict] = []
    current: dict | None = None

    for line in lines[1:]:
        key = line.strip().upper()
        if key in EN_HEADERS:
            if current:
                sections.append(current)
            sid, display = EN_HEADERS[key]
            current = {"id": sid, "title": display, "paragraphs": []}
            continue
        if current is not None:
            current["paragraphs"].append(line)

    if current:
        sections.append(current)

    return {"title": title, "sections": sections}


def parse_docx(doc: Document) -> dict:
    """Legacy parser for Hekayələr.docx (uppercase section headers)."""
    legacy_headers = {
        "NUR": ("nur", "NUR"),
        "VƏTƏN HİSSLƏRİ": ("veten-hissleri", "VƏTƏN HİSSLƏRİ"),
        "CIDIR DÜZÜ": ("cidir-duzu", "CIDIR DÜZÜ"),
        "XƏDİCƏ": ("xedice", "XƏDİCƏ"),
    }
    lines = [norm(p.text) for p in doc.paragraphs if norm(p.text)]
    title = lines[0] if lines else ""
    sections: list[dict] = []
    current: dict | None = None

    for line in lines[1:]:
        if line in legacy_headers:
            if current:
                sections.append(current)
            sid, display = legacy_headers[line]
            current = {"id": sid, "title": display, "paragraphs": []}
            continue
        if current is not None:
            current["paragraphs"].append(line)

    if current:
        sections.append(current)

    return {"title": title, "sections": sections}


def first_sentence(text: str) -> str:
    """First sentence of body copy for sidebar teasers."""
    t = text.strip()
    if not t:
        return ""
    m = re.match(r"^(.+?[.!?…])(?:\s|$)", t)
    return (m.group(1) if m else t).strip()


def story_teaser(section: dict) -> str:
    paras = section.get("paragraphs") or []
    i = 0
    while i < len(paras):
        line = paras[i]
        nxt = paras[i + 1] if i + 1 < len(paras) else None
        if is_pull_quote(line, nxt):
            i += 2
            continue
        return first_sentence(line)
    return ""


def is_pull_quote(line: str, next_line: str | None) -> bool:
    """Blockquote for a quoted passage with a separate — attribution line."""
    t = line.strip()
    if not (t.startswith('"') or t.startswith("\u201c") or t.startswith("'")):
        return False
    if re.search(r'["\u201d],\s*[-–]\s*dey', t, re.I):
        return False
    if " – deyə" in t or " - deyə" in t.lower():
        return False
    if not next_line:
        return False
    n = next_line.strip()
    return n.startswith("— ") or n.startswith("– ") or n.startswith("- ")


def paras_to_html(paragraphs: list[str]) -> str:
    parts: list[str] = []
    i = 0
    while i < len(paragraphs):
        line = paragraphs[i]
        nxt = paragraphs[i + 1] if i + 1 < len(paragraphs) else None
        if is_pull_quote(line, nxt):
            quote_text = line.strip().strip('"').strip("\u201c\u201d").strip()
            cite = nxt or ""
            i += 2
            block = f'<blockquote class="card-quote"><p>{esc(quote_text)}</p>'
            if cite:
                block += f"<footer>{esc(cite)}</footer>"
            block += "</blockquote>"
            parts.append(block)
            continue
        parts.append(f'<p class="card-text">{esc(line)}</p>')
        i += 1
    return "\n".join(parts)


def en_paras_to_html(paragraphs: list[dict]) -> str:
    parts: list[str] = []
    for item in paragraphs:
        if item["type"] == "quote":
            block = f'<blockquote class="card-quote"><p>{esc(item["text"])}</p>'
            if item.get("cite"):
                block += f'<footer>{esc(item["cite"])}</footer>'
            block += "</blockquote>"
            parts.append(block)
        else:
            parts.append(f'<p class="card-text">{esc(item["text"])}</p>')
    return "\n".join(parts)


def figure_html(section_id: str, alt: str) -> str:
    img = SECTION_IMAGES.get(section_id)
    if not img:
        return ""
    size_class = {
        "nur": "forum-story-figure--wide",
        "veten-hissleri": "forum-story-figure--compact",
        "cidir-duzu": "forum-story-figure--portrait",
        "xedice": "forum-story-figure--xedice",
    }.get(section_id, "")
    classes = " ".join(
        part
        for part in ("card-gallery", "single", "forum-story-figure", "forum-story-figure--float", size_class)
        if part
    )
    return (
        f'<figure class="{classes}">'
        f'<img src="{ASSET}images/forum/{esc(img)}" alt="{esc(alt)}" width="900" height="520" '
        f'loading="lazy" decoding="async"/>'
        f"</figure>"
    )


def split_paragraphs_for_wrap(section_id: str, paragraphs: list[str]) -> tuple[list[str], list[str]]:
    """Keep enough opening paragraphs beside the image to avoid empty side gaps."""
    if not paragraphs:
        return [], []

    min_paras = SECTION_WRAP_MIN_PARAS.get(section_id, 2)
    max_paras = SECTION_WRAP_MAX_PARAS.get(section_id, 5)
    target_chars = SECTION_WRAP_TARGET_CHARS.get(section_id, 2200)

    wrapped: list[str] = []
    char_total = 0
    para_count = 0
    i = 0

    while i < len(paragraphs):
        if para_count >= max_paras:
            break

        line = paragraphs[i]
        nxt = paragraphs[i + 1] if i + 1 < len(paragraphs) else None

        if is_pull_quote(line, nxt):
            block_len = len(line) + len(nxt or "")
            if wrapped and char_total + block_len > target_chars and para_count >= min_paras:
                break
            wrapped.append(line)
            if nxt:
                wrapped.append(nxt)
            char_total += block_len
            para_count += 1
            i += 2
            continue

        if wrapped and char_total + len(line) > target_chars and para_count >= min_paras:
            break

        wrapped.append(line)
        char_total += len(line)
        para_count += 1
        i += 1

    while i < len(paragraphs) and para_count < min_paras:
        line = paragraphs[i]
        nxt = paragraphs[i + 1] if i + 1 < len(paragraphs) else None
        if is_pull_quote(line, nxt):
            wrapped.append(line)
            if nxt:
                wrapped.append(nxt)
            char_total += len(line) + len(nxt or "")
            para_count += 1
            i += 2
            continue
        wrapped.append(line)
        char_total += len(line)
        para_count += 1
        i += 1

    return wrapped, paragraphs[i:]


def story_body_html(section: dict) -> str:
    paragraphs = list(section.get("paragraphs") or [])
    if not paragraphs:
        return figure_html(section["id"], section["title"])

    sid = section["id"]
    fig = figure_html(sid, section["title"])
    wrap_paras, rest_paras = split_paragraphs_for_wrap(sid, paragraphs)

    if not wrap_paras:
        return f"{fig}\n{paras_to_html(paragraphs)}"

    lead = f'<div class="forum-story-lead">\n{fig}\n{paras_to_html(wrap_paras)}\n</div>'
    rest_html = paras_to_html(rest_paras) if rest_paras else ""
    if rest_html:
        return f"{lead}\n{rest_html}"
    return lead


def eldar_title_photo_html(lang: str) -> str:
    alt = "Eldar Ahadov" if lang == "en" else "Eldar Əhədov"
    return (
        '<span class="story-title-photo-corner" aria-hidden="true">'
        f'<span class="story-title-photo-frame">'
        f'<img src="{ELDAR_PHOTO_SRC}" alt="{esc(alt)}" class="story-title-photo" '
        f'width="52" height="52" loading="lazy" decoding="async"/>'
        f"</span></span>"
    )


def story_card(section: dict, *, lang: str) -> str:
    title = section["title"]
    icon = SECTION_ICONS.get(section["id"], "📖")
    body = story_body_html(section)
    title_photo = eldar_title_photo_html(lang)
    return f"""
<article class="news-card forum-story-card" id="{esc(section["id"])}">
<div class="card-header">
<div class="story-card-title-block">
{title_photo}
<h2 class="card-title"><span class="card-title-icon" aria-hidden="true">{icon}</span><span class="card-title-text">{esc(title)}</span></h2>
</div>
</div>
<div class="card-body">
{body}
</div>
</article>"""


def toc_items(sections: list[dict], *, lang: str) -> str:
    parts: list[str] = []
    for sec in sections:
        teaser = story_teaser(sec)
        icon = SECTION_ICONS.get(sec["id"], "📖")
        teaser_html = (
            f'<span class="timeline-link-teaser">{esc(teaser)}</span>' if teaser else ""
        )
        parts.append(
            f'<li><a href="#{esc(sec["id"])}">'
            f'<span class="tl-icon" aria-hidden="true">{icon}</span>'
            f'<span class="timeline-link-body">'
            f'<span class="timeline-link-title">{esc(sec["title"])}</span>'
            f"{teaser_html}</span></a></li>"
        )
    return "\n".join(parts)


def page_html(data: dict, *, lang: str) -> str:
    if lang == "en":
        meta = STORIES_EN
        hero_h1 = meta["hero_h1"]
        panel_title = meta["panel_title"]
        panel_copy = meta["panel_copy"]
        breadcrumb = meta["breadcrumb"]
        sidebar_label = meta["sidebar_label"]
        sidebar_aria = meta["sidebar_aria"]
        page_title = meta["page_title"]
        meta_desc = meta["meta_description"]
        skip = "Skip to content"
        bc_home = "Home"
        bc_activities = "Activities"
        bc_forum = "Forum 2024"
        footer_brand = "World Association of Azerbaijani Scientists"
        footer_contact = "Contact"
        footer_address_title = "Address"
        footer_leadership = "Leadership"
        footer_rights = "© 2026 DAAB / WAAS — All Rights Reserved"
        bc_aria = "Breadcrumb"
        panel_aria = "Forum-related stories summary"
        doc_lead = meta["doc_title"]
    else:
        hero_h1 = "Hekayələr"
        panel_title = "Nur, vətən və yaddaş"
        panel_copy = (
            "Eldar Əhədovun şəxsi ədəbi yazıları — forum təəssüratları, Qarabağ səfəri və "
            "ömürlük xatirəyə çevrilən görüşlər."
        )
        breadcrumb = "Hekayələr"
        sidebar_label = "Hekayələr"
        sidebar_aria = "Hekayələr menyusunu aç"
        page_title = "Hekayələr — DAAB"
        meta_desc = (
            "Eldar Əhədovun Xaricdə Yaşayan Azərbaycanlı Alimlərin I Forumu haqqında ədəbi yazıları."
        )
        skip = "Məzmuna keç"
        bc_home = "Ana səhifə"
        bc_activities = "Fəaliyyətimiz"
        bc_forum = "Forum 2024"
        footer_brand = "Dünya Azərbaycanlı Alimlər Birliyi"
        footer_contact = "Əlaqə"
        footer_address_title = "Ünvan"
        footer_leadership = "Rəhbərlik"
        footer_rights = "© 2026 DAAB — Bütün hüquqlar qorunur"
        bc_aria = "Səhifə yolu"
        panel_aria = "Hekayələr haqqında qısa məlumat"
        doc_lead = data["title"]

    cards = "".join(story_card(s, lang=lang) for s in data["sections"])
    toc = toc_items(data["sections"], lang=lang)
    nav = forum_nav_strip(lang, active_nav_id="forum-2024")

    if lang == "en":
        footer_html = FORUM_FOOTER_EN
    else:
        footer_html = f"""<footer class="footer-pro">
<div class="footer-inner">
<div class="footer-brand"><h3>{esc(footer_brand)}</h3></div>
<div class="footer-grid">
<div class="footer-col"><h4 class="footer-title">{esc(footer_contact)}</h4><div class="footer-item">✉ <a href="mailto:info@daab-waas.com">info@daab-waas.com</a></div><div class="footer-item">☎ <span>+90 555 147 46 74</span></div><div class="footer-item">🌐 <a href="https://daab-waas.com" rel="noopener noreferrer" target="_blank">daab-waas.com</a></div></div>
<div class="footer-col"><h4 class="footer-title">{esc(footer_address_title)}</h4><p class="footer-address">Feneryolu Mahallesi<br/>Gazi Muhtar Paşa Sokak No:44<br/>Kadıköy, İstanbul, Türkiyə</p></div>
<div class="footer-col"><h4 class="footer-title">{esc(footer_leadership)}</h4><p class="footer-leader"><strong>Prof. Dr. Məsud Əfəndiyev</strong><br/>DAAB İdarə Heyətinin Sədri<br/>Germany — James D. Murray Distinguished Professor</p></div>
</div>
</div>
<div class="footer-bottom">{esc(footer_rights)}</div>
</footer>"""

    return f"""<!DOCTYPE html>
<html lang="{lang}" data-daab-lang="{lang}" data-daab-asset-root="{ASSET}" data-daab-page-id="{PAGE_ID}" data-daab-nav-mount="1">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0, viewport-fit=cover"/>
<title>{esc(page_title)}</title>
<meta name="description" content="{esc(meta_desc)}"/>
<link rel="preconnect" href="https://fonts.googleapis.com"/>
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin/>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&family=Playfair+Display:wght@700;800&display=swap" rel="stylesheet"/>
<link href="{ASSET}css/daab-common.css?v=44" rel="stylesheet"/>
<link href="{ASSET}css/daab-mobile.css?v=11" rel="stylesheet"/>
<link href="{ASSET}css/daab-search.css?v=4" rel="stylesheet"/>
<link href="{ASSET}css/daab-back-to-top.css?v=2" rel="stylesheet"/>
<link href="{ASSET}css/daab-lang.css?v=11" rel="stylesheet"/>
<link href="{ASSET}css/daab-nav-mega.css?v=23" rel="stylesheet"/>
<link href="{ASSET}css/daab-hero-summary.css?v=9" rel="stylesheet"/>
<link href="{ASSET}css/daab-sidebar-widget.css?v=4" rel="stylesheet"/>
<link href="{ASSET}css/daab-activities-layout.css?v=14" rel="stylesheet"/>
<link href="{ASSET}css/daab-forum-content.css?v=28" rel="stylesheet"/>
<script src="{ASSET}js/daab-mobile.js?v=5" defer></script>
<script src="{ASSET}js/daab-back-to-top.js?v=3" defer></script>
<script src="{ASSET}js/daab-i18n.js?v=18" defer></script>
<script src="{ASSET}js/daab-lang-position.js?v=7" defer></script>
<script src="{ASSET}js/daab-design-tokens.js?v=1" defer></script>
<script src="{ASSET}js/daab-nav.js?v=20" defer></script>
<script src="{ASSET}js/daab-primary-nav.js?v=17" defer></script>
<script src="{ASSET}js/daab-shell.js?v=12" defer></script>
<script src="{ASSET}js/daab-page-subtitle.js?v=2" defer></script>
<script src="{ASSET}js/daab-search.js?v=7" defer></script>
</head>
<body>
<a class="skip" href="#content">{esc(skip)}</a>
{nav}
<div class="breadcrumbs forum-breadcrumbs" role="navigation" aria-label="{esc(bc_aria)}">
<a href="../../index.html">{esc(bc_home)}</a><span aria-hidden="true">›</span><a href="../../activities.html">{esc(bc_activities)}</a><span aria-hidden="true">›</span><a href="index.html">{esc(bc_forum)}</a><span aria-hidden="true">›</span><span class="forum-breadcrumbs-current" aria-current="page">{esc(breadcrumb)}</span>
</div>
<header class="page-hero">
<div class="hero-wrap shell">
<section class="hero-copy">
<h1>{hero_h1}</h1>
<p class="hero-doc-lead">{esc(doc_lead)}</p>
</section>
<aside aria-label="{esc(panel_aria)}" class="hero-panel">
<div class="panel-card">
<h2 class="panel-title">{esc(panel_title)}</h2>
<div class="panel-copy">{esc(panel_copy)}</div>
</div>
</aside>
</div>
</header>
<div class="content-wrap">
<aside class="sidebar">
<div class="sidebar-widget">
<div class="widget-head"><span>{sidebar_label}</span><button aria-controls="hekayelerTOC" aria-expanded="false" aria-label="{esc(sidebar_aria)}" class="events-menu-toggle" type="button"><span></span><span></span><span></span></button></div>
<div class="widget-body">
<ul class="timeline-list" id="hekayelerTOC">
{toc}
</ul>
</div>
</div>
</aside>
<main class="news-feed main" id="content">
{cards}
</main>
</div>
{footer_html}
{SIDEBAR_SCRIPT}
</body>
</html>
"""


def main() -> None:
    az_path = DOCX_AZ if DOCX_AZ.is_file() else DOCX_LEGACY
    en_path = DOCX_EN if DOCX_EN.is_file() else DOCX_LEGACY
    az_data = parse_az_docx(Document(str(az_path))) if az_path == DOCX_AZ else parse_docx(Document(str(az_path)))
    en_data = parse_en_docx(Document(str(en_path))) if en_path == DOCX_EN else parse_docx(Document(str(en_path)))
    OUT_AZ.parent.mkdir(parents=True, exist_ok=True)
    OUT_AZ.write_text(page_html(az_data, lang="az"), encoding="utf-8", newline="\n")
    OUT_EN.write_text(page_html(en_data, lang="en"), encoding="utf-8", newline="\n")
    print(f"wrote {OUT_AZ.relative_to(ROOT)} ({len(az_data['sections'])} stories)")
    print(f"wrote {OUT_EN.relative_to(ROOT)} ({len(en_data['sections'])} stories)")


if __name__ == "__main__":
    main()
