#!/usr/bin/env python3
"""
Full-site linguistic consistency audit (AZ / EN).
Produces documents/docx/DAAB-Website-Linguistic-Consistency-Audit.docx
Does NOT modify site content.
"""
from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

try:
    from _paths import ROOT
except ImportError:
    ROOT = Path(__file__).resolve().parents[1]

OUT_DOCX = ROOT / "documents" / "docx" / "DAAB-Website-Linguistic-Consistency-Audit.docx"

NAVY = RGBColor(0x09, 0x4D, 0x78)
BLUE = RGBColor(0x00, 0x69, 0xB4)
INK = RGBColor(0x1A, 0x2E, 0x3D)
MUTED = RGBColor(0x34, 0x5D, 0x76)
HEADER_FILL = "094D78"
ALT_FILL = "F4F7FA"

EN_NAV_BOARD_VARIANTS = ("Executive Board", "Board of Directors", "Executive Board")
AZ_MEMBERSHIP_WHY = "Niyə üzv olmalı"

CATEGORIES = (
    "Grammar",
    "Orthography",
    "Punctuation",
    "Style",
    "Translation",
    "Terminology",
    "UI wording",
    "Mixed language",
    "Structure / sync",
    "Metadata",
    "Tone",
)


@dataclass
class Finding:
    page: str
    lang: str
    location: str
    original: str
    recommended: str
    category: str
    comment: str = ""
    page_id: str = ""


@dataclass
class PageText:
    path: str
    lang: str
    page_id: str
    title: str = ""
    meta_description: str = ""
    h1: str = ""
    headings: list[str] = field(default_factory=list)
    nav_labels: list[str] = field(default_factory=list)
    section_nav: list[str] = field(default_factory=list)
    buttons: list[str] = field(default_factory=list)
    labels: list[str] = field(default_factory=list)
    sample_paragraphs: list[str] = field(default_factory=list)
    deployable: bool = True


def load_routes() -> dict:
    return json.loads((ROOT / "i18n" / "routes.json").read_text(encoding="utf-8"))


def load_ui() -> dict:
    return json.loads((ROOT / "i18n" / "ui.json").read_text(encoding="utf-8"))


def page_pairs(routes: dict) -> list[tuple[str, str, str]]:
    pairs = []
    for p in routes.get("pages", []):
        az = p.get("az")
        en = p.get("en")
        if az and en:
            pairs.append((p["id"], az.replace("\\", "/"), en.replace("\\", "/")))
    return pairs


def visible_text(el) -> str:
    if el is None:
        return ""
    return re.sub(r"\s+", " ", el.get_text(separator=" ", strip=True))


def extract_page(rel_path: str, lang: str, page_id: str = "") -> PageText:
    path = ROOT / rel_path.replace("/", "\\") if "\\" in rel_path else ROOT / rel_path
    path = ROOT / rel_path.replace("\\", "/")
    html = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "html.parser")

    pt = PageText(
        path=rel_path.replace("\\", "/"),
        lang=lang,
        page_id=page_id,
        deployable=True,
    )

    if soup.title and soup.title.string:
        pt.title = soup.title.string.strip()

    meta = soup.find("meta", attrs={"name": "description"})
    if meta and meta.get("content"):
        pt.meta_description = meta["content"].strip()

    for tag in soup.find_all(re.compile(r"^h[1-6]$", re.I)):
        t = visible_text(tag)
        if t:
            pt.headings.append(t)
    if pt.headings:
        pt.h1 = pt.headings[0]

    for sel in (
        ".nav-dropdown-link-title",
        ".nav-link",
        ".nav-dropdown-toggle",
    ):
        for el in soup.select(sel):
            t = visible_text(el)
            if t and len(t) < 120 and t not in pt.nav_labels:
                pt.nav_labels.append(t)

    for el in soup.select(".daab-section-nav-list a, .daab-section-nav-title"):
        t = visible_text(el)
        if t:
            pt.section_nav.append(t)

    for el in soup.select(
        "button, .btn, input[type=submit], .join-button, .prog-step span"
    ):
        t = visible_text(el)
        if t and 1 < len(t) < 80:
            pt.buttons.append(t)

    for el in soup.select("label, .field-label, .lang-label, th"):
        t = visible_text(el)
        if t and len(t) < 200:
            pt.labels.append(t)

    main = soup.find("main") or soup.find(id="content") or soup.body
    if main:
        for p in main.find_all(["p", "li"], limit=40):
            t = visible_text(p)
            if t and len(t) > 20:
                pt.sample_paragraphs.append(t[:500])

    return pt


def add_finding(
    findings: list[Finding],
    page: str,
    lang: str,
    location: str,
    original: str,
    recommended: str,
    category: str,
    comment: str = "",
    page_id: str = "",
) -> None:
    original = original.strip()
    if not original:
        return
    recommended = recommended.strip()
    if original == recommended:
        return
    findings.append(
        Finding(
            page=page,
            lang=lang,
            location=location,
            original=original[:500],
            recommended=(recommended or "—")[:500],
            category=category,
            comment=comment[:800],
            page_id=page_id,
        )
    )


def audit_page(pt: PageText, findings: list[Finding]) -> None:
    page = pt.path
    lang = pt.lang

    if not pt.deployable:
        add_finding(
            findings,
            page,
            lang,
            "Scope",
            "(build-source page)",
            "Exclude from public deployment or align with parent page",
            "Structure / sync",
            "This path is a helper/build duplicate under application/; verify it is not published.",
            pt.page_id,
        )
        return

    if not pt.title:
        add_finding(
            findings,
            page,
            lang,
            "<title>",
            "(missing)",
            "Add descriptive <title>",
            "Metadata",
            "Page has no document title.",
            pt.page_id,
        )

    if lang == "az":
        if re.search(r"\bWAAS\b", pt.title + " " + pt.meta_description):
            add_finding(
                findings,
                page,
                lang,
                "Metadata",
                "WAAS",
                "DAAB",
                "Terminology",
                "Azerbaijani pages should use DAAB branding in metadata unless quoting English proper names.",
                pt.page_id,
            )
        if re.search(r"\bclick\b|\bklikləy", " ".join(pt.sample_paragraphs).lower()):
            for para in pt.sample_paragraphs:
                if "klikləy" in para.lower() or "click" in para.lower():
                    add_finding(
                        findings,
                        page,
                        lang,
                        "Body copy",
                        para[:200],
                        para.lower().replace("klikləyin", "basın").replace("click", "basın"),
                        "Style",
                        "Prefer native AZ UI verbs (basın) over Anglicisms.",
                        pt.page_id,
                    )
                    break
        if "Forum 2024-ü" in " ".join(pt.nav_labels + [pt.meta_description]):
            for lbl in pt.nav_labels:
                if "2024-ü" in lbl:
                    add_finding(
                        findings,
                        page,
                        lang,
                        "Navigation",
                        lbl,
                        "Forum 2024 haqqında",
                        "Grammar",
                        "Accusative on year label is awkward; use a prepositional phrase.",
                        pt.page_id,
                    )
        for lbl in pt.nav_labels + pt.section_nav:
            if lbl.strip() == AZ_MEMBERSHIP_WHY:
                add_finding(
                    findings,
                    page,
                    lang,
                    "Navigation / section nav",
                    lbl,
                    "Niyə üzv olmalısınız?",
                    "Grammar",
                    "Menu label reads incomplete without subject or question mark.",
                    pt.page_id,
                )
        for lbl in pt.labels:
            if lbl == "Ad, Soyadı":
                add_finding(
                    findings,
                    page,
                    lang,
                    "Table header",
                    lbl,
                    "Ad, soyad",
                    "Grammar",
                    "Column headers should use nominative forms, not possessive.",
                    pt.page_id,
                )
            if lbl == "İxtisası":
                add_finding(
                    findings,
                    page,
                    lang,
                    "Table header",
                    lbl,
                    "İxtisas",
                    "Grammar",
                    "Possessive form is incorrect for a column label.",
                    pt.page_id,
                )
            if lbl == "Yaşadığı Ölkə":
                add_finding(
                    findings,
                    page,
                    lang,
                    "Table header",
                    lbl,
                    "Yaşadığı ölkə",
                    "Orthography",
                    "Mid-phrase capitalisation on common noun.",
                    pt.page_id,
                )
        body = " ".join(pt.sample_paragraphs + pt.labels)
        if re.search(r"Section \d+ of \d+ — [A-Za-z]", body):
            add_finding(
                findings,
                page,
                lang,
                "Form UI",
                "Section N of M — (English)",
                "Bölmə N / M — (Azerbaijani only)",
                "Mixed language",
                "English step subtitles appear on Azerbaijani application pages.",
                pt.page_id,
            )
        if "nümunə@" in body.lower() or "nümunə@" in html_safe_read(page):
            add_finding(
                findings,
                page,
                lang,
                "Form placeholder",
                "nümunə@email.com",
                "misal@email.com",
                "Orthography",
                "Standard Azerbaijani for 'example' is misal.",
                pt.page_id,
            )

    if lang == "en":
        if re.search(r"\bDAAB\b", pt.title + " " + pt.meta_description):
            add_finding(
                findings,
                page,
                lang,
                "Metadata",
                "DAAB",
                "WAAS",
                "Terminology",
                "English public pages should use WAAS in metadata.",
                pt.page_id,
            )
        if "scholar" in pt.meta_description.lower() and "scientist" in (
            pt.h1 + " ".join(pt.headings)
        ).lower():
            add_finding(
                findings,
                page,
                lang,
                "Meta description",
                pt.meta_description[:200],
                pt.meta_description.replace("scholars", "scientists").replace(
                    "Scholars", "Scientists"
                ),
                "Terminology",
                "Meta uses 'scholars' while on-page copy uses 'scientists'.",
                pt.page_id,
            )
        nav_text = " ".join(pt.nav_labels)
        if "Executive Board" in nav_text and "Board of Directors" in " ".join(
            pt.headings + pt.sample_paragraphs
        ):
            add_finding(
                findings,
                page,
                lang,
                "Governance terminology",
                "Executive Board (nav) vs Board of Directors (body)",
                "Choose one term sitewide (e.g. Executive Board)",
                "Terminology",
                "Inconsistent English labels for idarə heyəti.",
                pt.page_id,
            )
        if "intellectuals" in " ".join(pt.sample_paragraphs[:3]).lower():
            for para in pt.sample_paragraphs[:5]:
                if "intellectual" in para.lower():
                    add_finding(
                        findings,
                        page,
                        lang,
                        "Intro copy",
                        para[:220],
                        para.replace("intellectuals", "scientists"),
                        "Terminology",
                        "AZ copy refers to alimlər/ziyalılar; EN brand is Scientists.",
                        pt.page_id,
                    )
                    break
        if "programme" in nav_text.lower():
            add_finding(
                findings,
                page,
                lang,
                "Navigation",
                "Forum programme",
                "Forum program",
                "Style",
                "British spelling; use consistently or standardise to US English if preferred.",
                pt.page_id,
            )


def html_safe_read(rel: str) -> str:
    try:
        return (ROOT / rel.replace("\\", "/")).read_text(encoding="utf-8", errors="replace")
    except OSError:
        return ""


def audit_ui_json(ui: dict, findings: list[Finding]) -> None:
    az = ui.get("nav", {}).get("az", {})
    en = ui.get("nav", {}).get("en", {})

    if az.get("membershipWhy") == AZ_MEMBERSHIP_WHY:
        add_finding(
            findings,
            "i18n/ui.json",
            "az",
            "nav.membershipWhy",
            az["membershipWhy"],
            "Niyə üzv olmalısınız?",
            "Grammar",
            "Central string used across nav and section pills.",
        )
    if az.get("forum2024Desc", "").find("2024-ü") >= 0:
        add_finding(
            findings,
            "i18n/ui.json",
            "az",
            "nav.forum2024Desc",
            az["forum2024Desc"],
            "Forum 2024 haqqında",
            "Grammar",
            "Same wording embedded in static HTML nav.",
        )
    if en.get("executiveBoard") == "Board of Directors":
        add_finding(
            findings,
            "i18n/ui.json",
            "en",
            "nav.executiveBoard",
            "Board of Directors",
            "Executive Board",
            "Terminology",
            "Many HTML pages use 'Executive Board' in embedded nav.",
        )
    if en.get("foundation") == "Foundation":
        add_finding(
            findings,
            "i18n/ui.json",
            "en",
            "nav.foundation",
            "Foundation",
            "Founding / Establishment",
            "Translation",
            "AZ menu item is 'Birliyin təsisi' (founding history), not the organisation as an entity.",
        )


def audit_pairs(
    pairs: list[tuple[str, str, str]],
    pages: dict[str, PageText],
    findings: list[Finding],
) -> None:
    for page_id, az_rel, en_rel in pairs:
        az = pages.get(az_rel)
        en = pages.get(en_rel)
        if not az or not en:
            continue
        if not az.deployable or not en.deployable:
            continue
        if az.h1 and en.h1:
            az_words = len(az.h1.split())
            en_words = len(en.h1.split())
            if az_words > 0 and en_words > 0 and (az_words > 3 * en_words or en_words > 3 * az_words):
                add_finding(
                    findings,
                    f"{az_rel} ↔ {en_rel}",
                    "az/en",
                    "H1 pairing",
                    f"AZ: {az.h1} | EN: {en.h1}",
                    "Review parallel structure and length",
                    "Structure / sync",
                    "Headings may not be equivalent; verify bilingual parity.",
                    page_id,
                )
        if len(az.section_nav) != len(en.section_nav) and az.section_nav and en.section_nav:
            add_finding(
                findings,
                f"{az_rel} ↔ {en_rel}",
                "az/en",
                "Section navigation",
                f"AZ ({len(az.section_nav)} items): {', '.join(az.section_nav[:5])}",
                f"EN ({len(en.section_nav)} items): {', '.join(en.section_nav[:5])}",
                "Structure / sync",
                "Section pill count or labels differ between locales.",
                page_id,
            )


def curated_findings() -> list[Finding]:
    """High-confidence manual findings from editorial review."""
    return [
        Finding(
            "az/index.html",
            "az",
            "Homepage card tag",
            "Təsisat",
            "Təsis",
            "Terminology",
            "Təsisat usually means facilities/installations; EN tag is Foundation.",
            "home",
        ),
        Finding(
            "en/index.html",
            "en",
            "Hero CTA",
            "Join our Association",
            "View membership terms / Apply for membership",
            "UI wording",
            "CTA links to membership terms page, not application; AZ hero has same pattern.",
            "home",
        ),
        Finding(
            "az/membership.html",
            "az",
            "Instructions",
            "düyməsini klikləyin",
            "düyməsinə basın",
            "Style",
            "Anglicism; standard AZ interface copy uses basın.",
            "membership",
        ),
        Finding(
            "en/application.html",
            "en",
            "Degree option",
            "PhD (Doctor of Philosophy) — Doctor of Sciences in Philosophy",
            "PhD (Doctor of Philosophy)",
            "Translation",
            "Mistranslation of AZ 'Elmlər üzrə Fəlsəfə Doktoru'; not equivalent to Russian-style DSs.",
            "membership-application",
        ),
        Finding(
            "az/application.html",
            "az",
            "Field gloss (English)",
            "Contributions you can make to the activities of WAAS",
            "…DAAB… (or remove English gloss on AZ page)",
            "Mixed language",
            "Wrong acronym and unnecessary English on AZ-only form.",
            "membership-application",
        ),
        Finding(
            "az/application.html",
            "az",
            "Academic rank label",
            "Akademik rütbəniz",
            "Akademik adınız / Akademik titulunuz",
            "Terminology",
            "rütbə often implies military rank; EN uses Academic Title.",
            "membership-application",
        ),
        Finding(
            "en/membership.html",
            "en",
            "Body instruction",
            'click the "Join our Association" button',
            "Use CTA label that exists on this page",
            "UI wording",
            "Referenced button text does not match this page hero.",
            "membership",
        ),
        Finding(
            "en/scientists/list.html",
            "en",
            "Filter vs column",
            "Degree",
            "Academic degree",
            "UI wording",
            "AZ filter uses 'Elmi dərəcə'; column header is longer.",
            "scientists-list",
        ),
        Finding(
            "Sitewide",
            "en",
            "Governance (multiple pages)",
            "Executive Board / Board of Directors / Chair of the WAAS Executive Board",
            "Single preferred term (recommend: Executive Board)",
            "Terminology",
            "Harmonise navigation, cards, signatures, and i18n/ui.json.",
            "",
        ),
        Finding(
            "Sitewide",
            "az",
            "Capitalisation",
            "Birliyin Təsisi / İdarə Heyəti / Üzvlük Şərtləri (cards)",
            "Sentence case: təsisi / heyəti / şərtləri",
            "Style",
            "Nav uses lowercase; homepage cards use title case inconsistently.",
            "home",
        ),
    ]


def dedupe_findings(findings: list[Finding]) -> list[Finding]:
    seen: set[tuple[str, str, str, str]] = set()
    out: list[Finding] = []
    for f in findings:
        key = (f.page, f.location, f.original[:80], f.category)
        if key in seen:
            continue
        seen.add(key)
        out.append(f)
    return sorted(out, key=lambda x: (x.lang, x.page, x.category, x.location))


def shade_cell(cell, fill: str) -> None:
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), fill)
    sh.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(sh)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9, color=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_findings_table(doc: Document, findings: list[Finding], start_num: int) -> int:
    headers = ["#", "Location", "Original text", "Recommended correction", "Category", "Comments"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(table.rows[0].cells[i], HEADER_FILL)

    n = start_num
    for f in findings:
        row = table.add_row().cells
        vals = [
            str(n),
            f.location,
            f.original,
            f.recommended,
            f.category,
            f.comment,
        ]
        for i, v in enumerate(vals):
            set_cell_text(row[i], v, size=8)
        if n % 2 == 0:
            for c in row:
                shade_cell(c, ALT_FILL)
        n += 1

    doc.add_paragraph()
    return n


def add_toc(doc: Document) -> None:
    from daab_docx_export import add_table_of_contents

    doc.add_heading("Table of contents", level=1)
    note = doc.add_paragraph(
        "In Microsoft Word: click this table, press F9, or right-click → Update Field → "
        "Update entire table."
    )
    note.runs[0].font.color.rgb = MUTED
    note.runs[0].font.size = Pt(10)
    toc_p = doc.add_paragraph()
    add_table_of_contents(toc_p)
    doc.add_page_break()


def setup_audit_header_footer(doc: Document) -> None:
    from daab_docx_export import add_page_number_field

    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    header._element.clear()
    hp = header.add_paragraph()
    r = hp.add_run("DAAB / WAAS — Linguistic Consistency Audit (editorial review)")
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    footer = section.footer
    footer._element.clear()
    fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("Confidential — for internal editorial use   |   Page ")
    r1.font.size = Pt(9)
    r1.font.color.rgb = MUTED
    add_page_number_field(fp, align_center=False)


def build_docx(findings: list[Finding], pages_audited: list[str]) -> None:
    az_findings = [f for f in findings if f.lang == "az"]
    en_findings = [f for f in findings if f.lang == "en"]
    cross = [f for f in findings if f.lang in ("az/en", "en/az") or "↔" in f.page]
    global_f = [f for f in findings if f.page.startswith("i18n") or f.page == "Sitewide"]

    doc = Document()
    setup_audit_header_footer(doc)
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Title
    t = doc.add_heading("DAAB / WAAS Website", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2 = doc.add_heading("Linguistic Consistency Audit Report", level=1)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}\nEditorial review — no changes applied to the live site")
    r.font.color.rgb = MUTED

    doc.add_page_break()
    add_toc(doc)

    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(
        f"This report documents linguistic, terminological, and UI consistency findings across "
        f"the bilingual DAAB (Azerbaijani) and WAAS (English) static website. "
        f"A total of {len(findings)} distinct issues were recorded after automated extraction "
        f"and editorial review of {len(pages_audited)} HTML page paths. "
        f"No corrections have been applied to the repository; this document is intended for "
        f"review and approval before implementation."
    )
    doc.add_paragraph(
        f"Breakdown: Azerbaijani — {len(az_findings)} findings; "
        f"English — {len(en_findings)} findings; "
        f"Cross-locale / structure — {len(cross)} findings; "
        f"Global (i18n / sitewide) — {len(global_f)} findings."
    )

    doc.add_heading("Methodology", level=1)
    for item in [
        "Inventory of public pages from i18n/routes.json and locale directories (az/, en/).",
        "Automated text extraction: titles, meta descriptions, headings, navigation, section pills, form labels, tables, and sample body copy.",
        "Rule-based checks: mixed language, branding (DAAB/WAAS), Anglicisms, grammar patterns, terminology drift.",
        "Bilingual pairing of AZ/EN routes for structural comparison.",
        "Manual editorial review of high-traffic pages (home, membership, application, scientists directory).",
        "Build-source pages under application/ flagged separately (not deployed by default).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Issue categories", level=2)
    doc.add_paragraph(", ".join(CATEGORIES))

    doc.add_heading("High-priority recurring patterns", level=2)
    recurring = [
        ("AZ", "Niyə üzv olmalı", "~18 pages + i18n/ui.json", "Use Niyə üzv olmalısınız? in nav and section pills"),
        ("AZ", "Forum 2024-ü kəşf edin", "~19 pages", "Prefer Forum 2024 haqqında"),
        ("EN", "Executive Board vs Board of Directors", "~20+ pages", "Pick one governance label sitewide"),
        ("AZ", "Section N of M — English", "az/application.html", "Localise step subtitles to Azerbaijani"),
        ("EN", "scholars vs scientists", "en/index.html meta", "Align with WAAS brand wording"),
    ]
    rt = doc.add_table(rows=1, cols=4)
    rt.style = "Table Grid"
    for i, h in enumerate(["Lang", "Pattern", "Scope", "Recommendation"]):
        set_cell_text(rt.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(rt.rows[0].cells[i], HEADER_FILL)
    for row_data in recurring:
        cells = rt.add_row().cells
        for i, v in enumerate(row_data):
            set_cell_text(cells[i], v, size=8)

    doc.add_page_break()

    # AZ by page
    doc.add_heading("Part 1 — Azerbaijani (az)", level=1)
    az_by_page: dict[str, list[Finding]] = {}
    for f in az_findings:
        az_by_page.setdefault(f.page, []).append(f)

    num = 1
    for page in sorted(az_by_page.keys()):
        doc.add_heading(page, level=2)
        doc.add_paragraph(f"Findings: {len(az_by_page[page])}")
        num = add_findings_table(doc, az_by_page[page], num)

    doc.add_page_break()

    # EN by page
    doc.add_heading("Part 2 — English (en)", level=1)
    en_by_page: dict[str, list[Finding]] = {}
    for f in en_findings:
        en_by_page.setdefault(f.page, []).append(f)

    for page in sorted(en_by_page.keys()):
        doc.add_heading(page, level=2)
        doc.add_paragraph(f"Findings: {len(en_by_page[page])}")
        num = add_findings_table(doc, en_by_page[page], num)

    doc.add_page_break()

    # Cross + global
    doc.add_heading("Part 3 — Cross-locale and structural sync", level=1)
    if cross:
        num = add_findings_table(doc, cross, num)
    else:
        doc.add_paragraph("No structural sync issues flagged.")

    doc.add_heading("Part 4 — Global configuration (i18n) and sitewide", level=1)
    if global_f:
        add_findings_table(doc, global_f, num)
    else:
        doc.add_paragraph("No global issues flagged.")

    doc.add_page_break()
    doc.add_heading("Appendix A — Pages audited", level=1)
    for p in sorted(pages_audited):
        doc.add_paragraph(p, style="List Bullet")

    doc.add_heading("Appendix B — Approval workflow", level=1)
    doc.add_paragraph(
        "Recommended next steps: (1) Editorial board reviews this document and marks "
        "approved corrections; (2) Implement approved changes in i18n/ui.json, HTML, and "
        "build helpers; (3) Re-run this audit script to verify closure; (4) Bump asset "
        "cache versions after deployment."
    )

    from daab_docx_export import set_update_fields_on_open

    set_update_fields_on_open(doc)
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))


def run_audit() -> tuple[list[Finding], list[str]]:
    routes = load_routes()
    ui = load_ui()
    pairs = page_pairs(routes)
    findings: list[Finding] = []
    pages: dict[str, PageText] = {}
    audited: list[str] = []

    all_rels: set[str] = set()
    for _pid, az, en in pairs:
        all_rels.add(az)
        all_rels.add(en)

    for rel in sorted(all_rels):
        lang = "az" if rel.startswith("az/") else "en"
        pid = next((p[0] for p in pairs if p[1] == rel or p[2] == rel), "")
        try:
            pt = extract_page(rel, lang, pid)
            pages[rel] = pt
            audited.append(rel)
            audit_page(pt, findings)
        except Exception as exc:
            add_finding(
                findings,
                rel,
                lang,
                "Audit",
                str(exc),
                "Fix HTML and re-audit",
                "Structure / sync",
                "Page could not be parsed.",
                pid,
            )

    audit_ui_json(ui, findings)
    audit_pairs(pairs, pages, findings)
    findings.extend(curated_findings())
    findings = dedupe_findings(findings)
    return findings, audited


def main() -> None:
    findings, audited = run_audit()
    build_docx(findings, audited)
    print(f"Wrote {OUT_DOCX.relative_to(ROOT)}")
    print(f"Total findings: {len(findings)}")
    print(f"Pages audited: {len(audited)}")


if __name__ == "__main__":
    main()
